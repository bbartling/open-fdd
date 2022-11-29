import operator, time
from datetime import datetime, timedelta

import pandas as pd
import numpy as np

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from docx import Document
from docx.shared import Inches


# required params taken from the screenshot above
VFD_SPEED_PERCENT_ERR_THRES = .05
VFD_SPEED_PERCENT_MAX = .99
DUCT_STATIC_INCHES_ERR_THRES = .1


def fault_condition_one(dataframe):
    return operator.and_(dataframe.duct_static < dataframe.duct_static_setpoint - dataframe.duct_static_inches_err_thres,
                         dataframe.supply_vfd_speed >= dataframe.vfd_speed_percent_max - dataframe.vfd_speed_percent_err_thres)


duct_pressure = pd.read_csv(
    './ahu_data/DA-P.csv',
    index_col='Date',
    parse_dates=True).fillna(method='ffill').dropna()
print(duct_pressure)
duct_pressure_avg = duct_pressure.rolling('5T').mean()


vfd_speed = pd.read_csv(
    './ahu_data/SF-O.csv',
    index_col='Date',
    parse_dates=True).fillna(method='ffill').dropna()
print(vfd_speed)
vfd_speed_avg = vfd_speed.rolling('5T').mean()


# combine duct pressure and fan speed datasets
df = duct_pressure_avg.join(vfd_speed_avg)
df['duct_static_setpoint'] = 1


start = df.head(1).index.date
print('Dataset start: ', start)

end = df.tail(1).index.date
print('Dataset end: ', end)


# make an entire column out of these params in the Pandas Dataframe
df['vfd_speed_percent_err_thres'] = VFD_SPEED_PERCENT_ERR_THRES
df['duct_static_inches_err_thres'] = DUCT_STATIC_INCHES_ERR_THRES
df['vfd_speed_percent_max'] = VFD_SPEED_PERCENT_MAX

df['fc1_flag'] = fault_condition_one(df)

df2 = df.copy()
df2['fc1_flag'] = df2['fc1_flag'].astype(int)

# drop params column for better plot
df2 = df2.drop(['vfd_speed_percent_err_thres',
                'duct_static_inches_err_thres',
                'vfd_speed_percent_max'], axis=1)

print(df2.columns)
print(df2.fc1_flag)


df2.plot(figsize=(25, 8), subplots=True,
         title='Duct Static Pressure and Setpoint Subplots')


plt.savefig('./static/ahu_fc1_fans_plot.png')


print("Starting ahu fc1 docx report")
document = Document()
document.add_heading('Fault Condition One Report', 0)

p = document.add_paragraph(
    'Fault condition one of ASHRAE Guideline 36 is related to flagging poor performance of a AHU variable supply fan attempting to control to a duct pressure setpoint. Fault condition equation as defined by ASHRAE:')

document.add_picture('./images/fc1_definition.png', width=Inches(6))

document.add_heading('Dataset Plot', level=2)

# ADD IN SUBPLOTS SECTION
document.add_picture('./static/ahu_fc1_fans_plot.png', width=Inches(6))

document.add_heading('Dataset Statistics', level=2)

# calculate dataset statistics
df["timedelta_alldata"] = df.index.to_series().diff()
seconds_alldata = df.timedelta_alldata.sum().seconds
days_alldata = df.timedelta_alldata.sum().days
hours_alldata = seconds_alldata//3600
minutes_alldata = (seconds_alldata//60) % 60
total_hours_calc = days_alldata * 24 + hours_alldata
df["timedelta_fddflag"] = df.index.to_series().diff().where(df["fc1_flag"] == 1)
percent_true = round(df.fc1_flag.mean() * 100, 2)
percent_false = round((1 - percent_true) * 100, 2)


# make hist plots
df['hour_of_the_day'] = df.index.hour.where(df["fc1_flag"] == 1)
df.hour_of_the_day.plot.hist(
    title='Hour Of Day When Flag is TRUE', figsize=(25, 8))
plt.savefig('./static/ahu_fc1_histogram.png')

flag_true_duct_pressure = round(
    df.duct_static.where(df["fc1_flag"] == 1).mean(), 2)


# add calcs to word doc
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time calculated in dataset: {df.timedelta_alldata.sum()}')
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time in hours calculated in dataset: {total_hours_calc}')
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time for when FDD flag is True: {df.timedelta_fddflag.sum()}')
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Percent of time in the dataset when the Fault FLAG is True: {percent_true}%')
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Percent of time in the dataset when FLAG is False: {percent_false}%')
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Average duct system pressure for when in fault condition (fan VFD speed > 95%): {flag_true_duct_pressure}"WC')
paragraph = document.add_paragraph()


# ADD in Summary Statistics of fan operation
document.add_heading('VFD Speed Statistics', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(str(df2.supply_vfd_speed.describe()))

# ADD in Summary Statistics of duct pressure
document.add_heading('Duct Pressure Statistics', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(str(df2.duct_static.describe()))

# ADD in Summary Statistics of duct pressure
document.add_heading('Duct Pressure Setpoints Statistics', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(str(df2.duct_static_setpoint.describe()))


document.add_heading('Suggestions based on data analysis', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'

if percent_true < 15:

    paragraph.add_run('The percent True of time in the dataset for when the variable fan is running at maximum speed and generating very little duct pressure is very low. This fan system appears to operate well from a perspective that the fan adiquelty meets a duct pressure setpoint.')

else:
    paragraph.add_run('The percent True of time in the dataset for when the variable fan is running at maximum speed and generating very little duct pressure is very high. This fan system appears to struggle and it could be recommended to further troubleshoot the system with a consulting engineer. A consulting engineer could be hired to redesign ventilation rates for the VAV system which can then be passed to a testing, adjusting, and balancing contractor (TAB) to implement where the TAB contractor would be responsible for any necessary mechanical adjustments for making the fan system operate to design.')


paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'

print('df.duct_static_setpoint.std: ',df.duct_static_setpoint.std())
if df.duct_static_setpoint.std() == 0:
    paragraph.add_run('The control programming doesnt appear to have a duct pressure reset strategy implemented as the standard deviation of the duct pressure setpoint data equals zero. It would be recommended to hire a consulting engineer to properly design, oversee, and validate a duct pressure reset strategy implemented by a controls contractor. A duct pressure reset can potentially save fan electrical energy consumption.')

else:
    paragraph.add_run('The control programming appears to have a duct pressure reset strategy implemented as the standard deviation of the duct pressure setpoint data does not equal zero. No further action maybe necessary if the faults are low and the fan system is delivering enough air under all conditions to VAV boxes.')

paragraph = document.add_paragraph()
run = paragraph.add_run(f'Report generated: {time.ctime()}')
run.style = 'Emphasis'

document.save('./final_report/ahu_fc1_report.docx')
print('All Done')
