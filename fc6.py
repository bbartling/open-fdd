from pandas import concat
from pandas import DataFrame
import operator
import time
from datetime import datetime, timedelta

import pandas as pd
import numpy as np

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from docx import Document
from docx.shared import Inches

import argparse
import math

# python 3.10 on Windows 10
# py .\fc6.py -i ./ahu_data/hvac_random_fake_data/fc6_fake_data1.csv -o fake1_ahu_fc6_report

parser = argparse.ArgumentParser(add_help=False)
args = parser.add_argument_group('Options')

args.add_argument('-h', '--help', action='help',
                  help='Show this help message and exit.')
args.add_argument('-i', '--input', required=True, type=str,
                  help='CSV File Input')
args.add_argument('-o', '--output', required=True, type=str,
                  help='Word File Output Name')
'''
args.add_argument('--use-flask', default=False, action='store_true')
args.add_argument('--no-flask', dest='use-flask', action='store_false')
'''
args = parser.parse_args()


def fault_condition_six_(df):
    return operator.and_(df.rat_minus_oat >= df.oat_rat_delta_min,
                          df.percent_oa_calc_minus_perc_OAmin > df.airflow_err_thres
                         )


df = pd.read_csv(args.input,
                 index_col='Date',
                 parse_dates=True).rolling('5T').mean()

print(df)
df_copy = df.copy()

df_copy.plot(figsize=(25, 8),
             title='AHU OA Fraction Data')
plt.savefig('./static/ahu_fc6_signals.png')

# make an entire column out of these params in the Pandas Dataframe
# required params taken from the screenshot above
AIRFLOW_ERR_THRES = .3
AHU_MIN_CFM_STP = 3000
OAT_DEGF_ERR_THRES = 5
RAT_DEGF_ERR_THRES = 2
DELTA_TEMP_MIN = 10

df['airflow_err_thres'] = AIRFLOW_ERR_THRES
df['ahu_min_cfm_stp'] = AHU_MIN_CFM_STP
df['oat_degf_err_thres'] = OAT_DEGF_ERR_THRES
df['rat_degf_err_thres'] = RAT_DEGF_ERR_THRES

df['oat_rat_delta_min'] = DELTA_TEMP_MIN
df['rat_minus_oat'] = abs(df['rat'] - df['oat'])

df['percent_oa_calc'] = (df['mat'] - df['rat']) / (df['oat'] - df['rat'])
df['percent_oa_calc'] = df['percent_oa_calc'].apply(lambda x : x if x > 0 else 0)

#df['actual_oa_calc'] = (df['percent_oa_calc']/100) * df['vav_total_flow']
#df['actual_oa_calc'] = (df['percent_oa_calc']) * df['vav_total_flow']
df['perc_OAmin'] = AHU_MIN_CFM_STP / df['vav_total_flow']
df['percent_oa_calc_minus_perc_OAmin'] = abs(df['percent_oa_calc'] - df['perc_OAmin'])


print('percent_oa_calc: ', df.percent_oa_calc)
print('percent_oa_calc_minus_perc_OAmin: ', df.percent_oa_calc_minus_perc_OAmin)
print('EQUAL?: ',df['percent_oa_calc'] == df['percent_oa_calc_minus_perc_OAmin'])

start = df.head(1).index.date
print('Dataset start: ', start)

end = df.tail(1).index.date
print('Dataset end: ', end)
print('COLUMNS: ', print(df.columns))

df['fc6_flag'] = fault_condition_six_(df)

df2 = df.copy().dropna()
df2['fc6_flag'] = df2['fc6_flag'].astype(int)

# drop params column for better plot
df2 = df2.drop(['airflow_err_thres',
                'oat_degf_err_thres',
                'rat_degf_err_thres',
                'oat_rat_delta_min'], axis=1)

print(df2)

fig, (ax1, ax2, ax3, ax4, ax5) = plt.subplots(5, 1, figsize=(25, 8))
plt.title('Fault Conditions 6 Plot')

plot1a, = ax1.plot(df2.index, df2.mat, color='r', label="Mix")  # red
plot1b, = ax1.plot(df2.index, df2.oat, color='b', label="Outside")  # blue
plot1c, = ax1.plot(df2.index, df2.rat, color='m', label="Return")  # magenta
ax1.legend(loc='best')
ax1.set_ylabel('AHU Temps °F')

ax2.plot(df2.index, df2.vav_total_flow, color='k',label="AHU Total Air Flow")  # black
ax2.legend(loc='best')
ax2.set_ylabel('CFM')

ax3.plot(df2.index, df2.rat_minus_oat, label="rat - oat")  # black
ax3.legend(loc='best')
ax3.set_ylabel('°F')

ax4.plot(df2.index, df2.percent_oa_calc, label="%OAfrac")  # yellow=
ax4.legend(loc='best')
ax4.set_ylabel('%')

ax5.plot(df2.index, df2.fc6_flag, color='g', label="Fault")  # green
ax5.set_xlabel('Date')
ax5.set_ylabel('Fault Flags')
ax5.legend(loc='best')

plt.legend()
plt.tight_layout()
plt.savefig('./static/ahu_fc6_fans_plot.png')
# plt.show()

print("Starting ahu fc6 docx report")
document = Document()
document.add_heading('Fault Condition Six Report', 0)

p = document.add_paragraph(
    'Fault condition six of ASHRAE Guideline 36 is an attempt at verifying that AHU design minimum outside air is close to the calculated outside air fraction through the outside, mix, and return air temperature sensors. Fault condition six equation as defined by ASHRAE:')
document.add_picture('./images/fc6_definition.png', width=Inches(6))

# ADD IN SUBPLOTS SECTION
document.add_heading('Dataset Plot', level=2)
document.add_picture('./static/ahu_fc6_fans_plot.png', width=Inches(6))
document.add_heading('Dataset Statistics', level=2)

# calculate dataset statistics
delta = df2.index.to_series().diff()
total_days = round(delta.sum() / pd.Timedelta(days=1), 2)
print('DAYS ALL DATA: ', total_days)
total_hours = delta.sum() / pd.Timedelta(hours=1)
print('TOTAL HOURS: ', total_hours)
hours_fc6_mode = (delta * df2["fc6_flag"]).sum() / pd.Timedelta(hours=1)
print('FALT FLAG TRUE TOTAL HOURS: ', hours_fc6_mode)
percent_true = round(df2.fc6_flag.mean() * 100, 2)
print('PERCENT TIME WHEN FLAG 6 TRUE: ', percent_true, '%')
percent_false = round((100 - percent_true), 2)
print('PERCENT TIME WHEN FLAG 6 FALSE: ', percent_false, '%')
df2['hour_of_the_day_fc6'] = df2.index.hour.where(df2["fc6_flag"] == 1)
flag_true_oafrac = round(
    df2.percent_oa_calc.where(df2["fc6_flag"] == 1).mean(), 2)
print('FLAG TRUE OA FRAC: ', flag_true_oafrac)

# make hist plots fc6
fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
ax.hist(df2.hour_of_the_day_fc6.dropna())
ax.set_xlabel('24 Hour Number in Day')
ax.set_ylabel('Frequency')
ax.set_title(f'Hour-Of-Day When Fault Flag 6 is TRUE')
fig.savefig('./static/ahu_fc6_histogram.png')

# add calcs to word doc
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time in days calculated in dataset: {total_days}')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time in hours calculated in dataset: {total_hours}')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time in hours for when fault flag 6 is True: {hours_fc6_mode}')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Percent of time in the dataset when the fault flag 6 is True: {percent_true}%')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Percent of time in the dataset when fault flag 6 is False: {percent_false}%')


if hours_fc6_mode != float(0):
    paragraph = document.add_paragraph()
    # ADD HIST Plots
    document.add_heading('Time-of-day Histogram Plots', level=2)
    document.add_picture('./static/ahu_fc6_histogram.png', width=Inches(6))

if not math.isnan(flag_true_oafrac):
    paragraph = document.add_paragraph()
    paragraph.style = 'List Bullet'
    paragraph.add_run(
        f'When fault condition 6 is True the average AHU xxx is {flag_true_oafrac}°F and the xxx is {flag_true_oafrac}°F. This could possibly help with pin pointing AHU operating conditions for when this fault is True.')

paragraph = document.add_paragraph()

# ADD in Summary Statistics
document.add_heading('Mix Temp Statistics', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(str(df2.mat.describe()))

# ADD in Summary Statistics
document.add_heading('Outside Temp Statistics', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(str(df2.oat.describe()))

# ADD in Summary Statistics
document.add_heading('Return Temp Statistics', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(str(df2.rat.describe()))

# ADD in Summary Statistics
document.add_heading('Calculated OA Fraction Statistics', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(str(df2.percent_oa_calc.describe()))

document.add_heading('Suggestions based on data analysis', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'

if percent_true > 5.0:
    paragraph.add_run(
        'The percent True metric that represents the amount of time for when the fault flag is True is high indicating the AHU temperature sensors are out of calibration')

else:
    paragraph.add_run(
        'The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the AHU temperature sensors are within calibration')

print('df2.mat.std: ', df2.mat.std())
print('df2.mat.min: ', df2.mat.min())
print('df2.mat.max: ', df2.mat.max())

paragraph = document.add_paragraph()
run = paragraph.add_run(f'Report generated: {time.ctime()}')
run.style = 'Emphasis'

document.save(f'./final_report/{args.output}.docx')
print('All Done')

# df2.to_csv('testdf2_fc6.csv')



