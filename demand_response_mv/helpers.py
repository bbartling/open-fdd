import pandas as pd
from pandas.tseries.offsets import BDay
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import datetime, os, time
from io import BytesIO


def generate_report(
    TEST_CASE_ID,
    formatted_date,
    filtered_data_plot,
    test_start_hour,
    test_end_hour,
    main_test_day_plot,
    main_average_net_diff_event,
    main_average_perc_diff_event,
    main_max_,
    main_max_test_day,
    main_total_energy_,
    main_total_energy_test_day,
    ahu_test_day_plot,
    ahu_average_net_diff_event,
    ahu_average_perc_diff_event,
    ahu_max_,
    ahu_max_test_day,
    ahu_total_energy_,
    ahu_total_energy_test_day,
    title
) -> Document:
    print(f"Starting docx report!")
    document = Document()
    document.add_heading(title, 0)

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"Test Case {TEST_CASE_ID} on {formatted_date}")
    run.bold = True
    run.font.size = Pt(12)

    filtered_data_plot_image = BytesIO()
    filtered_data_plot.savefig(filtered_data_plot_image, format="png")
    plt.close()  # This closes the specific Figure object associated with filtered_data_plot
    filtered_data_plot_image.seek(0)

    document.add_heading("Data used to calculate an averaged load profile", level=2)
    document.add_picture(
        filtered_data_plot_image,
        width=Inches(6),
    )

    document.add_heading(
        f"Main meter test day plot event from hour {test_start_hour} to hour {test_end_hour}",
        level=2,
    )

    main_test_day_plot_image = BytesIO()
    main_test_day_plot.savefig(main_test_day_plot_image, format="png")
    plt.close()  # This closes the specific Figure object associated with test_day_plot
    main_test_day_plot_image.seek(0)

    # ADD IN SUBPLOTS SECTION
    document.add_picture(
        main_test_day_plot_image,
        width=Inches(6),
    )

    document.add_heading("Main Meter Statistics", level=2)

    paragraph = document.add_paragraph()
    paragraph.style = "List Bullet"
    paragraph.add_run(
        f"Main meter net change power during event: {round(main_average_net_diff_event,2)} Watts"
    )

    paragraph = document.add_paragraph()
    paragraph.style = "List Bullet"
    paragraph.add_run(
        f"Main meter percent change in power during event: {round(main_average_perc_diff_event)} %"
    )

    paragraph = document.add_paragraph()
    paragraph.style = "List Bullet"
    paragraph.add_run(
        f"Averaged data main meter max power: {round(main_max_)} Watts"
    )

    paragraph = document.add_paragraph()
    paragraph.style = "List Bullet"
    paragraph.add_run(
        f"Test day main meter max power: {round(main_max_test_day)} Watts"
    )

    paragraph = document.add_paragraph()
    paragraph.style = "List Bullet"
    paragraph.add_run(
        f"Averaged data main meter calculated energy: {round(main_total_energy_)} Watt-Hr"
    )

    paragraph = document.add_paragraph()
    paragraph.style = "List Bullet"
    paragraph.add_run(
        f"Test day main meter calculated energy: {round(main_total_energy_test_day)} Watt-Hr"
    )

    # START AHU
    document.add_heading(
        f"AHU meter test day plot event from hour {test_start_hour} to hour {test_end_hour}",
        level=2,
    )

    ahu_test_day_plot_image = BytesIO()
    ahu_test_day_plot.savefig(ahu_test_day_plot_image, format="png")
    plt.close()  # This closes the specific Figure object associated with test_day_plot
    ahu_test_day_plot_image.seek(0)

    document.add_picture(
        ahu_test_day_plot_image,
        width=Inches(6),
    )

    document.add_heading("AHU Meter Statistics", level=2)

    paragraph = document.add_paragraph()
    paragraph.style = "List Bullet"
    paragraph.add_run(
        f"AHU meter net change power during event: {round(ahu_average_net_diff_event,2)} Watts"
    )

    paragraph = document.add_paragraph()
    paragraph.style = "List Bullet"
    paragraph.add_run(
        f"AHU meter percent change in power during event: {round(ahu_average_perc_diff_event)} %"
    )

    paragraph = document.add_paragraph()
    paragraph.style = "List Bullet"
    paragraph.add_run(f"Averaged data ahu meter max power: {round(ahu_max_)} Watts")

    paragraph = document.add_paragraph()
    paragraph.style = "List Bullet"
    paragraph.add_run(f"Test day ahu meter max power: {round(ahu_max_test_day)} Watts")

    paragraph = document.add_paragraph()
    paragraph.style = "List Bullet"
    paragraph.add_run(
        f"Averaged data ahu meter calculated energy: {round(ahu_total_energy_)} Watt-Hr"
    )

    paragraph = document.add_paragraph()
    paragraph.style = "List Bullet"
    paragraph.add_run(
        f"Test day ahu meter calculated energy: {round(ahu_total_energy_test_day)} Watt-Hr"
    )

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"Report generated: {time.ctime()}")
    run.style = "Emphasis"

    return document


def meters_stats_calcs(data, series):
    main_max_power = data[series].max()
    # Group by hour (4 readings per hour) and calculate the mean
    data["hour"] = data.index // 4
    hourly_data = data.groupby("hour")[series].mean()
    total_energy = hourly_data.sum()
    return main_max_power, total_energy


def plot_data_to_be_avg(filtered_data):
    filtered_data = filtered_data.reset_index().drop("index", axis=1)
    plot = filtered_data.plot(figsize=(25, 8))
    plot.set_xlabel("Weekdays Only")
    plot.set_ylabel("Watts")
    # Get the Figure object related to the Axes object 'plot'
    fig = plot.get_figure()
    return fig

def find_previous_10_days(suitable_baseline_no, all_data, test_case_date):
    weekdays_count = 0
    filtered_dates = []

    # convert to pandas datetime format if not already
    all_data["Date"] = pd.to_datetime(all_data["Date"])
    suitable_baseline_no["Date"] = pd.to_datetime(suitable_baseline_no["Date"])
    test_case_date = pd.to_datetime(test_case_date)

    # Convert date column of suitable_baseline_no to a set for efficient searching
    suitable_baseline_no_dates = set(suitable_baseline_no["Date"].dt.date)
    all_dates = set(all_data["Date"].dt.date)

    # Count backwards from test_case_date to find 10 complete weekdays (excluding weekends)
    current_date = test_case_date
    while weekdays_count < 10:
        current_date -= pd.Timedelta(days=1)
        if current_date.weekday() < 5:  # Monday to Friday
            if current_date.date() not in suitable_baseline_no_dates:  # Exclude suitable_baseline_no dates
                # Check if current_date is in all_dates before adding to filtered_dates
                if current_date.date() in all_dates:
                    filtered_dates.append(current_date.date())  # Use .date() to ensure format consistency
                    weekdays_count += 1

    # Filter all_data DataFrame for weekdays in filtered_dates
    filtered_data = all_data[all_data["Date"].dt.date.isin(filtered_dates)]
    filtered_data_copy = filtered_data.copy()
    return filtered_data_copy

def find_closest_weather_dates(all_data, suitable_baseline_no, test_case_date, num_dates=10):
    weekdays_count = 0
    filtered_dates = []

    # Convert to pandas datetime format if not already
    all_data["Date"] = pd.to_datetime(all_data["Date"])
    suitable_baseline_no["Date"] = pd.to_datetime(suitable_baseline_no["Date"])
    test_case_date = pd.to_datetime(test_case_date)

    # Convert date column of suitable_baseline_no to a set for efficient searching
    suitable_baseline_no_dates = set(suitable_baseline_no["Date"].dt.date)
    all_dates = set(all_data["Date"].dt.date)

    # Count backwards from test_case_date to find 10 complete weekdays (excluding weekends)
    current_date = test_case_date
    while weekdays_count < num_dates:
        current_date -= pd.Timedelta(days=1)
        if current_date.weekday() < 5:  # Monday to Friday
            if current_date.date() not in suitable_baseline_no_dates:  # Exclude suitable_baseline_no dates
                # Check if current_date is in all_dates before adding to filtered_dates
                if current_date.date() in all_dates:
                    filtered_dates.append(current_date.date())  # Use .date() to ensure format consistency
                    weekdays_count += 1

    # Filter all_data DataFrame for weekdays in filtered_dates
    filtered_data = all_data[all_data["Date"].dt.date.isin(filtered_dates)]

    # Calculate daily averages of temperature and relative humidity for filtered_data
    daily_averages = filtered_data.groupby(filtered_data['Date'].dt.date).mean()
    unique_dates = daily_averages.reset_index()
    unique_dates['Date'] = pd.to_datetime(unique_dates['Date'])
    unique_dates['Difference'] = abs(unique_dates['Date'] - test_case_date)
    unique_dates = unique_dates.sort_values('Difference')
    closest_dates = unique_dates.head(num_dates)
    closest_dates_only = closest_dates['Date'].dt.date
    print("10 Closest Weather Dates:\n", closest_dates_only)

    # Filter all_data DataFrame for closest_dates_only
    filtered_data = all_data[all_data["Date"].dt.date.isin(closest_dates_only)]

    return filtered_data


def calculate_power_averages(filtered_data_copy):
    # Create an empty dictionary to store the lists for each time step
    main_all_power_dict = {"main_all_power": [[] for _ in range(96)]}
    ahu_all_power_dict = {"ahu_all_power": [[] for _ in range(96)]}
    solar_all_power_dict = {"solar_all_power": [[] for _ in range(96)]}

    counter = 0
    # Loop through the rows of the DataFrame in groups of 96
    for i, row in enumerate(filtered_data_copy.iterrows()):
        main_all_power_val = row[1]["main_all_power"]
        ahu_all_power_val = row[1]["ahu_all_power"]
        solar_all_power_val = row[1]["solar_all_power"]

        main_all_power_dict["main_all_power"][counter].append(main_all_power_val)
        ahu_all_power_dict["ahu_all_power"][counter].append(ahu_all_power_val)
        solar_all_power_dict["solar_all_power"][counter].append(solar_all_power_val)

        counter += 1
        if counter == 96:
            counter = 0

    # Calculate the average of each list and store it in a new dictionary
    main_all_power_avg_dict = {
        "main_power_average": [
            sum(lst) / len(lst) for lst in main_all_power_dict["main_all_power"]
        ]
    }
    ahu_all_power_avg_dict = {
        "ahu_power_average": [
            sum(lst) / len(lst) for lst in ahu_all_power_dict["ahu_all_power"]
        ]
    }
    solar_all_power_avg_dict = {
        "solar_power_average": [
            sum(lst) / len(lst) for lst in solar_all_power_dict["solar_all_power"]
        ]
    }

    return main_all_power_avg_dict, ahu_all_power_avg_dict, solar_all_power_avg_dict


def generate_power_plot(DEVICE_DATA, data, test_case_data, dates, TEST_CASE_ID):
    # Generate time series data
    hours = np.arange(0, 24)
    minutes = np.arange(0, 60, 15)
    timestamps = []
    for hour in hours:
        for minute in minutes:
            timestamps.append(hour)

    # Create pandas DataFrame for hour of day
    hours = pd.DataFrame({"hour": timestamps})

    # Combine data and test_case_data into a single DataFrame
    test_case_main_power_combined_data = pd.DataFrame(
        test_case_data[f"{DEVICE_DATA}_all_power"].values,
        columns=[f"Test Case {TEST_CASE_ID} Data"],
    )
    combined_data = data.join(test_case_main_power_combined_data)
    combined_data = combined_data.join(hours)
    combined_data_copy = combined_data

    # Drop rows with missing 'Date' in 'dates'
    dates = dates.dropna(subset=["Date"])
    dates["Date"] = pd.to_datetime(
        dates["Date"], format="%A, %B %d, %Y", errors="coerce"
    )
    test_case_date = dates[dates["Test Case #"] == TEST_CASE_ID]["Date"].min()

    test_start_hour = int(
        dates[dates["Test Case #"] == TEST_CASE_ID]["Test Start"].values[0]
    )
    test_end_hour = int(
        dates[dates["Test Case #"] == TEST_CASE_ID]["Test End"].values[0]
    )

    start_time = test_case_date.replace(hour=0, minute=0, second=0)
    end_time = start_time + pd.DateOffset(days=1) - pd.DateOffset(minutes=15)
    timestamps = pd.date_range(start=start_time, end=end_time, freq="15T")

    # Set the timestamps as the index of the DataFrame
    combined_data.set_index(timestamps, inplace=True)
    combined_data = combined_data[
        [f"{DEVICE_DATA}_power_average", f"Test Case {TEST_CASE_ID} Data"]
    ]

    # Plot the combined DataFrame
    fig, ax = plt.subplots(figsize=(25, 8))
    ax.plot(
        combined_data.index,
        combined_data[f"{DEVICE_DATA}_power_average"],
        label="Averaged Data Weekday",
    )
    ax.plot(
        combined_data.index,
        combined_data[f"Test Case {TEST_CASE_ID} Data"],
        label=f"Test Case {TEST_CASE_ID} Data",
    )

    # Set the x-axis to show only the hour of the day
    ax.xaxis.set_major_locator(mdates.HourLocator(interval=1))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%H:%M"))

    # Rotate the x-axis tick labels by 45 degrees for better readability
    plt.xticks(rotation=45)

    plt.xlabel("Hour")
    plt.ylabel("Watts")
    plt.title(f"Averaged data {DEVICE_DATA} with Test Case {TEST_CASE_ID}")
    plt.legend(loc="upper left")  # Move the legend to the left-hand side

    return fig, combined_data_copy, test_start_hour, test_end_hour


def calculate_differences(
    combined_data, device, test_case_id, test_start_hour=None, test_end_hour=None
):
    combined_data["Percentage Difference"] = (
        (
            combined_data[f"Test Case {test_case_id} Data"]
            - combined_data[f"{device}_power_average"]
        )
        / combined_data[f"{device}_power_average"]
    ) * 100

    if test_start_hour is not None and test_end_hour is not None:
        filtered_data = combined_data[
            (combined_data["hour"] >= test_start_hour)
            & (combined_data["hour"] <= test_end_hour)
        ]
        filtered_data["Net Difference"] = (
            filtered_data[f"Test Case {test_case_id} Data"]
            - filtered_data[f"{device}_power_average"]
        )
        average_net_difference = filtered_data["Net Difference"].mean()

        filtered_data["Percentage Difference"] = (
            filtered_data["Net Difference"] / filtered_data[f"{device}_power_average"]
        ) * 100
        average_percentage_difference = filtered_data["Percentage Difference"].mean()
    else:
        average_net_difference = None
        average_percentage_difference = None

    return average_net_difference, average_percentage_difference
