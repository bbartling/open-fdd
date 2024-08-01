import os
from io import BytesIO
from docx import Document
from docx.shared import Inches
from open_fdd.air_handling_unit.reports.base_report import BaseReport
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import time
import pkg_resources

class FaultCodeSevenReport(BaseReport):
    """Class provides the definitions for Fault Code 7 Report.
        Very similar to FC 13 but uses heating valve
    """

    def __init__(self, config):
        super().__init__(config)
        self.sat_col = config['SAT_COL']
        self.sat_setpoint_col = config['SAT_SETPOINT_COL']
        self.heating_sig_col = config['HEATING_SIG_COL']
        self.supply_vfd_speed_col = config['SUPPLY_VFD_SPEED_COL']

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> BytesIO:
        if output_col is None:
            output_col = "fc7_flag"

        fig, (ax1, ax2, ax3) = plt.subplots(3, 1, figsize=(25, 8))
        plt.title('Fault Conditions 7 Plot')

        ax1.plot(df.index, df[self.sat_col], label="SAT")
        ax1.plot(df.index, df[self.sat_setpoint_col], label="SATsp")
        ax1.legend(loc='best')
        ax1.set_ylabel('AHU Supply Temps °F')

        ax2.plot(df.index, df[self.heating_sig_col], color="r", label="AHU Heat Vlv")
        ax2.legend(loc='best')
        ax2.set_ylabel('%')

        ax3.plot(df.index, df[output_col], label="Fault", color="k")
        ax3.set_xlabel('Date')
        ax3.set_ylabel('Fault Flags')
        ax3.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        plot_image = BytesIO()
        fig.savefig(plot_image, format="png")
        plot_image.seek(0)
        plt.close(fig)  # Close the figure

        return plot_image

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc7_flag"

        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)
        total_hours = delta.sum() / pd.Timedelta(hours=1)
        hours_fc7_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)
        percent_true = round(df[output_col].mean() * 100, 2)
        percent_false = round((100 - percent_true), 2)

        flag_true_satsp = round(df[self.sat_setpoint_col].where(df[output_col] == 1).mean(), 2)
        flag_true_sat = round(df[self.sat_col].where(df[output_col] == 1).mean(), 2)

        motor_on = df[self.supply_vfd_speed_col].gt(.01).astype(int)
        hours_motor_runtime = round((delta * motor_on).sum() / pd.Timedelta(hours=1), 2)

        # For summary stats on I/O data to make useful
        df_motor_on_filtered = df[df[self.supply_vfd_speed_col] > 0.1]

        return (
            total_days,
            total_hours,
            hours_fc7_mode,
            percent_true,
            percent_false,
            flag_true_satsp,
            flag_true_sat,
            hours_motor_runtime,
            df_motor_on_filtered
        )

    def create_hist_plot(self, df: pd.DataFrame, output_col: str = None) -> BytesIO:
        if output_col is None:
            output_col = "fc7_flag"

        # Calculate dataset statistics
        df["hour_of_the_day_fc7"] = df.index.hour.where(df[output_col] == 1)

        # Make hist plots fc7
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc7.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 7 is TRUE")

        hist_plot_image = BytesIO()
        fig.savefig(hist_plot_image, format="png")
        hist_plot_image.seek(0)
        plt.close(fig)  # Close the figure

        return hist_plot_image

    def create_report(self, path: str, df: pd.DataFrame, output_col: str = None) -> None:
        if output_col is None:
            output_col = "fc7_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Seven Report", 0)

        p = document.add_paragraph(
            """Fault condition seven of ASHRAE Guideline 36 is an AHU heating mode only with an attempt at verifying an AHU heating or cooling valve is not stuck or leaking by verifying AHU supply temperature to supply temperature setpoint. Fault condition seven equation as defined by ASHRAE:"""
        )

        image_path = pkg_resources.resource_filename('open_fdd', 'air_handling_unit/images/fc7_definition.png')
        document.add_picture(image_path, width=Inches(6))
        document.add_heading("Dataset Plot", level=2)

        plot_image = self.create_plot(df, output_col=output_col)
        document.add_picture(plot_image, width=Inches(6))

        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc7_mode,
            percent_true,
            percent_false,
            flag_true_satsp,
            flag_true_sat,
            hours_motor_runtime,
            df_motor_on_filtered
        ) = self.summarize_fault_times(df, output_col=output_col)

        stats = [
            f"Total time in days calculated in dataset: {total_days}",
            f"Total time in hours calculated in dataset: {total_hours}",
            f"Total time in hours for when fault flag is True: {hours_fc7_mode}",
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%",
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%",
            f"Calculated motor runtime in hours based off of VFD signal > zero: {hours_motor_runtime}"
        ]

        for stat in stats:
            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(stat)

        if int(total_hours) == int(hours_motor_runtime):
            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"This fan system appears to run 24/7 consider implementing occupancy schedules to reduce building fuel use through HVAC"
            )

        # If there are no faults, skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = self.create_hist_plot(df, output_col=output_col)
            document.add_picture(histogram_plot_image, width=Inches(6))

            paragraph = document.add_paragraph()
            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 7 is True the average AHU supply air setpoint is {flag_true_satsp} in °F and the supply air temperature is {flag_true_sat} in °F.'
            )
        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")
            paragraph = document.add_paragraph()
            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.'
            )

        document.add_heading('Summary Statistics filtered for when the AHU is running', level=1)

        document.add_heading('Supply Air Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.sat_col].describe()))

        document.add_heading('Supply Air Temp Setpoint', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.sat_setpoint_col].describe()))

        document.add_heading('Heating Coil Valve', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.heating_sig_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true > 5.0:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is high indicating the AHU heating valve maybe broken or there could be a flow issue with the amount of hot water flowing through the coil or that the boiler system reset is too aggressive and there isnt enough heat being produced by this coil. It could be worth viewing mechanical blue prints for this AHU design schedule to see what hot water temperature this coil was designed for and compare it to actual hot water supply temperatures. IE., an AHU hot water coil sized to have a 180°F water flowing through it may have a durastic reduction in performance the colder the hot water is flowing through it, if need be consult a mechanical design engineer to rectify.'
            )
        else:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the AHU heating valve operates Ok.'
            )

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        document.save(f"{path}/report_fc7.docx")
