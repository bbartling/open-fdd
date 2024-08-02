import os
from io import BytesIO
from docx import Document
from docx.shared import Inches
from open_fdd.air_handling_unit.reports.base_report import BaseReport
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import time
import pkg_resources



class FaultCodeEightReport(BaseReport):
    """Class provides the definitions for Fault Code 8 Report."""

    def __init__(self, config):
        super().__init__(config)
        self.sat_col = config['SAT_COL']
        self.mat_col = config['MAT_COL']
        self.supply_vfd_speed_col = config['SUPPLY_VFD_SPEED_COL']
        self.economizer_sig_col = config['ECONOMIZER_SIG_COL']

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> BytesIO:
        if output_col is None:
            output_col = "fc8_flag"

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(25, 8))
        plt.title('Fault Conditions 8 Plot')

        ax1.plot(df.index, df[self.sat_col], label="SAT")
        ax1.plot(df.index, df[self.mat_col], label="MAT")
        ax1.legend(loc='best')
        ax1.set_ylabel('AHU Temps °F')

        ax2.plot(df.index, df[output_col], label="Fault", color="k")
        ax2.set_xlabel('Date')
        ax2.set_ylabel('Fault Flags')
        ax2.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        plot_image = BytesIO()
        fig.savefig(plot_image, format="png")
        plot_image.seek(0)
        plt.close(fig)  # Close the figure

        return plot_image

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc8_flag"

        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)
        total_hours = delta.sum() / pd.Timedelta(hours=1)
        hours_fc8_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)
        percent_true = round(df[output_col].mean() * 100, 2)
        percent_false = round((100 - percent_true), 2)

        flag_true_mat = round(df[self.mat_col].where(df[output_col] == 1).mean(), 2)
        flag_true_sat = round(df[self.sat_col].where(df[output_col] == 1).mean(), 2)

        motor_on = df[self.supply_vfd_speed_col].gt(.01).astype(int)
        hours_motor_runtime = round((delta * motor_on).sum() / pd.Timedelta(hours=1), 2)

        # For summary stats on I/O data to make useful
        df_motor_on_filtered = df[df[self.supply_vfd_speed_col] > 0.1]

        return (
            total_days,
            total_hours,
            hours_fc8_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_sat,
            hours_motor_runtime,
            df_motor_on_filtered
        )

    def create_hist_plot(self, df: pd.DataFrame, output_col: str = None) -> BytesIO:
        if output_col is None:
            output_col = "fc8_flag"

        # Calculate dataset statistics
        df["hour_of_the_day_fc8"] = df.index.hour.where(df[output_col] == 1)

        # Make hist plots fc8
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc8.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 8 is TRUE")

        hist_plot_image = BytesIO()
        fig.savefig(hist_plot_image, format="png")
        hist_plot_image.seek(0)
        plt.close(fig)  # Close the figure

        return hist_plot_image

    def create_report(self, path: str, df: pd.DataFrame, output_col: str = None, report_name: str = "report_fc8.docx") -> None:
        if output_col is None:
            output_col = "fc8_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Eight Report", 0)

        p = document.add_paragraph(
            """Fault condition eight of ASHRAE Guideline 36 is an AHU economizer free cooling mode only with an attempt at flagging conditions when the AHU mixing air temperature the supply air temperature are not approximately equal. Fault condition eight equation as defined by ASHRAE:"""
        )

        image_path = pkg_resources.resource_filename('open_fdd', 'air_handling_unit/images/fc8_definition.png')
        document.add_picture(image_path, width=Inches(6))
        document.add_heading("Dataset Plot", level=2)

        plot_image = self.create_plot(df, output_col=output_col)
        document.add_picture(plot_image, width=Inches(6))

        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc8_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_sat,
            hours_motor_runtime,
            df_motor_on_filtered
        ) = self.summarize_fault_times(df, output_col=output_col)

        stats = [
            f"Total time in days calculated in dataset: {total_days}",
            f"Total time in hours calculated in dataset: {total_hours}",
            f"Total time in hours for when fault flag is True: {hours_fc8_mode}",
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%",
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%",
            f"Calculated motor runtime in hours based off of VFD signal > zero: {hours_motor_runtime}"
        ]

        for stat in stats:
            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(stat)

        if int(total_hours) == int(hours_motor_runtime):
            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"This fan system appears to run 24/7 consider implementing occupancy schedules to reduce building fuel use through HVAC"
            )

        # If there are no faults, skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = self.create_hist_plot(df, output_col=output_col)
            document.add_picture(histogram_plot_image, width=Inches(6))

            paragraph = document.add_paragraph()
            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 8 is True the average AHU mix air is {flag_true_mat} in °F and the supply air temperature is {flag_true_sat} in °F.'
            )
        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")
            paragraph = document.add_paragraph()
            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.'
            )

        document.add_heading('Summary Statistics filtered for when the AHU is running', level=1)

        document.add_heading('Supply Air Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.sat_col].describe()))

        document.add_heading('Mix Air Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.mat_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true > 5.0:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is high indicating temperature sensor error or the heating/cooling coils are leaking potentially creating simultenious heating/cooling which can be an energy penalty for running the AHU in this fashion. Verify AHU mix/supply temperature sensor calibration in addition to a potential mechanical issue of a leaking valve. A leaking valve can be troubleshot by isolating the valve closed by manual shut off valves where piping lines enter the AHU coil and then verifying any changes in the AHU discharge air temperature.'
            )
        else:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is low indicating the AHU components are within calibration for this fault equation Ok.'
            )

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        document.save(f"{path}/{report_name}")
