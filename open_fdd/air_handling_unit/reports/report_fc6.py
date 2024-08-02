import os
from io import BytesIO
from docx import Document
from docx.shared import Inches
from open_fdd.air_handling_unit.reports.base_report import BaseReport
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import time
import pkg_resources

class FaultCodeSixReport(BaseReport):
    """Class provides the definitions for Fault Code 6 Report."""

    def __init__(self, config):
        super().__init__(config)
        self.supply_fan_air_volume_col = config['SUPPLY_FAN_AIR_VOLUME_COL']
        self.mat_col = config['MAT_COL']
        self.oat_col = config['OAT_COL']
        self.rat_col = config['RAT_COL']
        self.supply_vfd_speed_col = config['SUPPLY_VFD_SPEED_COL']

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> BytesIO:
        if output_col is None:
            output_col = "fc6_flag"

        fig, (ax1, ax2, ax3, ax4, ax5) = plt.subplots(5, 1, figsize=(25, 8))
        plt.title('Fault Conditions 6 Plot')
        ax1.plot(df.index, df['rat_minus_oat'], label="Rat Minus Oat")
        ax1.legend(loc='best')
        ax1.set_ylabel("°F")

        ax2.plot(df.index, df[self.supply_fan_air_volume_col], label="Total Air Flow", color="r")
        ax2.set_xlabel('Date')
        ax2.set_ylabel('CFM')
        ax2.legend(loc='best')

        ax3.plot(df.index, df['percent_oa_calc'], label="OA Frac Calc", color="m")
        ax3.plot(df.index, df['perc_OAmin'], label="OA Perc Min Calc", color="y")
        ax3.set_xlabel('Date')
        ax3.set_ylabel('%')
        ax3.legend(loc='best')

        ax4.plot(df.index, df['percent_oa_calc_minus_perc_OAmin'], label="OA Error Frac Vs Perc Min Calc", color="g")
        ax4.set_xlabel('Date')
        ax4.set_ylabel('%')
        ax4.legend(loc='best')

        ax5.plot(df.index, df[output_col], label="Fault", color="k")
        ax5.set_xlabel('Date')
        ax5.set_ylabel('Fault Flags')
        ax5.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        plot_image = BytesIO()
        fig.savefig(plot_image, format="png")
        plot_image.seek(0)
        plt.close(fig)  # Close the figure

        return plot_image

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc6_flag"

        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)
        total_hours = delta.sum() / pd.Timedelta(hours=1)
        hours_fc6_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)
        percent_true = round(df[output_col].mean() * 100, 2)
        percent_false = round((100 - percent_true), 2)

        flag_true_mat = round(df[self.mat_col].where(df[output_col] == 1).mean(), 2)
        flag_true_rat = round(df[self.rat_col].where(df[output_col] == 1).mean(), 2)
        flag_true_oat = round(df[self.oat_col].where(df[output_col] == 1).mean(), 2)

        motor_on = df[self.supply_vfd_speed_col].gt(.01).astype(int)
        hours_motor_runtime = round((delta * motor_on).sum() / pd.Timedelta(hours=1), 2)

        # For summary stats on I/O data to make useful
        df_motor_on_filtered = df[df[self.supply_vfd_speed_col] > 0.1]

        return (
            total_days,
            total_hours,
            hours_fc6_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_rat,
            flag_true_oat,
            hours_motor_runtime,
            df_motor_on_filtered
        )

    def create_hist_plot(self, df: pd.DataFrame, output_col: str = None) -> BytesIO:
        if output_col is None:
            output_col = "fc6_flag"

        # Calculate dataset statistics
        df["hour_of_the_day_fc6"] = df.index.hour.where(df[output_col] == 1)

        # Make hist plots fc6
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc6.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 6 is TRUE")

        hist_plot_image = BytesIO()
        fig.savefig(hist_plot_image, format="png")
        hist_plot_image.seek(0)
        plt.close(fig)  # Close the figure

        return hist_plot_image

    def create_report(self, path: str, df: pd.DataFrame, output_col: str = None, report_name: str = "report_fc6.docx") -> None:
        if output_col is None:
            output_col = "fc6_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Six Report", 0)

        p = document.add_paragraph(
            """Fault condition six of ASHRAE Guideline 36 is an attempt at verifying that AHU design minimum outside air is close to the calculated outside air fraction through the outside, mix, and return air temperature sensors. A fault will get flagged in an AHU heating or mechanical cooling mode only if the calculated OA fraction is too low or too high as compared to percent Min calculation which is the AHU total air flow divided by the design minimum outdoor air expressed as a percent. Fault condition six equation as defined by ASHRAE:"""
        )

        image_path = pkg_resources.resource_filename('open_fdd', 'air_handling_unit/images/fc6_definition.png')
        document.add_picture(image_path, width=Inches(6))
        document.add_heading("Dataset Plot", level=2)

        plot_image = self.create_plot(df, output_col=output_col)
        document.add_picture(plot_image, width=Inches(6))

        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc6_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_rat,
            flag_true_oat,
            hours_motor_runtime,
            df_motor_on_filtered
        ) = self.summarize_fault_times(df, output_col=output_col)

        stats = [
            f"Total time in days calculated in dataset: {total_days}",
            f"Total time in hours calculated in dataset: {total_hours}",
            f"Total time in hours for when fault flag is True: {hours_fc6_mode}",
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%",
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%",
            f"Calculated motor runtime in hours based off of VFD signal > zero: {hours_motor_runtime}"
        ]

        for stat in stats:
            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(stat)

        if int(total_hours) == int(hours_motor_runtime):
            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"This fan system appears to run 24/7 consider implementing occupancy schedules to reduce building fuel use through HVAC"
            )

        # If there are no faults, skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = self.create_hist_plot(df, output_col=output_col)
            document.add_picture(histogram_plot_image, width=Inches(6))

            paragraph = document.add_paragraph()
            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 6 is True the average AHU mix air temperature {flag_true_mat}°F, outside air temperature is {flag_true_oat}°F, and the return air temperature is {flag_true_rat}°F. This could possibly help with pin pointing AHU operating conditions for when this AHU is drawing in excessive outside air.'
            )
        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")
            paragraph = document.add_paragraph()
            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.'
            )

        document.add_heading('Summary Statistics filtered for when the AHU is running', level=1)

        document.add_heading('Mix Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.mat_col].describe()))

        document.add_heading('Outside Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.oat_col].describe()))

        document.add_heading('Return Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.rat_col].describe()))

        document.add_heading('Total Air Flow', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.supply_fan_air_volume_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true > 5.0:
            paragraph.add_run(
                'The percent true metric maybe yielding sensors are out of calibration either on the AHU outside, mix, or return air temperature sensors that handle the OA fraction calculation or the totalized air flow calculation handled by a totalizing all VAV box air flows or AHU AFMS. Air flow and/or AHU temperature sensor may require recalibration.'
            )
        else:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is low indicating the sensors are within calibration'
            )

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        document.save(f"{path}/{report_name}")
