import os
from io import BytesIO
from docx import Document
from docx.shared import Inches
from open_fdd.air_handling_unit.reports.base_report import BaseReport
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import time
import pkg_resources
import sys

class FaultCodeFourReport(BaseReport):
    """Class provides the definitions for Fault Code 4 Report.
    Reporting the time series avg df that calculates control states per hour.
    """

    def __init__(self, config):
        super().__init__(config)
        self.delta_os_max = config['DELTA_OS_MAX']
        self.heating_mode_calc_col = 'heating_mode'
        self.econ_only_cooling_mode_calc_col = 'econ_only_cooling_mode'
        self.econ_plus_mech_cooling_mode_calc_col = 'econ_plus_mech_cooling_mode'
        self.mech_cooling_only_mode_calc_col = 'mech_cooling_only_mode'

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> BytesIO:
        if output_col is None:
            output_col = "fc4_flag"
        
        # Debug: print DataFrame columns to ensure they exist
        print("DataFrame columns:", df.columns)
        sys.stdout.flush()

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(25, 8))
        plt.title('Fault Conditions 4 Plots')

        ax1.plot(df.index, df[self.heating_mode_calc_col], label="Heat", color='orange')
        ax1.plot(df.index, df[self.econ_only_cooling_mode_calc_col], label="Econ Clg", color='olive')
        ax1.plot(df.index, df[self.econ_plus_mech_cooling_mode_calc_col], label="Econ + Mech Clg", color='c')
        ax1.plot(df.index, df[self.mech_cooling_only_mode_calc_col], label="Mech Clg", color='m')

        ax1.set_xlabel('Date')
        ax1.set_ylabel('Calculated AHU Operating States')
        ax1.legend(loc='best')

        ax2.plot(df.index, df[output_col], label="Fault", color="k")
        ax2.set_xlabel('Date')
        ax2.set_ylabel('Fault Flags')
        ax2.legend(loc='best')

        plt.tight_layout()

        plot_image = BytesIO()
        fig.savefig(plot_image, format="png")
        plot_image.seek(0)
        plt.close(fig)  # Close the figure

        return plot_image

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> tuple:
        if output_col is None:
            output_col = "fc4_flag"

        # Calculate dataset statistics
        delta_all_data = df.index.to_series().diff()

        total_days_all_data = round(delta_all_data.sum() / pd.Timedelta(days=1), 2)
        total_hours_all_data = delta_all_data.sum() / pd.Timedelta(hours=1)
        hours_fc4_mode = (delta_all_data * df[output_col]).sum() / pd.Timedelta(hours=1)
        percent_true_fc4 = round(df.fc4_flag.mean() * 100, 2)
        percent_false_fc4 = round((100 - percent_true_fc4), 2)

        # Heating mode runtime stats
        delta_heating = df[self.heating_mode_calc_col].index.to_series().diff()
        total_hours_heating = (delta_heating * df[self.heating_mode_calc_col]).sum() / pd.Timedelta(hours=1)
        percent_heating = round(df[self.heating_mode_calc_col].mean() * 100, 2)

        # Econ mode runtime stats
        delta_econ = df[self.econ_only_cooling_mode_calc_col].index.to_series().diff()
        total_hours_econ = (delta_econ * df[self.econ_only_cooling_mode_calc_col]).sum() / pd.Timedelta(hours(1))
        percent_econ = round(df[self.econ_only_cooling_mode_calc_col].mean() * 100, 2)

        # Econ plus mech cooling mode runtime stats
        delta_econ_clg = df[self.econ_plus_mech_cooling_mode_calc_col].index.to_series().diff()
        total_hours_econ_clg = (delta_econ_clg * df[self.econ_plus_mech_cooling_mode_calc_col]).sum() / pd.Timedelta(hours(1))
        percent_econ_clg = round(df[self.econ_plus_mech_cooling_mode_calc_col].mean() * 100, 2)

        # Mech clg mode runtime stats
        delta_clg = df[self.mech_cooling_only_mode_calc_col].index.to_series().diff()
        total_hours_clg = (delta_clg * df[self.mech_cooling_only_mode_calc_col]).sum() / pd.Timedelta(hours(1))
        percent_clg = round(df[self.mech_cooling_only_mode_calc_col].mean() * 100, 2)

        return (
            total_days_all_data,
            total_hours_all_data,
            hours_fc4_mode,
            percent_true_fc4,
            percent_false_fc4,
            percent_clg,
            percent_econ_clg,
            percent_econ,
            percent_heating,
            total_hours_heating,
            total_hours_econ,
            total_hours_econ_clg,
            total_hours_clg
        )

    def create_hist_plot(self, df: pd.DataFrame, output_col: str = None) -> BytesIO:
        if output_col is None:
            output_col = "fc4_flag"

        # Calculate dataset statistics
        df["hour_of_the_day_fc4"] = df.index.hour.where(df[output_col] == 1)

        # Make hist plots fc4
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc4.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 4 is TRUE")

        hist_plot_image = BytesIO()
        fig.savefig(hist_plot_image, format="png")
        hist_plot_image.seek(0)
        plt.close(fig)  # Close the figure

        return hist_plot_image

    def create_report(self, path: str, df: pd.DataFrame, output_col: str = None, report_name: str = "report_fc4.docx") -> None:
        if output_col is None:
            output_col = "fc4_flag"

        print(f"Starting {path} docx report!")
        sys.stdout.flush()
        document = Document()
        document.add_heading("Fault Condition Four Report", 0)

        p = document.add_paragraph(
            """Fault condition four of ASHRAE Guideline 36 is related to flagging AHU control programming that is hunting between heating, economizing, economizing plus mechanical cooling, and mechanical cooling operating states. This fault diagnostic does NOT flag simultaneous heating and cooling, just excessive cycling between the states or operating modes the AHU maybe going in and out of. Fault condition four equation as defined by ASHRAE:"""
        )

        image_path = pkg_resources.resource_filename('open_fdd', 'air_handling_unit/images/fc4_definition.png')
        document.add_picture(image_path, width=Inches(6))
        document.add_heading("Dataset Plot", level=2)

        plot_image = self.create_plot(df, output_col=output_col)
        document.add_picture(plot_image, width=Inches(6))

        document.add_heading("Dataset Statistics", level=2)

        (
            total_days_all_data,
            total_hours_all_data,
            hours_fc4_mode,
            percent_true_fc4,
            percent_false_fc4,
            percent_clg,
            percent_econ_clg,
            percent_econ,
            percent_heating,
            total_hours_heating,
            total_hours_econ,
            total_hours_econ_clg,
            total_hours_clg
        ) = self.summarize_fault_times(df, output_col=output_col)

        stats = [
            f"Total time in days calculated in dataset: {total_days_all_data}",
            f"Total time in hours calculated in dataset: {total_hours_all_data}",
            f"Total time in hours for when fault flag is True: {hours_fc4_mode}",
            f"Percent of time in the dataset when the fault flag is True: {percent_true_fc4}%",
            f"Percent of time in the dataset when the fault flag is False: {percent_false_fc4}%",
            f"Total time in hours while AHU is in a heating mode: {total_hours_heating}",
            f"Total percent time in while AHU is in a heating mode: {percent_heating}%",
            f"Total time in hours while AHU is in a economizing mode: {total_hours_econ}",
            f"Total percent time in while AHU is in a economizing mode: {percent_econ}%",
            f"Total time in hours while AHU is in a economizing plus mechanical cooling mode: {total_hours_econ_clg}",
            f"Total percent time in while AHU is in a economizing plus mechanical cooling mode: {percent_econ_clg}%",
            f"Total time in hours while AHU is in a mechanical cooling mode: {total_hours_clg}",
            f"Total percent time in while AHU is in a mechanical cooling mode: {percent_clg}%"
        ]

        for stat in stats:
            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(stat)

        # if there are no faults, skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = self.create_hist_plot(df, output_col=output_col)
            document.add_picture(histogram_plot_image, width=Inches(6))

            paragraph = document.add_paragraph()
            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'Fault condition 4 is True because of excessive cycling between different control system operation modes.'
            )
        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")
            sys.stdout.flush()
            paragraph = document.add_paragraph()
            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.'
            )

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if fc_max_faults_found >= self.delta_os_max:
            paragraph.add_run(
                f'The AHU control system needs tuning to reduce control loop hunting for setpoints. Its hunting or overshooting setpoints which can cause AHU systems to be oscillating (most likely too fast) between heating and cooling modes without never settling out. Low load conditions can also cause excessive cycling if heating or cooling setpoints are met very fast. Verify that the times when this fault is flagged that no occupant comfort issues persist. Fixing this fault may also improve energy efficiency and extend the mechanical equipment life span with the prevention of excessive cycling especially cooling compressors.'
            )
        else:
            paragraph.add_run(
                f'No control system tuning appears to be needed for the operating conditions of this AHU.'
            )

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        document.save(f"{path}/{report_name}")

    def print_summary_statistics(self, df: pd.DataFrame, output_col: str = None):
        (
            total_days_all_data,
            total_hours_all_data,
            hours_fc4_mode,
            percent_true_fc4,
            percent_false_fc4,
            percent_clg,
            percent_econ_clg,
            percent_econ,
            percent_heating,
            total_hours_heating,
            total_hours_econ,
            total_hours_econ_clg,
            total_hours_clg
        ) = self.summarize_fault_times(df, output_col=output_col)

        summary_labels = [
            "Total time in days calculated in dataset",
            "Total time in hours calculated in dataset",
            "Total time in hours for when fault flag is True",
            "Percent of time in the dataset when the fault flag is True",
            "Percent of time in the dataset when the fault flag is False",
            "Total percent time while AHU is in a mechanical cooling mode",
            "Total percent time while AHU is in an economizing plus mechanical cooling mode",
            "Total percent time while AHU is in an economizing mode",
            "Total percent time while AHU is in a heating mode",
            "Total time in hours while AHU is in a heating mode",
            "Total time in hours while AHU is in an economizing mode",
            "Total time in hours while AHU is in an economizing plus mechanical cooling mode",
            "Total time in hours while AHU is in a mechanical cooling mode"
        ]

        for label, stat in zip(summary_labels, [
            total_days_all_data,
            total_hours_all_data,
            hours_fc4_mode,
            percent_true_fc4,
            percent_false_fc4,
            percent_clg,
            percent_econ_clg,
            percent_econ,
            percent_heating,
            total_hours_heating,
            total_hours_econ,
            total_hours_econ_clg,
            total_hours_clg
        ]):
            print(f"{label}: {stat}")
            sys.stdout.flush()
