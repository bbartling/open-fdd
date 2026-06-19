"""Edge RCx DOCX — professional template with screenshot placeholders + programmatic data."""

from __future__ import annotations

from datetime import datetime, timezone
from typing import Any

from open_fdd.reports.fault_hours import aggregate_fault_hours
from open_fdd.reports.rcx_placeholders import (
    SCREENSHOT_LABEL,
    chart_placeholder_spec,
    disabled_chart_notes,
    rule_sensor_placeholder,
)

DEFAULT_SECTIONS = [
    "executive_summary",
    "mechanical_summary",
    "trend_charts",
    "fault_analytics",
    "ahu_analytics",
    "vav_analytics",
    "analyst_insights",
    "runtime_analytics",
    "model_health",
    "fdd_rule_trends",
    "recommendations",
    "appendix_faults",
    "appendix_missing_roles",
]


def _recommendations(rows: list[dict[str, Any]], warnings: list[str]) -> list[str]:
    recs: list[str] = []
    for row in rows[:5]:
        eq = row.get("equipment") or "equipment"
        name = row.get("fault_name") or row.get("fault_code") or "fault"
        recs.append(f"Review {eq} for {name} — ~{row.get('elapsed_hours', 0)} h elapsed in lookback.")
    for w in warnings[:5]:
        if "flatline" in w.lower():
            recs.append("Verify sensor wiring/calibration if flatline faults persist.")
        elif "static" in w.lower() or "duct" in w.lower():
            recs.append("Review duct static pressure reset if pressure exceeds setpoint for extended periods.")
        elif "comfort" in w.lower() or "zone" in w.lower():
            recs.append("Investigate VAV zones with high comfort fault hours and saturated dampers.")
    if not recs:
        recs.append("No critical recommendations from current snapshot; continue monitoring.")
    return recs[:12]


def _screenshot_placeholder(doc, *, title: str, subtitle: str = "", instruction: str = "") -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    if instruction:
        doc.add_paragraph(instruction)
    doc.add_paragraph()
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = box.add_run(SCREENSHOT_LABEL)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(title).italic = True
    if subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(subtitle).font.size = Pt(9)


def _kv_table(doc, rows: list[tuple[str, str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = str(k)
        table.rows[i].cells[1].text = str(v)


def _grid_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    if not headers:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(cell)


def build_rcx_docx(
    *,
    site_id: str,
    site_name: str,
    window: dict[str, str | None],
    fault_rows: list[dict[str, Any]],
    overview: dict[str, Any] | None = None,
    sections: list[str] | None = None,
    charts: list[str] | None = None,
    warnings: list[str] | None = None,
    equipment_notes: list[dict[str, Any]] | None = None,
    chart_previews: list[dict[str, Any]] | None = None,
    report_context: dict[str, Any] | None = None,
    equipment_bundle: dict[str, Any] | None = None,
    equipment_charts: list[dict[str, Any]] | None = None,
    available_charts: list[dict[str, Any]] | None = None,
    disabled_charts: list[dict[str, Any]] | None = None,
    ai_insights: dict[str, Any] | None = None,
) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt

    enabled_sections = set(sections or DEFAULT_SECTIONS)
    ctx = report_context or {}
    assigned_rules = ctx.get("assigned_rules") if isinstance(ctx.get("assigned_rules"), list) else []
    motor_rows = ctx.get("motor_runtime") if isinstance(ctx.get("motor_runtime"), list) else []
    overrides = ctx.get("overrides") if isinstance(ctx.get("overrides"), dict) else {}
    override_list = overrides.get("overrides") if isinstance(overrides.get("overrides"), list) else []

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Retro-Commissioning Analytics Report", level=0)
    doc.add_paragraph(f"Building: {site_name}")
    doc.add_paragraph(f"Site ID: {site_id}")
    if equipment_bundle:
        doc.add_paragraph(
            f"Equipment report: {equipment_bundle.get('label') or equipment_bundle.get('bundle_id') or '—'}"
        )
    if window.get("start") or window.get("end"):
        doc.add_paragraph(f"Analysis window: {window.get('start') or '—'} → {window.get('end') or '—'}")
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph(
        "Read-only analytics export. Paste Plotly trend screenshots into the marked placeholders. "
        "No BACnet writes or overrides are performed by this report."
    )

    ov = overview or {}
    by_sev = aggregate_fault_hours(fault_rows, group_by="severity")
    by_eq = aggregate_fault_hours(fault_rows, group_by="equipment")
    by_code = aggregate_fault_hours(fault_rows, group_by="fault_code")

    if "executive_summary" in enabled_sections:
        doc.add_heading("Executive summary", level=1)
        _kv_table(
            doc,
            [
                ("Active faults", str(ov.get("active_faults", len(fault_rows)))),
                ("Total elapsed fault hours (est.)", str(ov.get("total_fault_hours", 0))),
                ("BACnet P8 overrides", str(ctx.get("override_count", 0))),
                ("Assigned FDD rules", str(len(assigned_rules))),
                ("Motor / fan points modeled", str(len(motor_rows))),
            ],
        )
        if by_eq:
            doc.add_paragraph(
                f"Worst equipment: {by_eq[0].get('group')} (~{by_eq[0].get('elapsed_hours')} h)"
            )
        for w in (warnings or [])[:6]:
            doc.add_paragraph(f"• {w}")

    if "mechanical_summary" in enabled_sections:
        doc.add_heading("Mechanical summary", level=1)
        mech = ov.get("mechanical_summary") if isinstance(ov.get("mechanical_summary"), dict) else {}
        narrative = str(mech.get("narrative") or "").strip()
        if narrative:
            for para in narrative.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        counts = mech.get("counts") if isinstance(mech.get("counts"), dict) else {}
        if counts:
            doc.add_paragraph(
                f"Inventory: {counts.get('ahus', counts.get('ahu', '—'))} AHU(s), "
                f"{counts.get('vavs', counts.get('vav', '—'))} VAV(s), "
                f"{counts.get('zones', '—')} zone(s)."
            )
        elif not narrative:
            doc.add_paragraph("Mechanical summary unavailable for this Edge snapshot.")

    if "trend_charts" in enabled_sections:
        doc.add_heading("Trend charts (engineer screenshots)", level=1)
        doc.add_paragraph(
            "BRICK model resolves historian columns per equipment type (AHU, VAV, zone, plant). "
            "Open RCx Report Builder → Render chart gallery, or Trend plot with the listed columns. "
            "Snip each chart and paste into the placeholder below it."
        )
        bundle_chart_ids = equipment_bundle.get("chart_ids") if isinstance(equipment_bundle, dict) else None
        chart_ids = list(charts or bundle_chart_ids or ["ahu_sat_vs_setpoint", "ahu_duct_static_vs_setpoint", "vav_zone_temp"])
        for cid in chart_ids:
            spec = chart_placeholder_spec(
                str(cid),
                equipment_charts=equipment_charts,
                catalog=available_charts,
                chart_previews=chart_previews,
            )
            if spec.get("equipment_type") and spec.get("equipment_type") != "building":
                doc.add_paragraph(f"System type: {spec.get('equipment_type')}")
            _screenshot_placeholder(
                doc,
                title=str(spec.get("title") or cid),
                subtitle=str(spec.get("subtitle") or ""),
                instruction=str(spec.get("instruction") or ""),
            )
        gap_notes = disabled_chart_notes(disabled_charts)
        if gap_notes:
            doc.add_paragraph("Charts not included (model / historian gaps):")
            for note in gap_notes[:12]:
                doc.add_paragraph(f"• {note}")

    if "fault_analytics" in enabled_sections:
        doc.add_heading("Fault analytics", level=1)
        if by_sev:
            doc.add_paragraph("Fault hours by severity:")
            for row in by_sev:
                doc.add_paragraph(f"  {row.get('group')}: {row.get('elapsed_hours')} h")
        if by_eq:
            doc.add_paragraph("Fault hours by equipment:")
            for row in by_eq[:15]:
                doc.add_paragraph(f"  {row.get('group')}: {row.get('elapsed_hours')} h")
        if by_code:
            doc.add_paragraph("Fault hours by code:")
            for row in by_code[:15]:
                doc.add_paragraph(f"  {row.get('group')}: {row.get('elapsed_hours')} h")
        doc.add_paragraph("Active fault detail:")
        for row in fault_rows[:30]:
            doc.add_paragraph(
                f"{row.get('equipment')} ({row.get('equipment_type') or '—'}) — "
                f"{row.get('severity')}: {row.get('fault_name') or row.get('fault_code')} "
                f"[~{row.get('elapsed_hours')} h, "
                f"{row.get('samples_flagged')}/{row.get('samples_evaluated')} samples]"
            )

    if "ahu_analytics" in enabled_sections:
        doc.add_heading("AHU analytics", level=1)
        ahu_rows = [r for r in fault_rows if str(r.get("equipment_type") or "").upper() in ("AHU", "RTU")]
        if ahu_rows:
            for row in ahu_rows[:10]:
                doc.add_paragraph(
                    f"{row.get('equipment')}: {row.get('fault_name')} — ~{row.get('elapsed_hours')} h"
                )
        else:
            doc.add_paragraph("No AHU fault rows in selected window.")
        _screenshot_placeholder(doc, title="AHU overview trend", subtitle="SAT / static / OA-MAT-RAT")

    if "vav_analytics" in enabled_sections:
        doc.add_heading("VAV zone analytics", level=1)
        vav_rows = [r for r in fault_rows if "VAV" in str(r.get("equipment_type") or "").upper()]
        if vav_rows:
            for row in vav_rows[:15]:
                doc.add_paragraph(
                    f"{row.get('equipment')}: {row.get('fault_name')} — ~{row.get('elapsed_hours')} h"
                )
        else:
            doc.add_paragraph("No VAV fault rows in selected window.")
        _screenshot_placeholder(doc, title="VAV zone comfort trend", subtitle="Zone temp vs setpoints")

    if "runtime_analytics" in enabled_sections:
        doc.add_heading("Runtime analytics (motor / fan / pump)", level=1)
        if motor_rows:
            doc.add_paragraph("Estimated run-hours from BRICK motor/fan/pump points (PyArrow historian):")
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Equipment"
            hdr[1].text = "Point"
            hdr[2].text = "Column"
            hdr[3].text = "Hours (window)"
            hdr[4].text = "Weekly est. (h)"
            for m in motor_rows[:20]:
                row = table.add_row().cells
                row[0].text = str(m.get("equipment_name") or "—")
                row[1].text = str(m.get("label") or "—")
                row[2].text = str(m.get("column") or "—")
                row[3].text = str(m.get("runtime_hours") or 0)
                row[4].text = str(m.get("weekly_hours_est") or 0)
        else:
            doc.add_paragraph(
                "No motor/fan/pump points detected in BRICK model. "
                "Tag Fan_Status, Pump_Status, or VFD speed sensors to enable runtime estimates."
            )
        _screenshot_placeholder(doc, title="Motor / fan runtime trend", subtitle="Weekly run-hours verification")

    if "model_health" in enabled_sections:
        doc.add_heading("BACnet / model health", level=1)
        mh = ov.get("model_health") if isinstance(ov.get("model_health"), dict) else {}
        _kv_table(
            doc,
            [
                (label, str(mh[key]))
                for label, key in (
                    ("Devices", "device_count"),
                    ("Points", "point_count"),
                    ("Equipment", "equipment_count"),
                    ("Stale points", "stale_point_count"),
                )
                if key in mh
            ],
        )
        if override_list:
            p8_rows = [
                row
                for row in override_list
                if isinstance(row, dict) and (row.get("operator_p8") or row.get("operator_override"))
            ]
            if not p8_rows:
                p8_rows = [row for row in override_list if isinstance(row, dict)]
            doc.add_paragraph(f"Active BACnet operator overrides (P8): {len(p8_rows)}")
            by_device: dict[str, int] = {}
            for ov_row in p8_rows:
                dev = str(
                    ov_row.get("device_label")
                    or ov_row.get("device_name")
                    or ov_row.get("device_id")
                    or ov_row.get("device")
                    or ov_row.get("device_instance")
                    or "—"
                )
                by_device[dev] = by_device.get(dev, 0) + 1
            if by_device:
                doc.add_paragraph("Devices with P8 overrides:")
                for dev, cnt in sorted(by_device.items(), key=lambda x: (-x[1], x[0]))[:16]:
                    doc.add_paragraph(f"  {dev}: {cnt} point(s)")

            preview_rows = ctx.get("override_preview") if isinstance(ctx.get("override_preview"), list) else p8_rows[:8]
            table_rows: list[list[str]] = []
            for ov_row in preview_rows[:8]:
                if not isinstance(ov_row, dict):
                    continue
                table_rows.append(
                    [
                        str(ov_row.get("device_instance") or ov_row.get("device") or "—"),
                        str(ov_row.get("device_address") or "—"),
                        str(ov_row.get("object_name") or ov_row.get("point") or ov_row.get("object") or "—"),
                        str(ov_row.get("value_text") or ov_row.get("value") or "—"),
                        str(ov_row.get("scanned_at") or "—")[:19],
                    ]
                )
            if table_rows:
                doc.add_paragraph("Current P8 overrides (up to 8 shown on dashboard):")
                _grid_table(
                    doc,
                    ["Device", "Address", "Point", "Value", "Last scan (UTC)"],
                    table_rows,
                )

            by_device_rows = ctx.get("override_by_device") if isinstance(ctx.get("override_by_device"), list) else []
            device_table: list[list[str]] = []
            for dev_row in by_device_rows[:24]:
                if not isinstance(dev_row, dict):
                    continue
                device_table.append(
                    [
                        str(dev_row.get("device_instance") or "—"),
                        str(dev_row.get("device_address") or "—"),
                        str(dev_row.get("operator_override_count") or 0),
                        str(dev_row.get("total_override_count") or 0),
                        str(dev_row.get("last_scanned_at") or "—")[:19],
                    ]
                )
            if device_table:
                doc.add_paragraph("Override scan coverage by BACnet device:")
                _grid_table(
                    doc,
                    ["Device", "Address", "P8 points", "All priorities", "Last scanned (UTC)"],
                    device_table,
                )

            scan_health = ctx.get("override_scan_health") if isinstance(ctx.get("override_scan_health"), dict) else {}
            scan = ctx.get("override_scan") if isinstance(ctx.get("override_scan"), dict) else {}
            if scan_health or scan:
                doc.add_paragraph("Hourly supervisory override scan:")
                _kv_table(
                    doc,
                    [
                        (label, str(value))
                        for label, value in (
                            ("Scan health", scan_health.get("status") or "—"),
                            ("Detail", scan_health.get("detail") or "—"),
                            ("Interval (s)", scan.get("scan_interval_s") or "—"),
                            ("Devices in rotation", scan.get("device_count") or "—"),
                            ("Full rotation (h)", scan.get("full_rotation_hours") or "—"),
                            ("Last scan device", scan.get("last_scan_device") or "—"),
                            ("Last scan (UTC)", str(scan.get("last_scan_at") or "—")[:19]),
                        )
                        if value not in (None, "", "—")
                    ],
                )
            for ov_row in p8_rows[:12]:
                if isinstance(ov_row, dict):
                    dev = str(
                        ov_row.get("device_label")
                        or ov_row.get("device_name")
                        or ov_row.get("device_id")
                        or ov_row.get("device")
                        or ov_row.get("device_instance")
                        or "—"
                    )
                    point = str(ov_row.get("object_name") or ov_row.get("point") or ov_row.get("object") or "")
                    doc.add_paragraph(
                        f"  {dev} · {point} priority {ov_row.get('priority') or ov_row.get('priority_level') or '—'}"
                    )

    if "fdd_rule_trends" in enabled_sections:
        doc.add_heading("FDD rule trend screenshots", level=1)
        doc.add_paragraph(
            "For each assigned Rule Lab rule, paste a trend screenshot for the bound sensor columns "
            "(required even when no fault is active in the lookback window)."
        )
        if assigned_rules:
            for rule in assigned_rules:
                if not isinstance(rule, dict):
                    continue
                doc.add_heading(str(rule.get("rule_name") or rule.get("rule_id")), level=2)
                doc.add_paragraph(
                    f"Fault code: {rule.get('fault_code') or '—'} · Severity: {rule.get('severity') or '—'}"
                )
                for sensor in rule_sensor_placeholder(rule):
                    label = sensor.get("label") or "sensor"
                    col = sensor.get("column") or ""
                    brick = sensor.get("brick_type") or ""
                    doc.add_paragraph(f"Sensor: {label} ({col}) {brick}".strip())
                    doc.add_paragraph(sensor.get("instruction") or "")
                    _screenshot_placeholder(
                        doc,
                        title=f"{rule.get('rule_name')} — {label}",
                        subtitle=f"Historian column: {col}" if col else "Bind points in Model & assignments",
                    )
        else:
            doc.add_paragraph("No enabled Rule Lab rules found for this site.")

    if "recommendations" in enabled_sections:
        doc.add_heading("Recommendations", level=1)
        for rec in _recommendations(fault_rows, warnings or []):
            doc.add_paragraph(f"• {rec}")

    if "analyst_insights" in enabled_sections:
        doc.add_heading("AI analyst assessment", level=1)
        ai = ai_insights if isinstance(ai_insights, dict) else {}
        source = str(ai.get("source") or "deterministic")
        doc.add_paragraph(
            f"Automated interpretation of PyArrow historian trends, BRICK model context, "
            f"Rule Lab bindings, and BACnet override scan ({source} engine)."
        )

        paragraphs = [str(p).strip() for p in (ai.get("paragraphs") or []) if str(p).strip()]
        if paragraphs:
            for para in paragraphs[:10]:
                doc.add_paragraph(para)
        else:
            narratives: list[str] = []
            for prev in chart_previews or []:
                if not isinstance(prev, dict):
                    continue
                text = str(prev.get("narrative") or "").strip()
                if text:
                    title = str(prev.get("title") or prev.get("chart_id") or "Chart")
                    narratives.append(f"{title}: {text}")
            if narratives:
                doc.add_paragraph("Programmatic chart narratives from the selected window:")
                for para in narratives[:8]:
                    doc.add_paragraph(f"• {para}")
            else:
                doc.add_paragraph(
                    "Plain-language interpretation — add field notes here after reviewing trends and fault evidence."
                )

        chart_insights = ai.get("chart_insights") if isinstance(ai.get("chart_insights"), list) else []
        if chart_insights:
            doc.add_heading("Trend plot interpretation", level=2)
            for ci in chart_insights[:8]:
                if not isinstance(ci, dict):
                    continue
                title = str(ci.get("title") or "Chart")
                narrative = str(ci.get("narrative") or "").strip()
                if narrative:
                    doc.add_paragraph(f"{title}: {narrative}")
                for bullet in (ci.get("stats_bullets") or [])[:4]:
                    doc.add_paragraph(f"  • {bullet}")

        rule_notes = ai.get("rule_assessments") if isinstance(ai.get("rule_assessments"), list) else []
        if rule_notes:
            doc.add_heading("FDD rule lab assessment", level=2)
            for note in rule_notes[:8]:
                doc.add_paragraph(f"• {note}")

        override_notes = ai.get("override_notes") if isinstance(ai.get("override_notes"), list) else []
        if override_notes:
            doc.add_heading("BACnet override hygiene", level=2)
            for note in override_notes[:6]:
                doc.add_paragraph(note)

        err = str(ai.get("error") or "").strip()
        if err and source != "ollama":
            doc.add_paragraph(f"(AI engine note: {err[:200]})")

    if "smoke_validation" in enabled_sections:
        doc.add_heading("Half-hour smoke validation (bench 5007)", level=1)
        smoke = ctx.get("smoke_validation") if isinstance(ctx.get("smoke_validation"), dict) else {}
        doc.add_paragraph(
            "Automated bench smoke: PyArrow vs DataFusion SQL parity on device 5007, "
            "API/UI health probes, and BACnet P8 supervisory override scan mechanism."
        )
        _kv_table(
            doc,
            [
                ("Smoke PASS", str(smoke.get("pass"))),
                ("Mode", str(smoke.get("mode") or "—")),
                ("Health probe cycles", str(smoke.get("health_probe_count") or 0)),
                ("Override scan OK", str(smoke.get("override_scan_ok"))),
                ("Generated", str(smoke.get("generated_at") or "—")[:19]),
            ],
        )
        issues = smoke.get("issues") if isinstance(smoke.get("issues"), list) else []
        if issues:
            doc.add_paragraph("Smoke issues:")
            for issue in issues[:12]:
                doc.add_paragraph(f"• {issue}")
        else:
            doc.add_paragraph("No smoke issues recorded.")

        ov_scan = smoke.get("override_scan") if isinstance(smoke.get("override_scan"), dict) else {}
        ov_status = ov_scan.get("status") if isinstance(ov_scan.get("status"), dict) else {}
        if ov_status:
            doc.add_heading("BACnet override scan (hourly rotation)", level=2)
            _kv_table(
                doc,
                [
                    ("Scan interval (s)", str(ov_status.get("scan_interval_s") or "—")),
                    ("Devices in rotation", str(ov_status.get("device_count") or "—")),
                    ("Cursor", str(ov_status.get("cursor") or "—")),
                    ("Last scan device", str(ov_status.get("last_scan_device") or "—")),
                    ("Full rotation (h)", str(ov_status.get("full_rotation_hours") or "—")),
                    ("Operator priority", f"P{ov_status.get('operator_priority') or 8}"),
                ],
            )

        final_hp = smoke.get("health_probes_final") if isinstance(smoke.get("health_probes_final"), dict) else {}
        probes = final_hp.get("probes") if isinstance(final_hp.get("probes"), list) else []
        if probes:
            doc.add_heading("Health probes (final cycle)", level=2)
            for probe in probes[:8]:
                if not isinstance(probe, dict):
                    continue
                mark = "OK" if probe.get("ok") else "FAIL"
                doc.add_paragraph(f"{mark} — {probe.get('name')}: {probe.get('detail') or ''}")

        flagged = smoke.get("last_flagged") if isinstance(smoke.get("last_flagged"), dict) else {}
        if flagged:
            doc.add_heading("PyArrow vs SQL smoke rules (last cycle)", level=2)
            for rid, cnt in sorted(flagged.items()):
                if str(rid).startswith("smoke-paired"):
                    doc.add_paragraph(f"  {rid}: flagged={cnt}")

    if "appendix_faults" in enabled_sections:
        doc.add_heading("Appendix: fault table", level=1)
        for row in fault_rows:
            doc.add_paragraph(
                f"{row.get('fault_code')} | {row.get('equipment')} | {row.get('severity')} | "
                f"{row.get('elapsed_hours')} h"
            )

    if "appendix_missing_roles" in enabled_sections:
        doc.add_heading("Appendix: missing point roles", level=1)
        missing = ov.get("missing_roles") if isinstance(ov.get("missing_roles"), list) else []
        gap_notes = disabled_chart_notes(disabled_charts)
        if missing:
            doc.add_paragraph("Model issues from mechanical summary:")
            for m in missing[:40]:
                doc.add_paragraph(f"• {m}")
        if gap_notes:
            doc.add_paragraph("Charts disabled due to missing BRICK roles or historian data:")
            for note in gap_notes[:40]:
                doc.add_paragraph(f"• {note}")
        if not missing and not gap_notes:
            doc.add_paragraph("None recorded — all selected chart roles resolved.")

    if equipment_notes:
        doc.add_heading("Equipment notes", level=1)
        for note in equipment_notes[:20]:
            doc.add_paragraph(str(note.get("text") or note))

    import io

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
