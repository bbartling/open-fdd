"""
Generate Word (.docx) reports for FDD findings.

Target audience: non-technical facilities managers. Report includes
methodology, executive summary, charts, fault summary table, and recommendations.

Requires: python-docx  (pip install python-docx)
"""

from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd


def events_from_dataframe(
    df: pd.DataFrame,
    events: List[Tuple[int, int, str]],
    timestamp_col: str = "timestamp",
) -> List[Dict[str, Any]]:
    """
    Convert (start_iloc, end_iloc, flag_name) events to dicts with start/end timestamps.

    Use with fault_viz.all_fault_events() output.
    """
    result = []
    ts = df[timestamp_col] if timestamp_col in df.columns else df.index
    for start_iloc, end_iloc, flag_name in events:
        start_ts = ts.iloc[start_iloc] if start_iloc < len(ts) else None
        end_ts = ts.iloc[end_iloc] if end_iloc < len(ts) else None
        duration = end_iloc - start_iloc + 1
        result.append(
            {
                "flag": flag_name,
                "start": str(start_ts) if start_ts is not None else "—",
                "end": str(end_ts) if end_ts is not None else "—",
                "duration_samples": duration,
            }
        )
    return result


def _get_docx():
    """Lazy import to avoid hard dependency."""
    try:
        from docx import Document
        from docx.shared import Inches

        return Document, Inches
    except ImportError:
        raise ImportError(
            "python-docx is required for docx reports. Install with: pip install python-docx"
        )


def events_to_summary_table(
    events: List[Dict[str, Any]],
    flag_col: str = "flag",
    start_col: str = "start",
    end_col: str = "end",
    duration_col: str = "duration_samples",
) -> pd.DataFrame:
    """
    Convert event list to a summary DataFrame for table display.

    Args:
        events: List of event dicts with flag, start, end, duration_samples
        flag_col, start_col, end_col, duration_col: keys in each event dict

    Returns:
        DataFrame with columns: Flag, Start, End, Duration (samples)
    """
    rows = []
    for e in events:
        rows.append(
            {
                "Flag": e.get(flag_col, e.get("flag_name", "—")),
                "Start": e.get(start_col, "—"),
                "End": e.get(end_col, "—"),
                "Duration (samples)": e.get(duration_col, "—"),
            }
        )
    return pd.DataFrame(rows)


def build_report(
    result: pd.DataFrame,
    events: List[Dict[str, Any]],
    summaries: Dict[str, Dict[str, Any]],
    output_dir: Path,
    equipment_name: str = "Equipment",
    plot_paths: Optional[List[str]] = None,
    methodology: Optional[str] = None,
    executive_summary: Optional[str] = None,
    recommendations: Optional[List[str]] = None,
) -> Path:
    """
    Build a Word report with charts, tables, and summaries.

    Args:
        result: DataFrame with fault flags (from RuleRunner)
        events: List of fault event dicts (flag, start, end, duration_samples)
        summaries: Dict[flag_col, summary_dict] from fault_report.summarize_all_faults
        output_dir: Directory for report and plots
        equipment_name: e.g. "AHU7", "Chiller Plant"
        plot_paths: Paths to plot images to embed (e.g. from zoom_on_event)
        methodology: Optional methodology paragraph
        executive_summary: Optional pre-written summary
        recommendations: Optional list of recommended actions

    Returns:
        Path to saved .docx file
    """
    Document, Inches = _get_docx()
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    from datetime import datetime

    date_str = datetime.now().strftime("%Y-%m-%d")
    report_name = f"FDD_Report_{equipment_name}_{date_str}.docx"
    report_path = output_dir / report_name

    doc = Document()
    doc.add_heading("HVAC Fault Detection Report", 0)
    doc.add_paragraph(f"Equipment: {equipment_name}")
    doc.add_paragraph(f"Report Date: {date_str}")

    # Methodology
    doc.add_heading("Methodology", level=1)
    if methodology:
        doc.add_paragraph(methodology)
    else:
        doc.add_paragraph(
            "This report uses config-driven fault detection rules (YAML) run against "
            "time-series HVAC sensor data. Rules flag conditions such as stuck sensors, "
            "out-of-range values, and equipment-specific faults. Results were refined "
            "to reduce false positives before final reporting."
        )

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    if executive_summary:
        doc.add_paragraph(executive_summary)
    else:
        total_events = len(events)
        flag_counts = {}
        for e in events:
            f = e.get("flag", e.get("flag_name", "—"))
            flag_counts[f] = flag_counts.get(f, 0) + 1
        count_str = ", ".join(f"{k}: {v}" for k, v in sorted(flag_counts.items()))
        doc.add_paragraph(
            f"Total fault events identified: {total_events}. "
            f"Breakdown by type: {count_str}. "
            "See fault summary table and charts for details."
        )

    # Charts
    if plot_paths:
        doc.add_heading("Fault Event Charts", level=1)
        for p in plot_paths:
            path = Path(p)
            if path.exists():
                doc.add_paragraph(path.stem)
                doc.add_picture(str(path), width=Inches(6))

    # Fault Summary Table
    doc.add_heading("Fault Summary Table", level=1)
    table_df = events_to_summary_table(events)
    if not table_df.empty:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, col in enumerate(["Flag", "Start", "End", "Duration (samples)"]):
            hdr[i].text = col
        for _, row in table_df.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row.get("Flag", ""))
            cells[1].text = str(row.get("Start", ""))
            cells[2].text = str(row.get("End", ""))
            cells[3].text = str(row.get("Duration (samples)", ""))
    else:
        doc.add_paragraph("No fault events recorded.")

    # Key stats from summaries
    if summaries:
        doc.add_heading("Fault Analytics", level=1)
        for flag, s in summaries.items():
            doc.add_paragraph(f"{flag}:")
            for k, v in s.items():
                if k not in ("error",):
                    doc.add_paragraph(f"  {k}: {v}", style="List Bullet")

    # Recommendations
    doc.add_heading("Recommended Actions", level=1)
    if recommendations:
        for r in recommendations:
            doc.add_paragraph(r, style="List Bullet")
    else:
        doc.add_paragraph(
            "Review flagged events with on-site staff. Verify sensors and "
            "equipment operation during fault periods. Address persistent faults "
            "per facilities maintenance procedures."
        )

    doc.save(str(report_path))
    return report_path
