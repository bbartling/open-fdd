"""Fixture-driven RCx DOCX report validation."""

from __future__ import annotations

import io
import json
from pathlib import Path

import pytest

REPO = Path(__file__).resolve().parents[3]
FIXTURE_DIR = REPO / "tests" / "fixtures" / "rcx"


def _doc_text(doc) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


@pytest.fixture
def fault_rows():
    return json.loads((FIXTURE_DIR / "fault_rows.json").read_text(encoding="utf-8"))


@pytest.fixture
def report_context():
    return json.loads((FIXTURE_DIR / "report_context.json").read_text(encoding="utf-8"))


@pytest.fixture
def chart_catalog():
    return json.loads((FIXTURE_DIR / "chart_catalog.json").read_text(encoding="utf-8"))


def test_rcx_docx_fixture_sections(fault_rows, report_context, chart_catalog):
    pytest.importorskip("docx")
    from docx import Document
    from open_fdd.reports.rcx_docx import DEFAULT_SECTIONS, SCREENSHOT_LABEL, build_rcx_docx

    blob = build_rcx_docx(
        site_id="acme",
        site_name="Acme Lab Building",
        window={"start": "2026-01-01T00:00:00Z", "end": "2026-01-02T00:00:00Z"},
        fault_rows=fault_rows,
        overview={
            "active_faults": 2,
            "total_fault_hours": 4.5,
            "missing_roles": ["AHU-C: supply_air_temperature_setpoint"],
            "mechanical_summary": {
                "narrative": "Single AHU serves lab zones.",
                "counts": {"ahu": 1, "vav": 2, "zones": 4},
            },
        },
        sections=DEFAULT_SECTIONS,
        report_context=report_context,
        equipment_bundle={"bundle_id": "ahu-c", "label": "AHU-C", "chart_ids": ["ahu_sat_vs_setpoint"]},
        equipment_charts=chart_catalog["equipment_charts"],
        available_charts=chart_catalog["available_charts"],
        disabled_charts=chart_catalog["disabled_charts"],
        chart_previews=[
            {
                "chart_id": "ahu_sat_vs_setpoint",
                "title": "SAT vs setpoint",
                "narrative": "SAT tracking looks reasonable at this resolution.",
            }
        ],
    )
    doc = Document(io.BytesIO(blob))
    text = _doc_text(doc)
    assert "Acme Lab Building" in text
    assert "Site ID: acme" in text
    assert "AHU-C" in text
    assert "Retro-Commissioning" in text
    assert "Runtime analytics" in text
    assert "Recommendations" in text
    assert SCREENSHOT_LABEL in text
    assert "space_temperature_local" in text
    assert "supply_fan_status_local" in text
    assert "Weekly est" in text
    assert "SAT tracking looks reasonable" in text
    assert "zone_cooling_setpoint" in text
    assert "Traceback" not in text


def test_rcx_docx_active_fault_table(fault_rows):
    pytest.importorskip("docx")
    from docx import Document
    from open_fdd.reports.rcx_docx import build_rcx_docx

    blob = build_rcx_docx(
        site_id="demo",
        site_name="Demo Site",
        window={"start": "2026-01-01", "end": "2026-01-02"},
        fault_rows=fault_rows,
        sections=["executive_summary", "fault_analytics", "appendix_faults"],
    )
    doc = Document(io.BytesIO(blob))
    text = _doc_text(doc)
    assert "SAT-FLAT" in text or "flatline" in text.lower()


def test_rcx_docx_fdd_section_optional(fault_rows, report_context):
    pytest.importorskip("docx")
    from docx import Document
    from open_fdd.reports.rcx_docx import build_rcx_docx

    without = build_rcx_docx(
        site_id="demo",
        site_name="Demo",
        window={"start": "a", "end": "b"},
        fault_rows=fault_rows,
        sections=["executive_summary"],
        report_context=report_context,
    )
    text_without = "\n".join(p.text for p in Document(io.BytesIO(without)).paragraphs)
    assert "FDD rule trend" not in text_without

    with_rules = build_rcx_docx(
        site_id="demo",
        site_name="Demo",
        window={"start": "a", "end": "b"},
        fault_rows=fault_rows,
        sections=["fdd_rule_trends"],
        report_context=report_context,
    )
    text_with = "\n".join(p.text for p in Document(io.BytesIO(with_rules)).paragraphs)
    assert "Zone temperature out of bounds" in text_with
    assert "STAT ZN-T" in text_with
