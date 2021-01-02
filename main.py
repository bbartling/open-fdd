
import time, glob, os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import seaborn as sns
import ruptures as rpt
import calendar
from statsmodels.tsa.seasonal import seasonal_decompose
from statsmodels.graphics.tsaplots import plot_acf,plot_pacf
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxcompose.composer import Composer
from docx import Document as Document_compose
import os
import fnmatch, re
from datetime import datetime
import argparse
import glob

from electrical_defs import Electrical_Defs




# Create the parser
my_parser = argparse.ArgumentParser(description='Name of csv file without the .csv at the end in sample_data directory')
# Add the arguments
my_parser.add_argument('filename',
                       type=str,
                       help='Name of csv file without the .csv at the end in sample_data directory')
# Execute parse_args()
args = my_parser.parse_args()

file_name = args.filename
print('starting to process ' + file_name)


try:

    df = pd.read_csv('./sample_data/' + str(file_name) + '.csv', index_col='Date', parse_dates=True)

except Exception as error:
     print(error)

#delete all files from previous report
sumerFiles = glob.glob('./static_summer/*')
for s in sumerFiles:
    os.remove(s)

mainFiles = glob.glob('./static_main/*')
for m in mainFiles:
    os.remove(m)

winterFiles = glob.glob('./static_winter/*')
for w in winterFiles:
    os.remove(w)


print('old files removed')


try:


    if 'kW' in df.columns:

        #remove rows of data is kW == 0
        #df = df[(df[['kW']] != 0).all(axis=1)]
        #print('Removed columns that recorded zero kW')

        #clean dataset
        df = Electrical_Defs.clean_dataset(df)
        print(df.describe())
        df2 = pd.DataFrame()

        #used on other scripts for summer & winer.py
        df.to_csv('./static_summer/workingDataSet.csv')
        df.to_csv('./static_winter/workingDataSet.csv')

        sum = df.loc[df.index.month.isin([6,7,8])]
        win = df.loc[df.index.month.isin([1,2,3])]

        jan = df.loc[df.index.month.isin([1])]
        feb = df.loc[df.index.month.isin([2])]
        mar = df.loc[df.index.month.isin([3])]

        jun = df.loc[df.index.month.isin([6])]
        jul = df.loc[df.index.month.isin([7])]
        aug = df.loc[df.index.month.isin([8])]


        #Number of days for plots
        FebDays = 28
        AprJunSepNovDays = 30
        JanMarMayJulAugOctDecDays = 31

        maxDate = df.loc[df['kW'].idxmax()]
        maxy = df.kW.max()
        mediany= df.kW.median()
        stdy= df.kW.std()

        winMaxDate = win.loc[win['kW'].idxmax()]
        winMaxy = win.kW.max()
        winMediany= win.kW.median()
        winStdy= win.kW.std()

        sumMaxDate = sum.loc[sum['kW'].idxmax()]
        sumMaxy = sum.kW.max()
        sumMediany= sum.kW.median()
        sumStdy= sum.kW.std()

        maxDay = maxDate.name.day
        maxMon = maxDate.name.month
        maxHour = maxDate.name.hour
        maxMon = df.loc[df.index.month.isin([maxMon])]

        sumMaxDay = sumMaxDate.name.day
        sumMaxMon = sumMaxDate.name.month
        sumMaxHour = sumMaxDate.name.hour
        sumMaxMon = sum.loc[sum.index.month.isin([sumMaxMon])]

        winMaxDay = winMaxDate.name.day
        winMaxMon = winMaxDate.name.month
        winMaxHour = winMaxDate.name.hour
        winMaxMon = win.loc[win.index.month.isin([winMaxMon])]

        print(f"Max Demand Found On {maxDate}")


        df.plot(figsize=(20, 10), title='Entire Dataset Plot')
        plt.ylim(5, maxy)
        plt.savefig('./static_main/Entire_Dataset_Plot.png')


        fig2 = sns.boxplot(data=df, x=df.index.month, y=df.kW)
        fig2.set_ylabel('Units')
        fig2.set_xlabel('Month #')
        fig2.set_title('kW by Month Number')
        fig2.figure.savefig('./static_main/AllDatakWboxPlots.png')

        #rolling averages trending plots
        kW_28d = df.kW.rolling(28, center=True).mean()
        kW_365d = df.kW.rolling(window=365, center=True, min_periods=360).mean()


        # Plot daily, 7-day rolling mean, and 90-day rolling mean time series
        fig3, ax3 = plt.subplots(figsize=(20, 10))
        ax3.plot(df.kW, marker='.', markersize=2, color='0.6',
        linestyle='None', label='Daily')
        ax3.plot(kW_28d, linewidth=2, label='Trend (28-d Rolling Mean)')
        ax3.plot(kW_365d, color='0.2', linewidth=3,
        label='Trend (365-d Rolling Mean)')
        # Set x-ticks to yearly interval and add legend and labels
        ax3.xaxis.set_major_locator(mdates.YearLocator())
        ax3.legend()
        ax3.set_xlabel('Year')
        ax3.set_ylabel('Consumption (GWh)')
        ax3.set_title('Trends in Power Demand kW')
        fig3.savefig('./static_main/kWtrendsPlot.png')


        maxMon.plot(figsize=(20, 10), title='Month Plot Captured With Day of Maximum Recorded Demand')
        plt.ylim(5, maxy)
        plt.savefig('./static_main/Month_maximum_recorded_demand.png')


        print("Working on Winter Summer change point algorithm plots....")

        #January Plot
        jan.plot(figsize=(20, 10), title = 'Demand for Month of January')
        plt.ylim(5, winMaxy)
        plt.savefig('./static_winter/Demands_for_Winter_Month_January.png')
        Electrical_Defs.changPoint2(jan, np.random.randint(low=1, high=JanMarMayJulAugOctDecDays, size=1), winMaxy, winStdy)
        Electrical_Defs.changPoint2(jan, np.random.randint(low=1, high=JanMarMayJulAugOctDecDays, size=1), winMaxy, winStdy)
        Electrical_Defs.changPoint2(jan, np.random.randint(low=1, high=JanMarMayJulAugOctDecDays, size=1), winMaxy, winStdy)
        Electrical_Defs.changPoint2(jan, np.random.randint(low=1, high=JanMarMayJulAugOctDecDays, size=1), winMaxy, winStdy)

        #February Plot
        feb.plot(figsize=(20, 10), title = 'Demand for Month of February')
        plt.ylim(5, winMaxy)
        plt.savefig('./static_winter/Demands_for_Winter_Month_February.png')
        Electrical_Defs.changPoint2(feb, np.random.randint(low=1, high=FebDays, size=1), winMaxy, winStdy)
        Electrical_Defs.changPoint2(feb, np.random.randint(low=1, high=FebDays, size=1), winMaxy, winStdy)
        Electrical_Defs.changPoint2(feb, np.random.randint(low=1, high=FebDays, size=1), winMaxy, winStdy)
        Electrical_Defs.changPoint2(feb, np.random.randint(low=1, high=FebDays, size=1), winMaxy, winStdy)

        #March
        mar.plot(figsize=(20, 10), title = 'Demand for Month of March')
        plt.ylim(5, winMaxy)
        plt.savefig('./static_winter/Demands_for_Winter_Month_March.png')
        Electrical_Defs.changPoint2(mar, np.random.randint(low=1, high=JanMarMayJulAugOctDecDays, size=1), winMaxy, winStdy)
        Electrical_Defs.changPoint2(mar, np.random.randint(low=1, high=JanMarMayJulAugOctDecDays, size=1), winMaxy, winStdy)
        Electrical_Defs.changPoint2(mar, np.random.randint(low=1, high=JanMarMayJulAugOctDecDays, size=1), winMaxy, winStdy)
        Electrical_Defs.changPoint2(mar, np.random.randint(low=1, high=JanMarMayJulAugOctDecDays, size=1), winMaxy, winStdy)

        #June
        jun.plot(figsize=(20, 10), title = 'Demand for Month of June')
        plt.ylim(5, sumMaxy)
        plt.savefig('./static_summer/Demands_for_Summer_Month_June.png')
        Electrical_Defs.changPoint2(jun, np.random.randint(low=1, high=AprJunSepNovDays, size=1), sumMaxy, sumStdy)
        Electrical_Defs.changPoint2(jun, np.random.randint(low=1, high=AprJunSepNovDays, size=1), sumMaxy, sumStdy)
        Electrical_Defs.changPoint2(jun, np.random.randint(low=1, high=AprJunSepNovDays, size=1), sumMaxy, sumStdy)
        Electrical_Defs.changPoint2(jun, np.random.randint(low=1, high=AprJunSepNovDays, size=1), sumMaxy, sumStdy)

        #July
        jul.plot(figsize=(20, 10), title = 'Demand for Month of July')
        plt.ylim(5, sumMaxy)
        plt.savefig('./static_summer/Demands_for_Summer_Month_July.png')
        Electrical_Defs.changPoint2(jul, np.random.randint(low=1, high=JanMarMayJulAugOctDecDays, size=1), sumMaxy, sumStdy)
        Electrical_Defs.changPoint2(jul, np.random.randint(low=1, high=JanMarMayJulAugOctDecDays, size=1), sumMaxy, sumStdy)
        Electrical_Defs.changPoint2(jul, np.random.randint(low=1, high=JanMarMayJulAugOctDecDays, size=1), sumMaxy, sumStdy)
        Electrical_Defs.changPoint2(jul, np.random.randint(low=1, high=JanMarMayJulAugOctDecDays, size=1), sumMaxy, sumStdy)

        aug.plot(figsize=(20, 10), title = 'Demand for Month of August')
        plt.ylim(5, sumMaxy)
        plt.savefig('./static_summer/Demands_for_Summer_Month_August.png')
        Electrical_Defs.changPoint2(aug, np.random.randint(low=1, high=JanMarMayJulAugOctDecDays, size=1), sumMaxy, sumStdy)
        Electrical_Defs.changPoint2(aug, np.random.randint(low=1, high=JanMarMayJulAugOctDecDays, size=1), sumMaxy, sumStdy)
        Electrical_Defs.changPoint2(aug, np.random.randint(low=1, high=JanMarMayJulAugOctDecDays, size=1), sumMaxy, sumStdy)
        Electrical_Defs.changPoint2(aug, np.random.randint(low=1, high=JanMarMayJulAugOctDecDays, size=1), sumMaxy, sumStdy)


        '''
        result = seasonal_decompose(df, model='additive')
        result.savefig('./static_main/seasonal_decompose.png', dpi=150)

        title = 'Autocorrelation: kW data'
        lags = 40
        plot_acf(hourly_avg['kW'],title=title,lags=lags)
        plot_acf.savefig('./static_main/Autocorrelation.png', dpi=150)

        title2 = 'Partial Autocorrelation: kW data'
        lags2 = 20
        plot_pacf(hourly_avg['kW'],title=title2,lags=lags2)
        plot_pacf.savefig('./static_main/Partial_Autocorrelation.png', dpi=150)
        '''

except Exception as error:
     print(error)




try:
    print("Starting Main docx report")

    #start Docx report
    document = Document()
    document.add_heading('Electricity Dataset Visualation', 0)

    #p = document.add_paragraph('This document is a visual aid for electrical load profiles').italic = True
    s = document.add_heading(str(file_name) + '.csv', level=1)
    #s.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #add plots of entire dataset
    for file in os.listdir('./static_main/'):
        if fnmatch.fnmatch(file, 'Entire_Dataset_Plot.png'):
            print(file)
            z = f'{file}'
            s = document.add_paragraph(z, style='Intense Quote')
            document.add_picture('./static_main/' + z, width=Inches(6))

    s = document.add_heading('Plot of month with max demand recorded', level=1)
    for file in os.listdir('./static_main/'):
        if fnmatch.fnmatch(file, 'Month_maximum_recorded_demand.png'):
            print(file)
            z = f'{file}'
            s = document.add_paragraph(z, style='Intense Quote')
            document.add_picture('./static_main/' + z, width=Inches(6))


    s = document.add_heading('Power box plots per month', level=1)
    for file in os.listdir('./static_main/'):
        if fnmatch.fnmatch(file, 'AllDatakWboxPlots.png'):
            print(file)
            z = f'{file}'
            s = document.add_paragraph(z, style='Intense Quote')
            document.add_picture('./static_main/' + z, width=Inches(6))


    s = document.add_heading('Power consumption trends', level=1)
    for file in os.listdir('./static_main/'):
        if fnmatch.fnmatch(file, 'kWtrendsPlot.png'):
            print(file)
            z = f'{file}'
            s = document.add_paragraph(z, style='Intense Quote')
            document.add_picture('./static_main/' + z, width=Inches(6))

    document.add_page_break()
    document.save('./final_report/' + str(file_name) + '.docx')

except Exception as error:
     print(error)
     document.add_heading('Error on Main Report!', level=1)
     document.add_paragraph(str(error))
     document.save('./final_report/' + str(file_name) + '.docx')


print('All Done, trying Winter and Summer data/reports')


try:

    #read CSV file
    wint = pd.read_csv('./static_winter/workingDataSet.csv', index_col='Date', parse_dates=True)
    wint = wint.loc[wint.index.month.isin([1,2,3])]
    wint_copy2 = wint.copy()


    maxDate = wint.loc[wint['kW'].idxmax()]
    maxy = wint.kW.max()
    mediany= wint.kW.median()
    stdy= wint.kW.std()

    maxDay = maxDate.name.day
    maxMon = maxDate.name.month
    maxHour = maxDate.name.hour

    wintDesc = wint.describe()

    print(f"Max Demand Found On {maxDate}")
    print(wintDesc)


    wint.plot(figsize=(20, 10))
    plt.savefig('./static_winter/datasetPlot.png')



    #resample data for hourly avgerage
    hourly_avg = pd.DataFrame()
    hourly_avg['kW'] = wint['kW'].resample('H').mean()

    #calc units of energy conert kw to kwh
    daily_summary = pd.DataFrame()
    idx = wint.groupby(wint.index.normalize())['kW'].idxmax()
    daily_summary = wint.loc[idx].copy()

    # get the hour
    daily_summary['hour_max_demand'] = daily_summary.index.hour

    # set date as index
    daily_summary.index = daily_summary.index.normalize()


    daily_summary['kWH'] = hourly_avg['kW'].resample('D').sum()
    energySum = daily_summary['kWH'].sum()
    kWhInfo = f"Total Sum of calculated electrical energy {energySum} kWh"

    print(kWhInfo)

    # First figure.
    fig1, ax1 = plt.subplots(figsize=(25, 10))
    ax2 = ax1.twinx()

    ax1.bar(daily_summary.index, daily_summary['hour_max_demand'], width=20, alpha=0.2, color='orange')
    ax1.grid(b=False) # turn off grid #2

    ax2.plot(daily_summary.kW)
    ax2.set_title('Max Demand per Day and Max Demand Hour of Day')
    ax2.set_ylabel('Electric Demand kW')
    ax1.set_ylabel('Hour of Day')

    fig1.savefig('./static_winter/Max_Demand_and_Max_Hour_of_Day.png')

    # Figure 2.
    fig2, ax3 = plt.subplots(figsize=(25, 10))
    ax3.set_title(' 7 Day Rolling Average - kWh Per Day')
    data3 = daily_summary.kWH.rolling(7, center=True).mean()
    ax3.plot(data3)
    fig2.savefig('./static_winter/kWhRollingAvg.png')



    #create dummy variables
    daily_summary['month'] = daily_summary.index.month
    daily_summary['day_of_week'] = daily_summary.index.dayofweek

    daily_summary.columns

    first = daily_summary.first('1D').index.date[0]
    last = daily_summary.last('1D').index.date[0]

    print(first,last)

    dayInfo = f"Resampling the interval dataset to calculate units of energy KWh/day, the first day is {first} and the last day is {last}"
    print(dayInfo)


    totalDays = last - first
    totalDays = totalDays.days
    totalDays

    totalDayInfo = f"Total days in dataset {totalDays} days"
    print(totalDayInfo)


    #daily min & max demand kW
    #save these values for later, join on final wint
    daily_summary_Q = (wint['kW'].resample('D')
                           .agg(lambda x: x.quantile([.025, 0.25, 0.75, 0.975]))
                           .unstack()
                           .add_prefix('kW_Q'))

    daily_summary_Q.columns
    daily_summary_Q_plots = daily_summary_Q[['kW_Q0.025','kW_Q0.975']]
    daily_summary_Q_plots.plot(figsize=(25,10), title='High & Low Load Values kW per Day')
    plt.savefig('./static_winter/highLowLoadsPlot.png')


    daily_summary = daily_summary_Q.join(daily_summary)
    daily_summary.columns




    print("Thinking for a little while on Winter dataset....")
    print("Starting to comb thru data day by day with the change point algorithm...")



    s = (wint['kW'].ge(mediany)            # compare to mediany
            .groupby(wint.index.normalize())  # groupby day
            .transform('any')               # any time with value larger than median
        )

    #filter out days of data where day kW is less than entire kW mean value
    wint = wint[s]


    data = []


    for idx, days in wint.groupby(wint.index.date):
        stuff = {}
        listAll = Electrical_Defs.changPointDf(days)
        stuff['kW1diff'] = listAll[0]
        stuff['kW1Hrs'] = listAll[1]
        stuff['kW2diff'] = listAll[2]
        stuff['kW2Hrs'] = listAll[3]
        stuff['kW3diff'] = listAll[4]
        stuff['kW3Hrs'] = listAll[5]
        stuff['kW4diff'] = listAll[6]
        stuff['kW4Hrs'] = listAll[7]
        stuff['kW5diff'] = listAll[8]
        stuff['kW5Hrs'] = listAll[9]
        stuff['kW6diff'] = listAll[10]
        stuff['kW6Hrs'] = listAll[11]
        stuff['kW7diff'] = listAll[12]
        stuff['kW7Hrs'] = listAll[13]
        stuff['kW8diff'] = listAll[14]
        stuff['kW8Hrs'] = listAll[15]
        stuff['kW9diff'] = listAll[16]
        stuff['kW9Hrs'] = listAll[17]
        stuff['kW10diff'] = listAll[18]
        stuff['kW10Hrs'] = listAll[19]
        stuff['kW11diff'] = listAll[20]
        stuff['kW11Hrs'] = listAll[21]
        stuff['kW12diff'] = listAll[22]
        stuff['kW12Hrs'] = listAll[23]
        stuff['kW13diff'] = listAll[24]
        stuff['kW13Hrs'] = listAll[25]
        stuff['kW14diff'] = listAll[26]
        stuff['kW14Hrs'] = listAll[27]
        stuff['kW15diff'] = listAll[28]
        stuff['kW15Hrs'] = listAll[29]
        stuff['kW16diff'] = listAll[30]
        stuff['kW16Hrs'] = listAll[31]
        data.append(stuff)


    master_hrs = pd.DataFrame(data)


    master_hrs.columns = [
            'kW1diff', 'kW1Hrs',
            'kW2diff', 'kW2Hrs',
            'kW3diff', 'kW3Hrs',
            'kW4diff', 'kW4Hrs',
            'kW5diff', 'kW5Hrs',
            'kW6diff', 'kW6Hrs',
            'kW7diff', 'kW7Hrs',
            'kW8diff', 'kW8Hrs',
            'kW9diff', 'kW9Hrs',
            'kW10diff', 'kW10Hrs',
            'kW11diff', 'kW11Hrs',
            'kW12diff', 'kW12Hrs',
            'kW13diff', 'kW13Hrs',
            'kW14diff', 'kW14Hrs',
            'kW15diff', 'kW15Hrs',
            'kW16diff', 'kW16Hrs']

    master_hrsIdx = pd.period_range(first, periods=len(master_hrs))

    len(master_hrsIdx)

    master_hrs['Date'] = master_hrsIdx
    master_hrs = master_hrs.set_index('Date')


    master_hrs.head()
    master_hrs.filter(regex='diff$',axis=1).plot(figsize=(25, 25))


    sort_by_kW = master_hrs.filter(regex='diff$',axis=1)
    sort_by_kW = sort_by_kW.sort_values((
    [
            'kW1diff',
            'kW2diff',
            'kW3diff',
            'kW4diff',
            'kW5diff',
            'kW6diff',
            'kW7diff',
            'kW8diff',
            'kW9diff',
            'kW10diff',
            'kW11diff',
            'kW12diff',
            'kW13diff',
            'kW14diff',
            'kW15diff',
            'kW16diff']
    ))


    #sort_by_kW15 = sort_by_kW.index[0:15].to_list()
    sort_by_kW15 = sort_by_kW.index[0:15]
    sort_by_kW15 = pd.DataFrame(sort_by_kW15)

    rank = 1

    for index, row in sort_by_kW15.iterrows():


        x = row[0]
        x = wint_copy2.loc[str(x)]
        title= f'#{rank} Highest Ranked Electrical Demand Day by Changepoint Algorithm Detection'
        rank += 1
        x.plot(figsize=(25, 10),title=title)
        plt.savefig('./static_winter/' + title + '.png')


    #master_hrs.filter(regex='Hrs$',axis=1).plot(subplots=True, figsize=(25, 25))

    daily_summary.describe()


    master_hrs.to_csv('./static_winter/master_hrs.csv')
    daily_summary.to_csv('./static_winter/daily_summary.csv')



except Exception as error:
     print(error)



print("Starting Winter docx report")



try:


    document.add_heading('Data Analysis Report Winter', 0)
    p = document.add_paragraph('Winter Months Electrical Load Profiles').italic = True


    #add plots of entire dataset
    for file in os.listdir('./static_winter/'):
        if fnmatch.fnmatch(file, 'datasetPlot*.png'):
            print(file)
            z = f'{file}'
            s = document.add_paragraph(z, style='Intense Quote')
            document.add_picture('./static_winter/' + z, width=Inches(6))

    document.save('./final_report/' + str(file_name) + '.docx')


    #add in max demand to report #add in sum statistics
    document.add_heading('Max Demand Found In Dataset', level=1)
    document.add_paragraph(str(maxDate))


    document.add_heading('Dataset Summary Statistics', level=1)
    document.add_paragraph(str(wintDesc))

    document.save('./final_report/' + str(file_name) + '.docx')



    #Highest Ranked Spiked Demand Day
    document.add_heading('Highest Ranked Change Point Algorithm Detection', level=1)

    for file in os.listdir('./static_winter/'):
        if fnmatch.fnmatch(file, '*Highest Ranked Electrical Demand Day by Changepoint Algorithm Detection.png'):
            print(file)
            z = f'{file}'
            document.add_picture('./static_winter/' + z, width=Inches(6))


    document.save('./final_report/' + str(file_name) + '.docx')
    document.add_page_break()

    s = document.add_heading('Daily High and Low Load kW Values', level=1)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for file in os.listdir('./static_winter/'):
        if fnmatch.fnmatch(file, 'highLowLoadsPlot.png'):
            print(file)
            z = f'{file}'
            s = document.add_paragraph(z, style='Intense Quote')
            document.add_picture('./static_winter/' + z, width=Inches(6))


    s = document.add_heading('Max Demand and Hour of Day Plot', level=1)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for file in os.listdir('./static_winter/'):
        if fnmatch.fnmatch(file, 'Max_Demand_and_Max_Hour_of_Day.png'):
            print(file)
            z = f'{file}'
            s = document.add_paragraph(z, style='Intense Quote')
            document.add_picture('./static_winter/' + z, width=Inches(6))

    document.save('./final_report/' + str(file_name) + '.docx')
    document.add_page_break()


    document.add_paragraph(dayInfo, style='List Bullet')
    document.add_paragraph(totalDayInfo, style='List Bullet')
    document.add_paragraph(kWhInfo, style='List Bullet')


    s = document.add_heading('kWh Rolling 7 Day Avg', level=1)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for file in os.listdir('./static_winter/'):
        if fnmatch.fnmatch(file, 'kWhRollingAvg.png'):
            print(file)
            z = f'{file}'
            s = document.add_paragraph(z, style='Intense Quote')
            document.add_picture('./static_winter/' + z, width=Inches(6))

    document.save('./final_report/' + str(file_name) + '.docx')


    s = document.add_heading('Demand Plots By Month', level=1)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for file in os.listdir('./static_winter/'):
        if fnmatch.fnmatch(file, 'Demands_for_Winter_Month_*.png'):
            print(file)
            z = f'{file}'
            s = document.add_paragraph(z, style='Intense Quote')
            document.add_picture('./static_winter/' + z, width=Inches(6))

    document.save('./final_report/' + str(file_name) + '.docx')

    document.add_page_break()





except Exception as error:
     print(error)
     document.add_heading('Error on Winter Report!', level=1)
     document.add_paragraph(str(error))
     document.save('./final_report/' + str(file_name) + '.docx')





print("Starting Summer Analysis")



try:

    #read CSV file
    summer = pd.read_csv('./static_summer/workingDataSet.csv', index_col='Date', parse_dates=True)
    summer = summer.loc[summer.index.month.isin([6,7,8])]
    summer_copy2 = summer.copy()


    maxDate = summer.loc[summer['kW'].idxmax()]
    maxy = summer.kW.max()
    mediany= summer.kW.median()
    stdy= summer.kW.std()

    maxDay = maxDate.name.day
    maxMon = maxDate.name.month
    maxHour = maxDate.name.hour

    summerDesc = summer.describe()

    print(f"Max Demand Found On {maxDate}")
    print(summerDesc)


    summer.plot(figsize=(20, 10))
    plt.savefig('./static_summer/datasetPlot.png')


    #resample data for hourly avgerage
    hourly_avg = pd.DataFrame()
    hourly_avg['kW'] = summer['kW'].resample('H').mean()

    #calc units of energy conert kw to kwh
    daily_summary = pd.DataFrame()
    idx = summer.groupby(summer.index.normalize())['kW'].idxmax()
    daily_summary = summer.loc[idx].copy()

    # get the hour
    daily_summary['hour_max_demand'] = daily_summary.index.hour

    # set date as index
    daily_summary.index = daily_summary.index.normalize()


    daily_summary['kWH'] = hourly_avg['kW'].resample('D').sum()
    energySum = daily_summary['kWH'].sum()
    kWhInfo = f"Total Sum of calculated electrical energy {energySum} kWh"

    print(kWhInfo)

    # First figure.
    fig1, ax1 = plt.subplots(figsize=(25, 10))
    ax2 = ax1.twinx()

    ax1.bar(daily_summary.index, daily_summary['hour_max_demand'], width=20, alpha=0.2, color='orange')
    ax1.grid(b=False) # turn off grid #2

    ax2.plot(daily_summary.kW)
    ax2.set_title('Max Demand per Day and Max Demand Hour of Day')
    ax2.set_ylabel('Electric Demand kW')
    ax1.set_ylabel('Hour of Day')

    fig1.savefig('./static_summer/Max_Demand_and_Max_Hour_of_Day.png')

    # Figure 2.
    fig2, ax3 = plt.subplots(figsize=(25, 10))
    ax3.set_title(' 7 Day Rolling Average - kWh Per Day')
    data3 = daily_summary.kWH.rolling(7, center=True).mean()
    ax3.plot(data3)
    fig2.savefig('./static_summer/kWhRollingAvg.png')


    #create dummy variables
    daily_summary['month'] = daily_summary.index.month
    daily_summary['day_of_week'] = daily_summary.index.dayofweek

    daily_summary.columns

    first = daily_summary.first('1D').index.date[0]
    last = daily_summary.last('1D').index.date[0]

    print(first,last)

    dayInfo = f"Resampling the interval dataset to calculate units of energy KWh/day, the first day is {first} and the last day is {last}"
    print(dayInfo)


    totalDays = last - first
    totalDays = totalDays.days
    totalDays

    totalDayInfo = f"Total days in dataset {totalDays} days"
    print(totalDayInfo)


    #daily min & max demand kW
    #save these values for later, join on final summer
    daily_summary_Q = (summer['kW'].resample('D')
                           .agg(lambda x: x.quantile([.025, 0.25, 0.75, 0.975]))
                           .unstack()
                           .add_prefix('kW_Q'))

    daily_summary_Q.columns
    daily_summary_Q_plots = daily_summary_Q[['kW_Q0.025','kW_Q0.975']]
    daily_summary_Q_plots.plot(figsize=(25,10), title='High & Low Load Values kW per Day')
    plt.savefig('./static_summer/highLowLoadsPlot.png')


    daily_summary = daily_summary_Q.join(daily_summary)
    daily_summary.columns




    print("Thinking for a little while on Summer dataset....")
    print("Starting to comb thru data day by day with the change point algorithm...")



    s = (summer['kW'].ge(mediany)            # compare to mediany
            .groupby(summer.index.normalize())  # groupby day
            .transform('any')               # any time with value larger than median
        )

    #filter out days of data where day kW is less than entire kW mean value
    summer = summer[s]


    data = []


    for idx, days in summer.groupby(summer.index.date):
        stuff = {}
        listAll = Electrical_Defs.changPointDf(days)
        stuff['kW1diff'] = listAll[0]
        stuff['kW1Hrs'] = listAll[1]
        stuff['kW2diff'] = listAll[2]
        stuff['kW2Hrs'] = listAll[3]
        stuff['kW3diff'] = listAll[4]
        stuff['kW3Hrs'] = listAll[5]
        stuff['kW4diff'] = listAll[6]
        stuff['kW4Hrs'] = listAll[7]
        stuff['kW5diff'] = listAll[8]
        stuff['kW5Hrs'] = listAll[9]
        stuff['kW6diff'] = listAll[10]
        stuff['kW6Hrs'] = listAll[11]
        stuff['kW7diff'] = listAll[12]
        stuff['kW7Hrs'] = listAll[13]
        stuff['kW8diff'] = listAll[14]
        stuff['kW8Hrs'] = listAll[15]
        stuff['kW9diff'] = listAll[16]
        stuff['kW9Hrs'] = listAll[17]
        stuff['kW10diff'] = listAll[18]
        stuff['kW10Hrs'] = listAll[19]
        stuff['kW11diff'] = listAll[20]
        stuff['kW11Hrs'] = listAll[21]
        stuff['kW12diff'] = listAll[22]
        stuff['kW12Hrs'] = listAll[23]
        stuff['kW13diff'] = listAll[24]
        stuff['kW13Hrs'] = listAll[25]
        stuff['kW14diff'] = listAll[26]
        stuff['kW14Hrs'] = listAll[27]
        stuff['kW15diff'] = listAll[28]
        stuff['kW15Hrs'] = listAll[29]
        stuff['kW16diff'] = listAll[30]
        stuff['kW16Hrs'] = listAll[31]
        data.append(stuff)


    master_hrs = pd.DataFrame(data)


    master_hrs.columns = [
            'kW1diff', 'kW1Hrs',
            'kW2diff', 'kW2Hrs',
            'kW3diff', 'kW3Hrs',
            'kW4diff', 'kW4Hrs',
            'kW5diff', 'kW5Hrs',
            'kW6diff', 'kW6Hrs',
            'kW7diff', 'kW7Hrs',
            'kW8diff', 'kW8Hrs',
            'kW9diff', 'kW9Hrs',
            'kW10diff', 'kW10Hrs',
            'kW11diff', 'kW11Hrs',
            'kW12diff', 'kW12Hrs',
            'kW13diff', 'kW13Hrs',
            'kW14diff', 'kW14Hrs',
            'kW15diff', 'kW15Hrs',
            'kW16diff', 'kW16Hrs']

    master_hrsIdx = pd.period_range(first, periods=len(master_hrs))

    len(master_hrsIdx)

    master_hrs['Date'] = master_hrsIdx
    master_hrs = master_hrs.set_index('Date')


    master_hrs.head()
    master_hrs.filter(regex='diff$',axis=1).plot(figsize=(25, 25))


    sort_by_kW = master_hrs.filter(regex='diff$',axis=1)
    sort_by_kW = sort_by_kW.sort_values((
    [
            'kW1diff',
            'kW2diff',
            'kW3diff',
            'kW4diff',
            'kW5diff',
            'kW6diff',
            'kW7diff',
            'kW8diff',
            'kW9diff',
            'kW10diff',
            'kW11diff',
            'kW12diff',
            'kW13diff',
            'kW14diff',
            'kW15diff',
            'kW16diff']
    ))


    #sort_by_kW15 = sort_by_kW.index[0:15].to_list()
    sort_by_kW15 = sort_by_kW.index[0:15]
    sort_by_kW15 = pd.DataFrame(sort_by_kW15)

    rank = 1

    for index, row in sort_by_kW15.iterrows():


        x = row[0]
        x = summer_copy2.loc[str(x)]
        title= f'#{rank} Summer Highest Ranked Changepoint Algorithm Detection'
        rank += 1
        x.plot(figsize=(25, 10),title=title)
        plt.savefig('./static_summer/' + title + '.png')


    #master_hrs.filter(regex='Hrs$',axis=1).plot(subplots=True, figsize=(25, 25))

    daily_summary.describe()


    master_hrs.to_csv('./static_summer/master_hrs.csv')
    daily_summary.to_csv('./static_summer/daily_summary.csv')



except Exception as error:
     print(error)



print("Starting Summer docx report")



try:


    document.add_heading('Data Analysis Report Summer', 0)
    p = document.add_paragraph('Summer Months Electrical Load Profiles').italic = True


    #add plots of entire dataset
    for file in os.listdir('./static_summer/'):
        if fnmatch.fnmatch(file, 'datasetPlot*.png'):
            print(file)
            z = f'{file}'
            s = document.add_paragraph(z, style='Intense Quote')
            document.add_picture('./static_summer/' + z, width=Inches(6))

    document.save('./final_report/' + str(file_name) + '.docx')


    #add in max demand to report #add in sum statistics
    document.add_heading('Max Demand Found In Dataset', level=1)
    document.add_paragraph(str(maxDate))


    document.add_heading('Dataset Summary Statistics', level=1)
    document.add_paragraph(str(summerDesc))

    document.save('./final_report/' + str(file_name) + '.docx')


    #Highest Ranked Spiked Demand Day
    document.add_heading('Highest Ranked Change Point Algorithm Detection', level=1)

    for file in os.listdir('./static_summer/'):
        if fnmatch.fnmatch(file, '*Summer Highest Ranked Changepoint Algorithm Detection.png'):
            print(file)
            z = f'{file}'
            document.add_picture('./static_summer/' + z, width=Inches(6))


    document.save('./final_report/' + str(file_name) + '.docx')
    document.add_page_break()

    s = document.add_heading('Daily High and Low Load kW Values', level=1)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for file in os.listdir('./static_summer/'):
        if fnmatch.fnmatch(file, 'highLowLoadsPlot.png'):
            print(file)
            z = f'{file}'
            s = document.add_paragraph(z, style='Intense Quote')
            document.add_picture('./static_summer/' + z, width=Inches(6))


    s = document.add_heading('Max Demand and Hour of Day Plot', level=1)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for file in os.listdir('./static_summer/'):
        if fnmatch.fnmatch(file, 'Max_Demand_and_Max_Hour_of_Day.png'):
            print(file)
            z = f'{file}'
            s = document.add_paragraph(z, style='Intense Quote')
            document.add_picture('./static_summer/' + z, width=Inches(6))

    document.save('./final_report/' + str(file_name) + '.docx')
    document.add_page_break()


    document.add_paragraph(dayInfo, style='List Bullet')
    document.add_paragraph(totalDayInfo, style='List Bullet')
    document.add_paragraph(kWhInfo, style='List Bullet')


    s = document.add_heading('kWh Rolling 7 Day Avg', level=1)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for file in os.listdir('./static_summer/'):
        if fnmatch.fnmatch(file, 'kWhRollingAvg.png'):
            print(file)
            z = f'{file}'
            s = document.add_paragraph(z, style='Intense Quote')
            document.add_picture('./static_summer/' + z, width=Inches(6))

    document.save('./final_report/' + str(file_name) + '.docx')


    s = document.add_heading('Demand Plots By Month', level=1)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for file in os.listdir('./static_summer/'):
        if fnmatch.fnmatch(file, 'Demands_for_Summer_Month_*.png'):
            print(file)
            z = f'{file}'
            s = document.add_paragraph(z, style='Intense Quote')
            document.add_picture('./static_summer/' + z, width=Inches(6))

    document.save('./final_report/' + str(file_name) + '.docx')


    now = datetime.now()
    print('data_analysis report done!')
    dt_string = now.strftime("%m/%d/%Y %H:%M:%S")
    print("date and time =", dt_string)
    document.add_paragraph('report compiled on:')
    document.add_paragraph(str(dt_string))

    document.save('./final_report/' + str(file_name) + '.docx')



    print('All Done final ' + str(file_name) + ' docx saved!')

except Exception as error:
     print(error)
     document.add_heading('Error on Summer Report!', level=1)
     document.add_paragraph(str(error))
     document.save('./final_report/' + str(file_name) + '.docx')



'''
CONCAT SUMMER AND WINTER WORD DOCs


files = ["./static_main/data_analysis_report_main.docx", "./static_summer/data_analysis_report_summer.docx",
         "./static_winter/data_analysis_report_wint.docx"]

composed = "./static_main/final.docx"

result = Document(files[0])
result.add_page_break()
composer = Composer(result)

for i in range(1, len(files)):
    doc = Document(files[i])

    if i != len(files) - 1:
        doc.add_page_break()

    composer.append(doc)

composer.save(composed)
print('All Done final docx saved!')
'''
