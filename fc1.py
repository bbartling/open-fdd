import operator, time
from datetime import datetime, timedelta

import os

import os

import pandas as pd
import numpy as np

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from docx import Document
from docx.shared import Inches

import argparse, math

from faults import FaultConditionOne

# python 3.10 on Windows 10
# py .\fc1.py -i ./ahu_data/hvac_random_fake_data/fc1_fake_data1.csv -o fake1_ahu_fc1_report

parser = argparse.ArgumentParser(add_help=False)
args = parser.add_argument_group("Options")

args.add_argument(
    "-h", "--help", action="help", help="Show this help message and exit."
)
args.add_argument("-i", "--input", required=True, type=str, help="CSV File Input")
args.add_argument(
    "-o", "--output", required=True, type=str, help="Word File Output Name"
)
"""
FUTURE 
 * incorporate an arg for SI units 
 * °C on temp sensors
 * piping pressure sensor PSI conversion
 * air flow CFM conversion
 * AHU duct static pressure "WC

args.add_argument('--use-SI-units', default=False, action='store_true')
args.add_argument('--no-SI-units', dest='use-SI-units', action='store_false')
"""
args = parser.parse_args()

# required params taken from the screenshot above
VFD_SPEED_PERCENT_ERR_THRES = 0.05
VFD_SPEED_PERCENT_MAX = 0.99
DUCT_STATIC_INCHES_ERR_THRES = 0.1

_fc1 = FaultConditionOne(
    VFD_SPEED_PERCENT_ERR_THRES,
    VFD_SPEED_PERCENT_MAX,
    DUCT_STATIC_INCHES_ERR_THRES,
    "duct_static",
    "supply_vfd_speed",
    "duct_static_setpoint",
)


# def fault_condition_one(df):
#     return operator.and_(
#         df.duct_static < (df.duct_static_setpoint - df.duct_static_inches_err_thres),
#         df.supply_vfd_speed
#         > (df.vfd_speed_percent_max - df.vfd_speed_percent_err_thres),
#     )


df = pd.read_csv(args.input, index_col="Date", parse_dates=True).rolling("5T").mean()

df["duct_static_setpoint"] = 1

start = df.head(1).index.date
print("Dataset start: ", start)

end = df.tail(1).index.date
print("Dataset end: ", end)


# make an entire column out of these params in the Pandas Dataframe
# df['vfd_speed_percent_err_thres'] = VFD_SPEED_PERCENT_ERR_THRES
# df['duct_static_inches_err_thres'] = DUCT_STATIC_INCHES_ERR_THRES
# df['vfd_speed_percent_max'] = VFD_SPEED_PERCENT_MAX

for col in df.columns:
    print("df column: ", col, "max len: ", df[col].size)

# df['fc1_flag'] = fault_condition_one(df)
df["fc1_flag"] = _fc1.apply(df)

df2 = df.copy().dropna()
df2["fc1_flag"] = df2["fc1_flag"].astype(int)

# drop params column for better plot
# df2 = df2.drop(
#     [
#         "vfd_speed_percent_err_thres",
#         "duct_static_inches_err_thres",
#         "vfd_speed_percent_max",
#     ],
#     axis=1,
# )

print(df2.columns)
print(df2.fc1_flag)

fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(25, 8))
plt.title("Fault Conditions 1 Plots")

ax1b = ax1.twinx()
(plot1a,) = ax1.plot(df2.index, df2.duct_static, color="r")  # red
(plot1b,) = ax1b.plot(df2.index, df2.supply_vfd_speed, color="b")  # blue
ax1.set_ylabel("Duct Static Pressure and Setpoint Subplots")

ax1.set_ylabel("Duct Static")
ax1b.set_ylabel("Fan Speed")

ax2.plot(df2.index, df2.fc1_flag, color="g")  # green
ax2.set_xlabel("Date")
ax2.set_ylabel("Fault Flag")

red_patch = mpatches.Patch(color="red", label="Duct Static")
blue_patch = mpatches.Patch(color="blue", label="Supply Fan Speed")
green_patch = mpatches.Patch(color="green", label="FC1 Flag")
plt.legend(handles=[red_patch, blue_patch, green_patch])
plt.tight_layout()
plt.savefig(os.path.join(os.path.curdir, "static", "ahu_fc1_fans_plot.png"))

print(f"Starting {args.output} docx report!")
document = Document()
document.add_heading("Fault Condition One Report", 0)

p = document.add_paragraph(
    "Fault condition one of ASHRAE Guideline 36 is related to flagging poor performance of a AHU variable supply fan attempting to control to a duct pressure setpoint. Fault condition equation as defined by ASHRAE:"
)

document.add_picture(
    os.path.join(os.path.curdir, "images", "fc1_definition.png"), width=Inches(6)
)
document.add_heading("Dataset Plot", level=2)

# ADD IN SUBPLOTS SECTION
document.add_picture(
    os.path.join(os.path.curdir, "static", "ahu_fc1_fans_plot.png"), width=Inches(6)
)
document.add_heading("Dataset Statistics", level=2)

# calculate dataset statistics
delta = df2.index.to_series().diff()
total_days = round(delta.sum() / pd.Timedelta(days=1), 2)
print("DAYS ALL DATA: ", total_days)
total_hours = delta.sum() / pd.Timedelta(hours=1)
print("TOTAL HOURS: ", total_hours)
hours_fc1_mode = (delta * df2["fc1_flag"]).sum() / pd.Timedelta(hours=1)
print("FALT FLAG TRUE TOTAL HOURS: ", hours_fc1_mode)
percent_true = round(df2.fc1_flag.mean() * 100, 2)
print("PERCENT TIME WHEN FLAG IS TRUE: ", percent_true, "%")
percent_false = round((100 - percent_true), 2)
print("PERCENT TIME WHEN FLAG 5 FALSE: ", percent_false, "%")
df2["hour_of_the_day_fc1"] = df2.index.hour.where(df2["fc1_flag"] == 1)
flag_true_duct_static = round(df2.duct_static.where(df2["fc1_flag"] == 1).mean(), 2)


# make hist plots fc3
fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
ax.hist(df2.hour_of_the_day_fc1.dropna())
ax.set_xlabel("24 Hour Number in Day")
ax.set_ylabel("Frequency")
ax.set_title(f"Hour-Of-Day When Fault Flag 1 is TRUE")
fig.savefig(os.path.join(os.path.curdir, "static", "ahu_fc1_histogram.png"))


# add calcs to word doc
paragraph = document.add_paragraph()
paragraph.style = "List Bullet"
paragraph.add_run(f"Total time in days calculated in dataset: {total_days}")

paragraph = document.add_paragraph()
paragraph.style = "List Bullet"
paragraph.add_run(f"Total time in hours calculated in dataset: {total_hours}")

paragraph = document.add_paragraph()
paragraph.style = "List Bullet"
paragraph.add_run(f"Total time in hours for when fault flag is True: {hours_fc1_mode}")

paragraph = document.add_paragraph()
paragraph.style = "List Bullet"
paragraph.add_run(
    f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
)

paragraph = document.add_paragraph()
paragraph.style = "List Bullet"
paragraph.add_run(
    f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
)

paragraph = document.add_paragraph()
# ADD HIST Plots
document.add_heading("Time-of-day Histogram Plots", level=2)
document.add_picture(
    os.path.join(os.path.curdir, "static", "ahu_fc1_histogram.png"), width=Inches(6)
)

if not math.isnan(flag_true_duct_static):
    paragraph = document.add_paragraph()
    paragraph.style = "List Bullet"
    paragraph.add_run(
        f'Average duct system pressure for when in fault condition (fan VFD speed > 95%): {flag_true_duct_static}"WC'
    )


paragraph = document.add_paragraph()

# ADD in Summary Statistics of fan operation
document.add_heading("VFD Speed Statistics", level=2)
paragraph = document.add_paragraph()
paragraph.style = "List Bullet"
paragraph.add_run(str(df2.supply_vfd_speed.describe()))

# ADD in Summary Statistics of duct pressure
document.add_heading("Duct Pressure Statistics", level=2)
paragraph = document.add_paragraph()
paragraph.style = "List Bullet"
paragraph.add_run(str(df2.duct_static.describe()))

# ADD in Summary Statistics of duct pressure
document.add_heading("Duct Pressure Setpoints Statistics", level=2)
paragraph = document.add_paragraph()
paragraph.style = "List Bullet"
paragraph.add_run(str(df2.duct_static_setpoint.describe()))


document.add_heading("Suggestions based on data analysis", level=2)
paragraph = document.add_paragraph()
paragraph.style = "List Bullet"

if percent_true > 5.0:
    paragraph.add_run(
        "The percent True metric that represents the amount of time for when the fault flag is True is high indicating the fan is running at high speeds and appearing to not generate good duct static pressure"
    )

else:
    paragraph.add_run(
        "The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the fan appears to generate good duct static pressure"
    )


paragraph = document.add_paragraph()
paragraph.style = "List Bullet"

print("df2.duct_static_setpoint.std: ", df2.duct_static_setpoint.std())
if df2.duct_static_setpoint.std() == 0:
    paragraph.add_run("No duct pressure setpoint reset detected (BAD)")

else:
    paragraph.add_run("Duct pressure reset detected (Good)")

paragraph = document.add_paragraph()
run = paragraph.add_run(f"Report generated: {time.ctime()}")
run.style = "Emphasis"

document.save(os.path.join(os.path.curdir, "final_report", f"{args.output}.docx"))
print("All Done")
