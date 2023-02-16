import math
import os
import time
from io import BytesIO

import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches

from faults import FaultConditionOne


class FaultCodeOneReport:
    """Class provides the definitions for Fault Code 1 Report."""

    def __init__(
        self,
        vfd_speed_percent_err_thres: float,
        vfd_speed_percent_max: float,
        duct_static_inches_err_thres: float,
        duct_static_col: str,
        supply_vfd_speed_col: str,
        duct_static_setpoint_col: str,
    ):
        self.vfd_speed_percent_err_thres = vfd_speed_percent_err_thres
        self.vfd_speed_percent_max = vfd_speed_percent_max
        self.duct_static_inches_err_thres = duct_static_inches_err_thres
        self.duct_static_col = duct_static_col
        self.supply_vfd_speed_col = supply_vfd_speed_col
        self.duct_static_setpoint_col = duct_static_setpoint_col

    def create_fan_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:
        if output_col is None:
            output_col = "fc1_flag"

        df[output_col] = df[output_col].astype(int)

        fig, (ax1, ax2, ax3) = plt.subplots(3, 1, figsize=(25, 8))
        plt.title('Fault Conditions 1 Plot')

        ax1.plot(df.index, df[self.duct_static_col], label="STATIC")
        ax1.legend(loc='best')
        ax1.set_ylabel("Inch WC")

        ax2.plot(df.index, df[self.supply_vfd_speed_col],
                 color="g", label="FAN")
        ax2.legend(loc='best')
        ax2.set_ylabel('%')

        ax3.plot(df.index, df[output_col], label="Fault", color="k")
        ax3.set_xlabel('Date')
        ax3.set_ylabel('Fault Flags')
        ax3.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc1_flag"
        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)
        # print("DAYS ALL DATA: ", total_days)
        total_hours = delta.sum() / pd.Timedelta(hours=1)
        # print("TOTAL HOURS: ", total_hours)
        hours_fc1_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)
        # print("FALT FLAG TRUE TOTAL HOURS: ", hours_fc1_mode)
        percent_true = round(df[output_col].mean() * 100, 2)
        # print("PERCENT TIME WHEN FLAG IS TRUE: ", percent_true, "%")
        percent_false = round((100 - percent_true), 2)
        # print("PERCENT TIME WHEN FLAG 5 FALSE: ", percent_false, "%")
        flag_true_duct_static = round(
            df[self.duct_static_col].where(df[output_col] == 1).mean(), 2
        )
        return (
            total_days,
            total_hours,
            hours_fc1_mode,
            percent_true,
            percent_false,
            flag_true_duct_static,
        )

    def create_hist_plot(
        self, df: pd.DataFrame, output_col: str = None, duct_static_col: str = None
    ) -> plt:
        if output_col is None:
            output_col = "fc1_flag"
        if duct_static_col is None:
            duct_static_col = "duct_static"
        # calculate dataset statistics
        df["hour_of_the_day_fc1"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc3
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc1.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 1 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None,
        duct_static_col: str = None,
        flag_true_duct_static: bool = None,
    ) -> None:
        
        if output_col is None:
            output_col = "fc1_flag"

        df[output_col] = df[output_col].astype(int)
        
        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition One Report", 0)

        p = document.add_paragraph(
            """Fault condition one of ASHRAE Guideline 36 is related to flagging poor performance of a AHU variable supply fan attempting to control to a duct pressure setpoint. Fault condition equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc1_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_fan_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)
        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc1_mode,
            percent_true,
            percent_false,
            flag_true_duct_static,
        ) = self.summarize_fault_times(df, output_col=output_col)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc1_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:

            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f'Average duct system pressure for when in fault condition (fan VFD speed > 95%): {flag_true_duct_static}"WC'
            )

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.')

        paragraph = document.add_paragraph()

        # ADD in Summary Statistics of fan operation
        document.add_heading("VFD Speed Statistics", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(str(df[self.supply_vfd_speed_col].describe()))

        # ADD in Summary Statistics of duct pressure
        document.add_heading("Duct Pressure Statistics", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(str(df[self.duct_static_col].describe()))

        # ADD in Summary Statistics of duct pressure
        document.add_heading("Duct Pressure Setpoints Statistics", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(str(df[self.duct_static_setpoint_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true > 5.0:
            paragraph.add_run(
                "The percent True metric that represents the amount of time for when the fault flag is True is high indicating the fan is running at high speeds and appearing to not generate good duct static pressure"
            )

        else:
            paragraph.add_run(
                "The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the fan appears to generate good duct static pressure"
            )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if df[self.duct_static_setpoint_col].std() == 0:
            paragraph.add_run("No duct pressure setpoint reset detected (BAD)")

        else:
            paragraph.add_run("Duct pressure reset detected (Good)")

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document


class FaultCodeTwoReport:
    """Class provides the definitions for Fault Code 2 Report."""

    def __init__(
        self,
        mix_degf_err_thres: float,
        return_degf_err_thres: float,
        outdoor_degf_err_thres: float,
        mat_col: str,
        rat_col: str,
        oat_col: str,
    ):
        self.mix_degf_err_thres = mix_degf_err_thres
        self.return_degf_err_thres = return_degf_err_thres
        self.outdoor_degf_err_thres = outdoor_degf_err_thres
        self.mat_col = mat_col
        self.rat_col = rat_col
        self.oat_col = oat_col

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:
        if output_col is None:
            output_col = "fc2_flag"

        df[output_col] = df[output_col].astype(int)

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(25, 8))
        plt.title('Fault Conditions 2 Plot')

        plot1a, = ax1.plot(df.index, df[self.mat_col],
                           color='r', label="Mix Temp")  # red

        plot1b, = ax1.plot(df.index, df[self.rat_col],
                           color='b', label="Return Temp")  # blue

        plot1c, = ax1.plot(df.index, df[self.oat_col],
                           color='g', label="Out Temp")  # green

        ax1.legend(loc='best')
        ax1.set_ylabel("°F")

        ax2.plot(df.index, df[output_col], label="Fault", color="k")
        ax2.set_xlabel('Date')
        ax2.set_ylabel('Fault Flags')
        ax2.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc2_flag"
        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)
        # print("DAYS ALL DATA: ", total_days)
        total_hours = delta.sum() / pd.Timedelta(hours=1)
        # print("TOTAL HOURS: ", total_hours)
        hours_fc2_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)
        # print("FALT FLAG TRUE TOTAL HOURS: ", hours_fc1_mode)
        percent_true = round(df[output_col].mean() * 100, 2)
        # print("PERCENT TIME WHEN FLAG IS TRUE: ", percent_true, "%")
        percent_false = round((100 - percent_true), 2)
        # print("PERCENT TIME WHEN FLAG 5 FALSE: ", percent_false, "%")

        flag_true_mat = round(
            df[self.mat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_oat = round(
            df[self.oat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_rat = round(
            df[self.rat_col].where(df[output_col] == 1).mean(), 2
        )
        return (
            total_days,
            total_hours,
            hours_fc2_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_oat,
            flag_true_rat
        )

    def create_hist_plot(
        self, df: pd.DataFrame,
        output_col: str = None,
        mat_col: str = None
    ) -> plt:

        if output_col is None:
            output_col = "fc2_flag"

        if mat_col is None:
            mat_col = "mat"

        # calculate dataset statistics
        df["hour_of_the_day_fc2"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc3
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc2.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 2 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None,
        mat_col: str = None,
        flag_true_mat: bool = None,
    ) -> None:

        if output_col is None:
            output_col = "fc2_flag"

        df[output_col] = df[output_col].astype(int)

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Two Report", 0)

        p = document.add_paragraph(
            """Fault condition two and three of ASHRAE Guideline 36 is related to flagging mixing air temperatures of the AHU that are out of acceptable ranges. Fault condition 2 flags mixing air temperatures that are too low and fault condition 3 flags mixing temperatures that are too high when in comparision to return and outside air data. The mixing air temperatures in theory should always be in between the return and outside air temperatures ranges. Fault condition two equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc2_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc2_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_oat,
            flag_true_rat
        ) = self.summarize_fault_times(df, output_col=output_col)

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc2_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:

            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()
            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 2 is True the average mix air temp is {flag_true_mat}°F, outside air temp is {flag_true_oat}°F, and return air temp is {flag_true_rat}°F. This could possibly help with pin pointing AHU operating conditions for when this fault is True.')

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.')

        paragraph = document.add_paragraph()

        # ADD in Summary Statistics
        document.add_heading('Mix Temp Statistics', level=2)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df[self.mat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Return Temp Statistics', level=2)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df[self.rat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Outside Temp Statistics', level=2)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df[self.oat_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true < 5:

            paragraph.add_run(
                'The percent True of time in fault condition 2 or 3 is high indicating the AHU temperature sensors are out of calibration')

        else:
            paragraph.add_run(
                'The percent True of time is low inidicating the AHU temperature sensors are within calibration')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document


class FaultCodeThreeReport:
    """Class provides the definitions for Fault Code 3 Report."""

    def __init__(
        self,
        mix_degf_err_thres: float,
        return_degf_err_thres: float,
        outdoor_degf_err_thres: float,
        mat_col: str,
        rat_col: str,
        oat_col: str,
    ):
        self.mix_degf_err_thres = mix_degf_err_thres
        self.return_degf_err_thres = return_degf_err_thres
        self.outdoor_degf_err_thres = outdoor_degf_err_thres
        self.mat_col = mat_col
        self.rat_col = rat_col
        self.oat_col = oat_col

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:
        if output_col is None:
            output_col = "fc3_flag"

        df[output_col] = df[output_col].astype(int)

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(25, 8))
        plt.title('Fault Conditions 3 Plot')

        plot1a, = ax1.plot(df.index, df[self.mat_col],
                           color='r', label="Mix Temp")  # red

        plot1b, = ax1.plot(df.index, df[self.rat_col],
                           color='b', label="Return Temp")  # blue

        plot1c, = ax1.plot(df.index, df[self.oat_col],
                           color='g', label="Out Temp")  # green

        ax1.legend(loc='best')
        ax1.set_ylabel("°F")

        ax2.plot(df.index, df[output_col], label="Fault", color="k")
        ax2.set_xlabel('Date')
        ax2.set_ylabel('Fault Flags')
        ax2.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc3_flag"
        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)
        # print("DAYS ALL DATA: ", total_days)
        total_hours = delta.sum() / pd.Timedelta(hours=1)
        # print("TOTAL HOURS: ", total_hours)
        hours_fc3_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)
        # print("FALT FLAG TRUE TOTAL HOURS: ", hours_fc1_mode)
        percent_true = round(df[output_col].mean() * 100, 2)
        # print("PERCENT TIME WHEN FLAG IS TRUE: ", percent_true, "%")
        percent_false = round((100 - percent_true), 2)
        # print("PERCENT TIME WHEN FLAG 5 FALSE: ", percent_false, "%")

        flag_true_mat = round(
            df[self.mat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_oat = round(
            df[self.oat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_rat = round(
            df[self.rat_col].where(df[output_col] == 1).mean(), 2
        )
        return (
            total_days,
            total_hours,
            hours_fc3_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_oat,
            flag_true_rat
        )

    def create_hist_plot(
        self, df: pd.DataFrame,
        output_col: str = None,
        mat_col: str = None
    ) -> plt:

        if output_col is None:
            output_col = "fc3_flag"

        if mat_col is None:
            mat_col = "mat"

        # calculate dataset statistics
        df["hour_of_the_day_fc3"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc3
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc3.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 3 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None,
        mat_col: str = None,
        flag_true_mat: bool = None,
    ) -> None:

        if output_col is None:
            output_col = "fc3_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Three Report", 0)

        p = document.add_paragraph(
            """Fault condition two and three of ASHRAE Guideline 36 is related to flagging mixing air temperatures of the AHU that are out of acceptable ranges. Fault condition 2 flags mixing air temperatures that are too low and fault condition 3 flags mixing temperatures that are too high when in comparision to return and outside air data. The mixing air temperatures in theory should always be in between the return and outside air temperatures ranges. Fault condition three equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc3_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc3_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_oat,
            flag_true_rat
        ) = self.summarize_fault_times(df, output_col=output_col)

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc3_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:

            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 3 is True the average mix air temp is {flag_true_mat}°F, outside air temp is {flag_true_oat}°F, and return air temp is {flag_true_rat}°F. This could possibly help with pin pointing AHU operating conditions for when this fault is True.')

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.')

        paragraph = document.add_paragraph()

        # ADD in Summary Statistics
        document.add_heading('Mix Temp Statistics', level=2)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df[self.mat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Return Temp Statistics', level=2)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df[self.rat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Outside Temp Statistics', level=2)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df[self.oat_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true < 5:

            paragraph.add_run(
                'The percent True of time in fault condition 2 or 3 is high indicating the AHU temperature sensors are out of calibration')

        else:
            paragraph.add_run(
                'The percent True of time is low inidicating the AHU temperature sensors are within calibration')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document


class FaultCodeFourReport:
    """Class provides the definitions for Fault Code 4 Report.
        Reporting the time series avg df that calculates control states per hour

    """

    def __init__(self,
                 delta_os_max: float
                 ):
        self.delta_os_max = delta_os_max
        self.heating_mode_calc_col = 'heating_mode'
        self.econ_only_cooling_mode_calc_col = 'econ_only_cooling_mode'
        self.econ_plus_mech_cooling_mode_calc_col = 'econ_plus_mech_cooling_mode'
        self.mech_cooling_only_mode_calc_col = 'mech_cooling_only_mode'

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:

        if output_col is None:
            output_col = "fc4_flag"

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(25, 8))
        plt.title('Fault Conditions 4 Plots')

        plot1a, = ax1.plot(df.index,
                           df[self.heating_mode_calc_col],
                           label="Heat",
                           color='orange')  # orange

        plot1b, = ax1.plot(df.index,
                           df[self.econ_only_cooling_mode_calc_col],
                           label="Econ Clg",
                           color='olive')  # olive

        plot1c, = ax1.plot(df.index,
                           df[self.econ_plus_mech_cooling_mode_calc_col],
                           label="Econ + Mech Clg",
                           color='c')  # cyan

        plot1d, = ax1.plot(df.index,
                           df[self.mech_cooling_only_mode_calc_col],
                           label="Mech Clg",
                           color='m')  # black

        ax1.set_xlabel('Date')
        ax1.set_ylabel('Calculated AHU Operating States')
        ax1.legend(loc='best')

        ax2.plot(df.index, df[output_col], label="Fault", color="k")
        ax2.set_xlabel('Date')
        ax2.set_ylabel('Fault Flags')
        ax2.legend(loc='best')

        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:

        if output_col is None:
            output_col = "fc4_flag"

        # calculate dataset statistics
        delta_all_data = df.index.to_series().diff()

        total_days_all_data = round(
            delta_all_data.sum() / pd.Timedelta(days=1), 2)

        total_hours_all_data = delta_all_data.sum() / pd.Timedelta(hours=1)

        hours_fc4_mode = (
            delta_all_data * df[output_col]).sum() / pd.Timedelta(hours=1)

        percent_true_fc4 = round(df.fc4_flag.mean() * 100, 2)
        percent_false_fc4 = round((100 - percent_true_fc4), 2)

        # heating mode runtime stats
        delta_heating = df[self.heating_mode_calc_col].index.to_series(
        ).diff()
        total_hours_heating = (
            delta_heating * df[self.heating_mode_calc_col]).sum() / pd.Timedelta(hours=1)

        percent_heating = round(
            df[self.heating_mode_calc_col].mean() * 100, 2)

        # econ mode runtime stats
        delta_econ = df[self.econ_only_cooling_mode_calc_col].index.to_series(
        ).diff()
        total_hours_econ = (
            delta_econ * df[self.econ_only_cooling_mode_calc_col]).sum() / pd.Timedelta(hours=1)

        percent_econ = round(
            df[self.econ_only_cooling_mode_calc_col].mean() * 100, 2)

        # econ plus mech cooling mode runtime stats
        delta_econ_clg = df[self.econ_plus_mech_cooling_mode_calc_col].index.to_series(
        ).diff()

        total_hours_econ_clg = (
            delta_econ_clg * df[self.econ_plus_mech_cooling_mode_calc_col]).sum() / pd.Timedelta(hours=1)

        percent_econ_clg = round(
            df[self.econ_plus_mech_cooling_mode_calc_col].mean() * 100, 2)

        # mech clg mode runtime stats
        delta_clg = df[self.mech_cooling_only_mode_calc_col].index.to_series(
        ).diff()

        total_hours_clg = (
            delta_clg * df[self.mech_cooling_only_mode_calc_col]).sum() / pd.Timedelta(hours=1)

        percent_clg = round(
            df[self.mech_cooling_only_mode_calc_col].mean() * 100, 2)

        return (
            total_days_all_data,
            total_hours_all_data,
            hours_fc4_mode,
            percent_true_fc4,
            percent_false_fc4,
            percent_clg,
            percent_econ_clg,
            percent_econ,
            percent_heating,
            total_hours_heating,
            total_hours_econ,
            total_hours_econ_clg,
            total_hours_clg
        )

    def create_hist_plot(
        self, df: pd.DataFrame,
        output_col: str = None,
    ) -> plt:

        if output_col is None:
            output_col = "fc4_flag"

        # calculate dataset statistics
        df["hour_of_the_day_fc4"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc4
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc4.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 4 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None,
    ) -> None:
        if output_col is None:
            output_col = "fc4_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Four Report", 0)

        p = document.add_paragraph(
            """Fault condition four of ASHRAE Guideline 36 is related to flagging AHU control programming that is hunting between heating, economizing, economizing plus mechanical cooling, and mechanical cooling operating states. This fault diagnostic does NOT flag simultaneous heating and cooling, just excessive cycling between the states or operating modes the AHU maybe going in and out of. Fault condition four equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc4_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days_all_data,
            total_hours_all_data,
            hours_fc4_mode,
            percent_true_fc4,
            percent_false_fc4,
            percent_clg,
            percent_econ_clg,
            percent_econ,
            percent_heating,
            total_hours_heating,
            total_hours_econ,
            total_hours_econ_clg,
            total_hours_clg
        ) = self.summarize_fault_times(df, output_col=output_col)

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in days calculated in dataset: {total_days_all_data}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours calculated in dataset: {total_hours_all_data}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc4_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true_fc4}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false_fc4}%"
        )

        document.add_heading("Calculated AHU Mode Statistics", level=2)

        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(
            f'Total time in hours while AHU is in a heating mode: {total_hours_heating}')

        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(
            f'Total percent time in while AHU is in a heating mode: {percent_heating}%')

        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(
            f'Total time in hours while AHU is in a economizing mode: {total_hours_econ}')

        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(
            f'Total percent time in while AHU is in a economizing mode: {percent_econ}%')

        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(
            f'Total time in hours while AHU is in a economizing plus mechanical cooling mode: {total_hours_econ_clg}')

        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(
            f'Total percent time in while AHU is in a economizing plus mechanical cooling mode: {percent_econ_clg}%')

        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(
            f'Total time in hours while AHU is in a mechanical cooling mode: {total_hours_clg}')

        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(
            f'Total percent time in while AHU is in a mechanical cooling mode: {percent_clg}%')

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:

            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'Fault condition 4 is True because of excessive cycling between different control system operation modes.')

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.')

        paragraph = document.add_paragraph()

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if fc_max_faults_found >= self.delta_os_max:

            paragraph.add_run(
                f'The AHU control system needs tuning to reduce control loop hunting for setpoints. Its hunting or overshooting setpoints which can cause AHU systems to be oscillating (most likely too fast) between heating and cooling modes without never settling out. Low load conditions can also cause excessive cycling if heating or cooling setpoints are met very fast. Verify that the times when this fault is flagged that no occupant comfort issues persist. Fixing this fault may also improve energy efficiency and extend the mechanical equipment life span with the prevention of excessive cycling especially cooling compressors.')

        else:
            paragraph.add_run(
                f'No control system tuning appears to be needed for the operating conditions of this AHU.')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document
    
    
class FaultCodeFiveReport:
    """Class provides the definitions for Fault Code 5 Report."""

    def __init__(
        self,
        mix_degf_err_thres: float,
        supply_degf_err_thres: float,
        delta_t_supply_fan: float,
        mat_col: str,
        sat_col: str,
    ):
        self.mix_degf_err_thres = mix_degf_err_thres
        self.supply_degf_err_thres = supply_degf_err_thres
        self.delta_t_supply_fan = delta_t_supply_fan
        self.mat_col = mat_col
        self.sat_col = sat_col

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:
        if output_col is None:
            output_col = "fc5_flag"

        df[output_col] = df[output_col].astype(int)

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(25, 8))
        plt.title('Fault Conditions 2 Plot')

        plot1a, = ax1.plot(df.index, df[self.mat_col],
                           color='r', label="Mix Temp")  # red

        plot1b, = ax1.plot(df.index, df[self.sat_col],
                           color='b', label="Supply Temp")  # blue

        ax1.legend(loc='best')
        ax1.set_ylabel("°F")

        ax2.plot(df.index, df[output_col], label="Fault", color="k")
        ax2.set_xlabel('Date')
        ax2.set_ylabel('Fault Flags')
        ax2.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc5_flag"
            
        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)

        total_hours = delta.sum() / pd.Timedelta(hours=1)

        hours_fc5_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)

        percent_true = round(df[output_col].mean() * 100, 2)
        percent_false = round((100 - percent_true), 2)


        flag_true_mat = round(
            df[self.mat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_sat = round(
            df[self.sat_col].where(df[output_col] == 1).mean(), 2
        )

        return (
            total_days,
            total_hours,
            hours_fc5_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_sat,
        )

    def create_hist_plot(
        self, df: pd.DataFrame,
        output_col: str = None,
        mat_col: str = None
    ) -> plt:

        if output_col is None:
            output_col = "fc5_flag"

        if mat_col is None:
            mat_col = "mat"

        # calculate dataset statistics
        df["hour_of_the_day_fc5"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc5
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc5.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 5 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None,
        mat_col: str = None,
        flag_true_mat: bool = None,
    ) -> None:

        if output_col is None:
            output_col = "fc5_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Five Report", 0)

        p = document.add_paragraph(
            """Fault condition five of ASHRAE Guideline 36 is (an AHU heating mode or winter time conditions only fault equation) related to flagging supply air temperatures that are out of acceptable ranges based on the mix air temperature and an assumption for heat created by the AHU supply fan in the air stream. Fault condition five equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc5_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc5_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_sat,

        ) = self.summarize_fault_times(df, output_col=output_col)

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc5_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:

            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 5 is True the average mix air temp is {flag_true_mat}°F and the outside air temp is {flag_true_sat}°F. This could possibly help with pin pointing AHU operating conditions for when this fault is True.')

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.')

        paragraph = document.add_paragraph()

        # ADD in Summary Statistics
        document.add_heading('Mix Temp Statistics', level=2)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df[self.mat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading('Supply Temp Statistics', level=2)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df[self.sat_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true > 5.0:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is high indicating the AHU temperature sensors for either the supply or mix temperature are out of calibration. Verify the mixing temperature sensor is not a probe type sensor but a long averaging type sensor that is installed properly inside the AHU mixing chamber to get a good solid true reading of the actual air mixing temperature. Poor duct design may also contribute to not having good air mixing, to troubleshoot install data loggers inside the mixing chamber or take measurements when the AHU is running of different locations in the mixing chamber to spot where better air blending needs to take place.')

        else:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the AHU temperature sensors are within calibration')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document

