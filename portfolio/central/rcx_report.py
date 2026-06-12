"""RCx report builder — python-docx + in-memory matplotlib charts."""

from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Any

from portfolio.central.fault_hours import aggregate_fault_hours, fault_summary_from_validation


def _chart_bytes(title: str, labels: list[str], values: list[float]) -> bytes:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(6, 3))
    ax.bar(labels[:12], values[:12], color="#2563eb")
    ax.set_title(title)
    ax.tick_params(axis="x", rotation=45, labelsize=8)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=120)
    plt.close(fig)
    return buf.getvalue()


def build_rcx_docx(
    *,
    site_id: str,
    site_name: str,
    validation: dict[str, Any] | None = None,
    rollups: list[dict[str, Any]] | None = None,
    warnings: list[str] | None = None,
) -> bytes:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading(f"Open-FDD RCx Report — {site_name}", level=0)
    doc.add_paragraph(f"Site ID: {site_id}")
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph(
        "Read-only portfolio report from edge REST APIs. No BACnet writes or equipment commands."
    )

    doc.add_heading("Executive summary", level=1)
    if validation:
        ok = validation.get("ok")
        doc.add_paragraph(f"Latest validation: {'PASS' if ok else 'FAIL'}")
        for err in validation.get("errors") or []:
            doc.add_paragraph(f"• {err}")
    else:
        doc.add_paragraph("No live validation snapshot included.")

    doc.add_heading("Active faults by equipment", level=1)
    if validation:
        rows = fault_summary_from_validation(validation)
        if not rows:
            doc.add_paragraph("No active FDD faults at snapshot time.")
        for row in rows:
            doc.add_paragraph(
                f"{row['equipment']} ({row.get('equipment_type') or '—'}) — "
                f"{row['severity']}: {row['title'] or row['code']}"
            )
    else:
        doc.add_paragraph("Equipment context unavailable without validation snapshot.")

    doc.add_heading("Fault-hour estimates", level=1)
    hours_rows = aggregate_fault_hours(rollups or [])
    if hours_rows:
        for row in hours_rows[:20]:
            doc.add_paragraph(
                f"{row['fault_code']}: ~{row['elapsed_hours']} h (site {row['site_id']})"
            )
        labels = [r["fault_code"] for r in hours_rows[:10]]
        values = [r["elapsed_hours"] for r in hours_rows[:10]]
        if labels:
            png = _chart_bytes("Fault hours by code", labels, values)
            doc.add_picture(io.BytesIO(png), width=Inches(5.5))
    else:
        doc.add_paragraph("Insufficient rollup history for fault-hour chart.")

    doc.add_heading("Warnings", level=1)
    for w in warnings or []:
        doc.add_paragraph(f"• {w}")
    if not warnings:
        doc.add_paragraph("None recorded.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
