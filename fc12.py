
import operator
import time
from datetime import datetime, timedelta

import pandas as pd
import numpy as np

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from docx import Document
from docx.shared import Inches

import argparse
import math

# python 3.10 on Windows 10
# py .\fc12.py -i ./ahu_data/hvac_random_fake_data/fc12_fake_data1.csv -o fake1_ahu_fc12_report

parser = argparse.ArgumentParser(add_help=False)
args = parser.add_argument_group('Options')

args.add_argument('-h', '--help', action='help',
                  help='Show this help message and exit.')
args.add_argument('-i', '--input', required=True, type=str,
                  help='CSV File Input')
args.add_argument('-o', '--output', required=True, type=str,
                  help='Word File Output Name')
'''
FUTURE 
 * incorporate an arg for SI units 
 * °C on temp sensors
 * piping pressure sensor PSI conversion
 * air flow CFM conversion
 * AHU duct static pressure "WC

args.add_argument('--use-SI-units', default=False, action='store_true')
args.add_argument('--no-SI-units', dest='use-SI-units', action='store_false')
'''
args = parser.parse_args()


def fault_condition_twelve(df):
    return df.sat_minus_saterr_delta_supply_fan >= df.mat_plus_materr


df = pd.read_csv(args.input,
                 index_col='Date',
                 parse_dates=True).rolling('5T').mean()

print(df)
df_copy = df.copy()

df_copy.plot(figsize=(25, 8),
             title='AHU Temp Sensors')
plt.savefig('./static/ahu_fc12_signals.png')

# make an entire column out of these params in the Pandas Dataframe
# required params taken from the screenshot above

DELTA_SUPPLY_FAN = 2
MIX_DEGF_ERR_THRES = 5
SUPPLY_DEGF_ERR_THRES = 2
df['delta_supply_fan'] = DELTA_SUPPLY_FAN
df['mix_err_thres'] = MIX_DEGF_ERR_THRES
df['supply_err_thres'] = SUPPLY_DEGF_ERR_THRES
df['sat_minus_saterr_delta_supply_fan'] = df.sat - df.supply_err_thres - df.delta_supply_fan
df['mat_plus_materr'] = df.mat + df.mix_err_thres


start = df.head(1).index.date
print('Dataset start: ', start)

end = df.tail(1).index.date
print('Dataset end: ', end)
print('COLUMNS: ', print(df.columns))

df['fc12_flag'] = fault_condition_twelve(df)

df2 = df.copy().dropna()
df2['fc12_flag'] = df2['fc12_flag'].astype(int)

# drop params column for better plot
df2 = df2.drop(['delta_supply_fan',
                'mix_err_thres',
                'supply_err_thres'], axis=1)

print(df2)

fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(25, 8))
plt.title('Fault Conditions 12 Plot')

plot1a, = ax1.plot(df2.index, df2.sat, label="SAT")
plot1b, = ax1.plot(df2.index, df2.mat, label="MAT")
ax1.legend(loc='best')
ax1.set_ylabel('AHU Supply & Mix Temps °F')

ax2.plot(df2.index, df2.fc12_flag, label="Fault", color="k")
ax2.set_xlabel('Date')
ax2.set_ylabel('Fault Flags')
ax2.legend(loc='best')

plt.legend()
plt.tight_layout()
plt.savefig('./static/ahu_fc12_fans_plot.png')
# plt.show()

print("Starting ahu fc12 docx report")
document = Document()
document.add_heading('Fault Condition Twelve Report', 0)

p = document.add_paragraph(
    'Fault condition twelve of ASHRAE Guideline 36 is an AHU economizing mode only fault equation with an attempt at verifying an AHU mixing and supply temperature accuracy. Fault condition twelve equation as defined by ASHRAE:')
document.add_picture('./images/fc12_definition.png', width=Inches(6))

# ADD IN SUBPLOTS SECTION
document.add_heading('Dataset Plot', level=2)
document.add_picture('./static/ahu_fc12_fans_plot.png', width=Inches(6))
document.add_heading('Dataset Statistics', level=2)

# calculate dataset statistics
delta = df2.index.to_series().diff()
total_days = round(delta.sum() / pd.Timedelta(days=1), 2)
print('DAYS ALL DATA: ', total_days)
total_hours = delta.sum() / pd.Timedelta(hours=1)
print('TOTAL HOURS: ', total_hours)
hours_fc12_mode = (delta * df2["fc12_flag"]).sum() / pd.Timedelta(hours=1)
print('FALT FLAG TRUE TOTAL HOURS: ', hours_fc12_mode)
percent_true = round(df2.fc12_flag.mean() * 100, 2)
print('PERCENT TIME WHEN Flag 12 TRUE: ', percent_true, '%')
percent_false = round((100 - percent_true), 2)
print('PERCENT TIME WHEN Flag 12 FALSE: ', percent_false, '%')
df2['hour_of_the_day_fc12'] = df2.index.hour.where(df2["fc12_flag"] == 1)

flag_true_sat = round(
    df2.sat.where(df2["fc12_flag"] == 1).mean(), 2)
print('FLAG TRUE SAT DEGF: ', flag_true_sat)

# make hist plots fc12
fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
ax.hist(df2.hour_of_the_day_fc12.dropna())
ax.set_xlabel('24 Hour Number in Day')
ax.set_ylabel('Frequency')
ax.set_title(f'Hour-Of-Day When Fault Flag 12 is TRUE')
fig.savefig('./static/ahu_fc12_histogram.png')

# add calcs to word doc
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time in days calculated in dataset: {total_days}')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time in hours calculated in dataset: {total_hours}')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time in hours for when fault Flag 12 is True: {hours_fc12_mode}')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Percent of time in the dataset when the fault Flag 12 is True: {percent_true}%')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Percent of time in the dataset when fault Flag 12 is False: {percent_false}%')


if hours_fc12_mode != float(0):
    paragraph = document.add_paragraph()
    # ADD HIST Plots
    document.add_heading('Time-of-day Histogram Plots', level=2)
    document.add_picture('./static/ahu_fc12_histogram.png', width=Inches(6))

if not math.isnan(flag_true_sat):
    paragraph = document.add_paragraph()
    paragraph.style = 'List Bullet'
    paragraph.add_run(
        f'When fault condition 12 is True the average supply temperature is {flag_true_sat}°F. This data along with time-of-day could possibly help with pin pointing AHU operating conditions for when this fault is True.')

paragraph = document.add_paragraph()

# ADD in Summary Statistics
document.add_heading('Supply Temp Statistics', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(str(df2.sat.describe()))

# ADD in Summary Statistics
document.add_heading('Mix Temp Statistics', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(str(df2.mat.describe()))

document.add_heading('Suggestions based on data analysis', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'

if percent_true > 5.0:
    paragraph.add_run(
        'The percent True metric that represents the amount of time for when the fault flag is True is high indicating the AHU temperature sensors are out of calibration')

else:
    paragraph.add_run(
        'The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the AHU temperature sensors are within calibration')

print('df2.sat.std: ', df2.sat.std())
print('df2.sat.min: ', df2.sat.min())
print('df2.sat.max: ', df2.sat.max())

paragraph = document.add_paragraph()
run = paragraph.add_run(f'Report generated: {time.ctime()}')
run.style = 'Emphasis'

document.save(f'./final_report/{args.output}.docx')
print('All Done')

# df2.to_csv('testdf2_fc12.csv')
