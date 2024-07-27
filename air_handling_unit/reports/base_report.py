import time
from io import BytesIO

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches



class BaseReport:
    def __init__(self, config):
        self.config = config

    def create_plot(self, df: pd.DataFrame, cols: list, labels: list, title: str) -> plt:
        fig, axes = plt.subplots(len(cols), 1, figsize=(25, 8))
        plt.title(title)

        for ax, col, label in zip(axes, cols, labels):
            ax.plot(df.index, df[col], label=label)
            ax.legend(loc='best')

        plt.tight_layout()
        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str) -> dict:
        delta = df.index.to_series().diff().dt.total_seconds()
        total_days = round(delta.sum() / 86400, 2)
        total_hours = round(delta.sum() / 3600, 2)
        hours_fault_mode = (delta * df[output_col]).sum() / 3600
        percent_true = round(df[output_col].mean() * 100, 2)
        percent_false = round((100 - percent_true), 2)

        # Calculate motor runtime
        motor_on = df[self.config['SUPPLY_VFD_SPEED_COL']].gt(.01).astype(int)
        hours_motor_runtime = round((delta * motor_on).sum() / 3600, 2)

        summary = {
            'total_days': total_days,
            'total_hours': total_hours,
            'hours_fault_mode': hours_fault_mode,
            'percent_true': percent_true,
            'percent_false': percent_false,
            'hours_motor_runtime': hours_motor_runtime
        }
        print("SUMMARY: ", summary)
        return summary

    def create_hist_plot(self, df: pd.DataFrame, output_col: str) -> plt:
        df["hour_of_the_day"] = df.index.hour.where(df[output_col] == 1)
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day.dropna())
        ax.set_xlabel("Hour of the Day")
        ax.set_ylabel("Frequency")
        ax.set_title("Hour-Of-Day When Fault Flag is TRUE")
        return fig

    def add_summary_to_document(self, document: Document, summary: dict):
        document.add_heading("Dataset Statistics", level=2)
        for key, value in summary.items():
            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(f"{key.replace('_', ' ').title()}: {value}")

    def create_report(self, path: str, df: pd.DataFrame, output_col: str) -> None:
        document = Document()
        document.add_heading("Fault Condition Report", 0)

        self.add_summary_to_document(document, self.summarize_fault_times(df, output_col))

        fig = self.create_hist_plot(df, output_col)
        hist_plot_image = BytesIO()
        fig.savefig(hist_plot_image, format="png")
        hist_plot_image.seek(0)
        document.add_picture(hist_plot_image, width=Inches(6))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        percent_true = self.summarize_fault_times(df, output_col)['percent_true']
        if percent_true > 5:
            paragraph.add_run("The percent True is high, indicating issues.")
        else:
            paragraph.add_run("The percent True is low, indicating normal operation.")

        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        document.save(f"{path}/report.docx")
