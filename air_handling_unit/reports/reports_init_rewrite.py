import pandas as pd
import numpy as np
from pydantic import BaseModel
from typing import Optional
import matplotlib as mpl
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import math
import os
import time
from io import BytesIO
from num2words import num2words
import itertools

class Sensor(BaseModel):
    """Gives the different attributes for relevant sensors used in each fault.

    Atrributes:
        col_name            The standard column name any time this sensor is in a Pandas data frame
        long_name           The standard full sensor name, how it's referred to in fault reports
        measurement         How the sensor is measured (temperature, percent, pressure, etc.)
        short_name          The sensor's corresponding short name, mainly used for legends in graphs
    """
    col_name: str
    long_name: str
    measurement: str
    units: str
    short_name: Optional[str] = None
    avg_fault_val: Optional[float] = None
    op_time: Optional[list] = None

    def __init__(self,**data):
        super().__init__(**data)
        if 'short_name' not in data.keys():
            super().__init__(short_name = data['long_name'].title(),**data)

class Fault(BaseModel):
    """
    Gives a fault's relevant attributes for generating final reports.

    Attributes:
        num                     The fault number as an integer
        definition              The text description for the fault's definition
        sensors                 All of the fault's relevant sensors (not including supply_vfd_speed, which is always assumed to exist)
        suggestion_fault_high   Text to display if this fault occurs too often
        suggestion_fault_low    Text to display if this fault does not occur too often

    """
    num: int
    definition: str
    sensors: list
    suggestion_high_fault: str
    suggestion_low_fault: str

    def __init__(self, col_names, **data):
        # do this so we can create the sensor objects from ALL_SENSORS whose col_names match the provided col_name
        super().__init__(sensors = [sensor for sensor in ALL_SENSORS if sensor.col_name in col_names],**data)

        # super().__init__(sensors = [Sensor(**dict(zip(sensor_attrs, sensor_def))) for sensor_def in sensor_defs if sensor_def[0] in col_names], **data)

class Calculator:
    """
    Calculates the data fed to the report. Assumes 'supply_vfd_speed' col exists in the df.
    """
    def __init__(self, fault: Fault, df: pd.DataFrame):
        self.df = df
        self.fault_col = f"fc{fault.num}_flag"

        self.compile_stats(fault)

    def summarize_avg_faults(self,fault: Fault):
        """
        Store each sensor's average fault value as an attribute of that sensor.
        """
        for sensor in fault.sensors:
            sensor.avg_fault_val = round(self.df[sensor.col_name].where(self.df[self.fault_col] == 1).mean(), 2)

    def summarize_operational_time(self, df: pd.DataFrame):
        """
        Summarize the times this df is in the provided modes. Only works for boolean cols indexing when the mode is active, where 1 is "active" and 0 is "not active".
        """
        delta = df.index.to_series().diff()
        time_in_mode = (delta * df).sum()

        days_in_mode = round(time_in_mode / pd.Timedelta(days=1), 2)
        hours_in_mode = time_in_mode / pd.Timedelta(hours=1)
        percent_in_mode = round(df.mean() * 100, 2)

        return(days_in_mode, hours_in_mode, percent_in_mode)

    def compile_stats(self, fault: Fault):
        """
        Compile the relevant stats. These are summary stats for relevant data, as well as operational times for different operation modes.
        """
        df = self.df

        df['ident'] = 1 # to get total op time in the df
        df['motor_on'] = df['supply_vfd_speed'].gt(1.).astype(int)

        self.total_days, self.total_hours, __ = self.summarize_operational_time(df['ident'])
        __, self.hours_in_fault_mode, self.percent_in_fault_mode = self.summarize_operational_time(df[self.fault_col])
        __, self.hours_motor_runtime, __ = self.summarize_operational_time(df['motor_on'])

        for sensor in fault.sensors:
            if sensor.measurement == 'operating state':
                sensor.op_time = self.summarize_operational_time(df[sensor.col_name])
            else:
                sensor.avg_fault_val = round(self.df[sensor.col_name].where(self.df[self.fault_col] == 1).mean(), 2)

        self.summarize_avg_faults(fault)

        self.df_motor_on_filtered = self.df[self.df['supply_vfd_speed'] > 1.0]

class DocumentGenerator:
    """Class provides the skeleton for creating a report document."""

    def __init__(self, fault: int, df: pd.DataFrame, calculator: Calculator):
        self.fault = fault
        self.df = df
        self.df['fault_flag'] = self.df[f'fc{self.fault.num}_flag']
        self.document = Document()
        self.calculator = calculator

    def create_dataset_plot(self) -> plt:
        """
        Creates timeseries data plots. Does so by grouping this fault's sensors 
        according to their measurement type, creates one plot per measurement type, and
        flags the fault locations in blue.
        """

        # group sensors of interest by measurement type
        grouped_sensors = []
        for key, group in itertools.groupby(self.fault.sensors, key=lambda x: x.measurement):
            grouped_sensors.append(list(group))

        n_groups = len(grouped_sensors)

        fig, data_axes = plt.subplots(n_groups, 1, figsize=(25, 4*n_groups))
        # plt.title(f'Fault Condition {self.fault.num} Plot')

        # group df by runs of values in a col: 
        # ttps://towardsdatascience.com/pandas-dataframe-group-by-consecutive-certain-values-a6ed8e5d8cc
        fault_grouping = self.df[self.df['fault_flag'] == 1].groupby((self.df['fault_flag'] != 1).cumsum())

        # this is for changing the color palette choice for each subplot
        # plt.rcParams["axes.prop_cycle"] = plt.cycler("color", plt.cm.tab10c.colors)
        # cmap_list = ['Dark2', 'Set1', 'Set2', 'tab10'] # reasonable contrasting colormaps to use

        # data_axes groups the different axes by measurement, so they can be graphed on similar scales
        if n_groups == 1:
            data_axes = [data_axes]

        for i in range(n_groups):
            group = grouped_sensors[i]
            for sensor in group:
                data_axes[i].plot(self.df.index, self.df[sensor.col_name], label=sensor.short_name)

                # highlight each span of faults
                for k, v in fault_grouping:
                    data_axes[i].axvspan(v.index[0], v.index[-1], alpha=0.1, color = 'blue')

            measurement = list(set([sensor.measurement for sensors in group]))[0]
            units = list(set([sensor.units for sensors in group]))[0]

            data_axes[i].legend(loc='best')
            data_axes[i].set_ylabel(f'{measurement} ({units})')

        plt.legend()
        plt.tight_layout()

        return fig

    # creates a histogram plot for times when fault condition is true
    def create_hist_plot(self) -> plt:
        # calculate dataset statistics
        self.df[f"hour_of_the_day_fc{self.fault.num}"] = self.df.index.hour.where(self.df['fault_flag'] == 1)

        # make hist plots fc10
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(self.df[f"hour_of_the_day_fc{self.fault.num}"].dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag {self.fault.num} is TRUE")
        return fig

    def format_title_and_def(self):
        # add title and fault definition
        self.document.add_heading(f"Fault Condition {num2words(self.fault.num).title()} Report", 0)

        fault_def_str = f'Fault condition {num2words(self.fault.num)} of ASHRAE Guideline 36 {self.fault.definition}. Fault condition {num2words(self.fault.num)} equation as defined by ASHRAE:'

        p = self.document.add_paragraph(fault_def_str)

        self.document.add_picture(
            os.path.join(os.path.curdir, "images", f"fc{self.fault.num}_definition.png"),
            width=Inches(6),
        )

    def format_dataset_plot(self):
        # add dataset plot
        self.document.add_heading("Dataset Plot", level=2)

        fig = self.create_dataset_plot()
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        self.document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )

    def format_dataset_stats(self):
        # add dataset statistics
        self.document.add_heading("Dataset Statistics", level=2)

        stats_lines = [
            f"Total time in days calculated in dataset: {self.calculator.total_days}",
            f"Total time in hours calculated in dataset: {self.calculator.total_hours}",
            f"Total time in hours for when fault flag is True: {self.calculator.hours_in_fault_mode}",
            f"Percent of time in the dataset when the fault flag is True: {self.calculator.percent_in_fault_mode}%",
            f"Percent of time in the dataset when the fault flag is False: {round((100 - self.calculator.percent_in_fault_mode), 2)}%",
            f"Calculated motor runtime in hours based off of VFD signal > zero: {self.calculator.hours_motor_runtime}"
        ]

        for line in stats_lines:
            self.document.add_paragraph(line, style="List Bullet")

    def format_hist_plot(self):
        # if there are faults, add the histogram plot
        fc_max_faults_found = self.df['fault_flag'].max()

        if fc_max_faults_found != 0:

            self.document.add_heading("Time-of-day Histogram Plots", level=2)

            max_faults_line = f'When fault condition {self.fault.num} is True, the'

            for sensor in self.fault.sensors:
                if sensor.measurement == 'temperature': # only temp sensors
                    max_faults_line =  f'{max_faults_line} average {sensor.long_name} is {sensor.avg_fault_val} {sensor.units}, ' # need to add {units}

            max_faults_line = max_faults_line[0:-2] + "."

            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot()
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            self.document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )
        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")
            max_faults_line = f'No faults were found in this given dataset for the equation defined by ASHRAE.'

        self.document.add_paragraph(max_faults_line, style='List Bullet')

    def format_summary_stats(self):
        data_types = set([sensor.measurement for sensor in self.fault.sensors])

        if 'operating state' in data_types: # summarize operating state times
            self.document.add_heading("Calculated AHU Mode Statistics")
        
            for sensor in self.fault.sensors:
                (days, hours, percent) = sensor.op_time

                self.document.add_paragraph(
                    f'Total time in hours while AHU is in {sensor.long_name}: {hours}',
                    style='List Bullet')
                
                self.document.add_paragraph(
                    f'Total percent time while AHU is in {sensor.long_name}: {percent}%',
                    style="List Bullet")
        else: # just summarize temperature levels for different temp sensors
            self.document.add_heading('Summary Statistics filtered for when the AHU is running', level=1)

            for sensor in self.fault.sensors:
                if sensor.measurement == 'temperature': # only temp sensors
                    self.document.add_heading(sensor.short_name, level=3)
                    self.document.add_paragraph(
                        str(self.calculator.df_motor_on_filtered[sensor.col_name].describe()),
                        style = 'List Bullet'
                        )

    def format_suggestion(self):
        self.document.add_heading("Suggestions based on data analysis", level=3)
 
        suggestion_string = 'The amount of time this fault is True is'

        if self.calculator.percent_in_fault_mode > 5.0:
            suggestion_string = f'{suggestion_string} high, indicating {self.fault.suggestion_high_fault}.'
                
        else:
            suggestion_string = f'{suggestion_string} low, indicating {self.fault.suggestion_low_fault}.'

        self.document.add_paragraph(suggestion_string, style="List Bullet")

        paragraph = self.document.add_paragraph()
        paragraph.add_run(f"Report generated: {time.ctime()}", style="Emphasis")

    def create_document(self, path: str) -> Document:
        self.format_title_and_def()
        self.format_dataset_plot()
        self.format_dataset_stats()
        self. format_hist_plot()
        self.format_summary_stats()
        self.format_suggestion()

        return self.document


sensor_attrs = ['col_name','long_name','measurement', 'units', 'short_name']

# this is just a simple way to store the sensors before we make them into objects
# this is ordered as: col_name, long_name, measurement, units, short_name (with short_name optional)
# the following should probably not live in code, would be better as a csv or json or database or whatever
slist = [['mat', 'mixing air temperature', 'temperature', '°F', 'Mix Temp'],
        ['rat', 'return air temperature', 'temperature', '°F', 'Return Temp'],
        ['oat', 'outside air temperature', 'temperature', '°F', 'Out Temp'],
        ['sat', 'supply air temperature', 'temperature', '°F', 'Supply Temp'],
        ['satsp', 'supply air temperature setpoint', 'temperature', '°F', 'Out Temp',],
        ['economizer_sig', 'outside air damper position', 'percent', '%', 'AHU Dpr Cmd'],
        ['clg', 'AHU cooling valve', 'percent', '%', 'AHU Cool Valv'],
        ['htg', 'AHU heating valve', 'percent', '%', 'AHU Htg Valv'],
        ['cooling_sig', 'cooling signal', 'operating state', 'flag', 'Mech Clg'],
        ['heating_sig', 'heating signal', 'operating state', 'flag', 'Heat'],
        ['vav_total_flow', 'VAV total air flow', 'volume flow', 'cfm', 'VAV total flow'],
        ['duct_static', 'duct static pressure', 'pressure','inches WC', 'Static'],
        ['duct_static_setpoint', 'duct static pressure setpoint', 'pressure', 'inches WC', 'Static SP'],
        ['heating_mode', 'heating mode','operating state', 'flag','Heat'],
        ['econ_only_cooling_mode','economizing mode','operating state', 'flag','Econ Clg'],
        ['econ_plus_mech_cooling_mode','economizing plus mechanical cooling mode','operating state', 'flag','Econ + Mech Clg'],
        ['mech_cooling_only_mode','mechanical cooling mode','operating state', 'flag','Mech Clg']
    ]

ALL_SENSORS = [Sensor(**dict(zip(sensor_attrs, sensor))) for sensor in slist]

fault_attrs = ['num', 'col_names', 'definition', 'suggestion_high_fault', 'suggestion_low_fault']

# same thing here -- this shouldn't be stored as code, should be separate files that can be easily tweaked
fault_defs = [
    [1, ['duct_static', 'duct_static_setpoint', 'supply_vfd_speed'], 
        'flags poor performance of a AHU variable supply fan attempting to control to a duct pressure setpoint',
        'the fan is running at high speeds and appearing to not generate good duct static pressure',
        'the fan appears to generate good duct static pressure'
        ],

    [2, ['mat','rat','oat'],
        'flags mixing air temperatures of the AHU that are too low compared to return and outside air data. The mixing air temperatures in theory should always be in between the return and outside air temperatures ranges',
        'the AHU temperature sensors are out of calibration',
        'the AHU temperature sensors are within calibration'
        ],

    [3, ['mat','rat','oat'], 
        'flags mixing air temperatures of the AHU that are too high compared to return and outside air data. The mixing air temperatures in theory should always be in between the return and outside air temperatures ranges',
        'the AHU temperature sensors are out of calibration',
        'the AHU temperature sensors are within calibration'
        ],

    [4, ['heating_mode', 'econ_only_cooling_mode', 'econ_plus_mech_cooling_mode', 'mech_cooling_only_mode'],
        'flags AHU control programming that is hunting between heating, economizing, economizing plus mechanical cooling, and mechanical cooling operating states. This fault diagnostic does NOT flag simultaneous heating and cooling, just excessive cycling between AHU operating modes',
        'the AHU control system needs tuning to reduce control loop hunting for setpoints. It is hunting or overshooting setpoints which can cause AHU systems to be oscillating (most likely too fast) between heating and cooling modes without never settling out. Low load conditions can also cause excessive cycling if heating or cooling setpoints are met very fast. Verify that the times when this fault is flagged that no occupant comfort issues persist. Fixing this fault may also improve energy efficiency and extend the mechanical equipment life span with the prevention of excessive cycling especially cooling compressors',
        'no control system tuning appears to be needed for the operating conditions of this AHU'
        ],

    [5, ['mat','sat','htg'], 
        'supply air temperatures that are out of acceptable ranges based on the mix air temperature and an assumption for heat created by the AHU supply fan in the air stream',
        'the AHU temperature sensors for either the supply or mix temperature are out of calibration. Verify the mixing temperature sensor is not a probe type sensor but a long averaging type sensor that is installed properly inside the AHU mixing chamber to get a good solid true reading of the actual air mixing temperature. Poor duct design may also contribute to not having good air mixing, to troubleshoot install data loggers inside the mixing chamber or take measurements when the AHU is running of different locations in the mixing chamber to spot where better air blending needs to take place',
        'the AHU temperature sensors are within calibration'
        ],

    [6, ['vav_total_flow','mat','oat','rat'],
        'attempts to verify that AHU design minimum outside air is close to the calculated outside air fraction through the outside, mix, and return air temperature sensors. A fault will get flagged if the OA fraction is too low or too high as compared to design OA',
        'sensors are out of calibration either on the AHU outside, mix, or return air temperature sensors that handle the OA fraction calculation or the totalized air flow calculation handled by a totalizing all VAV box air flows or AHU AFMS. Air flow and/or AHU temperature sensor may require recalibration',
        'the sensors are within calibration'],

    [7, ['sat','satsp','htg'],
        'applies to AHU heating mode only, and attempts to verify an AHU heating or cooling valve is not stuck or leaking by verifying AHU supply temperature to supply temperature setpoint',
        'the AHU heating valve maybe broken or there could be a flow issue with the amount of hot water flowing through the coil or that the boiler system reset is too aggressive and there isnt enough heat being produced by this coil. It could be worth viewing mechanical blue prints for this AHU design schedule to see what hot water temperature this coil was designed for and compare it to actual hot water supply temperatures. IE., an AHU hot water coil sized to have a 180°F water flowing through it may have a durastic reduction in performance the colder the hot water is flowing through it, if need be consult a mechanical design engineer to rectify',
        'the AHU heating valve is operating appropriately'
        ],

    [8, ['sat','mat','economizer_sig'], 
        'is an AHU economizer free cooling mode only with an attempt at flagging conditions when the AHU mixing air temperature the supply air temperature are not approximately equal',
        'temperature sensor error or the heating/cooling coils are leaking potentially creating simultenious heating/cooling which can be an energy penalty for running the AHU in this fashion. Verify AHU mix/supply temperature sensor calibration in addition to a potential mechanical issue of a leaking valve. A leaking valve can be troubleshot by isolating the valve closed by manual shut off valves where piping lines enter the AHU coil and then verifying any changes in the AHU discharge air temperature',
        'the AHU components are within calibration for this fault equation'],

    [9, ['satsp','oat'], 
        'is an AHU economizer free cooling mode only with an attempt at flagging conditions where the outside air temperature is too warm for cooling without additional mechanical cooling',
        'temperature sensor error or the cooling valve is stuck open or leaking causing overcooling. Trouble shoot a leaking valve by isolating the coil with manual shutoff valves and verify a change in AHU discharge air temperature with the AHU running',
        'the AHU components are within calibration for this fault equation'
        ],

    [10, ['oat','mat','clg','economizer_sig'], 
        'is an AHU economizer + mechanical cooling mode only with an attempt at flagging conditions where the outside air temperature and mixing air temperatures are not approximetely equal when the AHU is in a 100% outside air mode',
        'temperature sensor error or the mixing air dampers are stuck or broken with the inability for the AHU to go into a proper 100 percent outside air mode. If the outside air temperature is a global variable on the BAS verify (IE, installed to the boiler plant controller and then shared via supervisory level logic on the BAS to the AHU controllers on the BAS network) that where the actual OA temperature is installed that is on the North side of the building in the shade. On the AHU verify mix temperature sensor calibration and that the mixing dampers have good proper rotation with good seals when in the closed position. When testing AHU systems operating in a 100 percent outside air mode it could be worth verifying exhaust systems or return fans are operating properly. In thoery if alot of air is being pumped into the building and it is allowed to be exhaust or relieved properly, a balanced building will not have any issues of closing or opening egress doors to the building due to excess positive building pressure',
        'the AHU components are within calibration for this fault equation'
        ],

    [11, ['satsp','oat','clg','economizer_sig'], 
        'is an AHU economizer + mechanical cooling mode only with an attempt at flagging conditions where the outside air temperature is too low for 100% outside air AHU operating mode',
        'temperature sensor error or the heating coil could be leaking potentially creating simultenious heating/cooling scenorio which can be an energy penalty for running the AHU in this fashion. Also visually verify with the AHU off via lock-out-tag-out that the mixing dampers operates effectively. To do this have one person the BAS sending operator override commands to drive the damper back and forth. The other person should put on eyes on the operation of the actuator motor driving the OA dampers 100 percent open and then closed and visually verify the dampers rotate effectively per BAS command where to also visually verify the dampers have a good seal when in the closed position. Also consider looking into BAS programming that may need tuning or parameter adjustments for the staging between OS state changes between AHU modes of operation',
        'the AHU components are within calibration for this fault equation'
        ],

    [12, ['sat','mat','clg','economizer_sig'], 
        'is an AHU economizer + mechanical cooling mode and AHU mechanical cooling mode only with an attempt at flagging conditions when the AHU mixing air temperature is warmer than the supply air temperature',
        'temperature sensor error or the heating/cooling coils are leaking potentially creating simultenious heating/cooling which can be an energy penalty for running the AHU in this fashion. Verify AHU mix/supply temperature sensor calibration in addition to a potential mechanical issue of a leaking valve. A leaking valve can be troubleshot by isolating the valve closed by manual shut off valves where piping lines enter the AHU coil and then verifying any changes in the AHU discharge air temperature',
        'the AHU components are within calibration for this fault equation'],

    [13, ['sat','satsp','clg','economizer_sig'], 
        'is an AHU cooling mode only with an attempt at verifying an AHU cooling valve is not stuck or leaking by verifying AHU supply temperature to supply temperature setpoint',
        'the AHU cooling valve maybe broken or there could be a flow issue with the amount of cold water flowing through the coil or that the chiller system leaving temperature reset is too aggressive and there isnt enough cold air being produced by this cooling coil. If this AHU has a DX cooling coil there could be a problem with the refrigerant charge. It could be worth viewing mechanical blue prints for this AHU design schedule to see what cold water temperature this coil was designed for and compare it to actual cold water supply temperatures. IE., an AHU cooling coil sized to have a 44°F water flowing through it may have significant performance reduction with 48°F water flowing through it and under design day type high load conditions this AHU may not meet setpoint or properly dehumidify the air for the building which could potentially also lead to IAQ or mold issues if %RH levels in the zones are kept within tollerance. Also check excessive outside air faults in fault condition 6 that the AHU isnt taking in too much outdoor air which could also cause coil performance issues if the load on the coil is higher than what it was intended for',
        'the AHU cooling valve operates as expected']
]

class Report:
    def __init__(self, fault_num: int, df: pd.DataFrame, path: str, op_mode_cols: list = None):

        fault_def = next((fault_def for fault_def in fault_defs if fault_def[0] == fault_num))

        self.fault = Fault(**dict(zip(fault_attrs, fault_def))) 

        self.calculator = Calculator(self.fault, df)
        self.document_generator = DocumentGenerator(self.fault, df, self.calculator)

        self.document = self.document_generator.create_document(path)

        self.path = path

    def save_report(self):
        self.document.save(self.path)


# to run an example of this, run the following:

# ADJUST this param for the AHU MIN OA damper stp
AHU_MIN_OA = 20

# G36 params shouldnt need adjusting
# error threshold parameters
OAT_DEGF_ERR_THRES = 5
MAT_DEGF_ERR_THRES = 5

from faults import FaultConditionTen

_fc10 = FaultConditionTen(
    OAT_DEGF_ERR_THRES,
    MAT_DEGF_ERR_THRES,
    "mat",
    "oat",
    "clg",
    "economizer_sig",
)

FAULT_NUM = 10

df = pd.read_csv(f'fc{FAULT_NUM}_data.csv')

path = os.path.join(os.path.curdir, "final_report")
if not os.path.exists(path):
    os.makedirs(path)

report = Report(FAULT_NUM, df2, os.path.join(path, f"test_fc{FAULT_NUM}.docx"))
report.save_report()