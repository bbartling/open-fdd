# air_handling_unit/reports/report_fc2.py
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches
from air_handling_unit.reports.base_report import BaseReport
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import time

class FaultCodeTwoReport(BaseReport):
    def __init__(self, config):
        super().__init__(config)
        self.mix_degf_err_thres = config['MIX_DEGF_ERR_THRES']
        self.return_degf_err_thres = config['RETURN_DEGF_ERR_THRES']
        self.outdoor_degf_err_thres = config['OUTDOOR_DEGF_ERR_THRES']
        self.mat_col = config['MAT_COL']
        self.rat_col = config['RAT_COL']
        self.oat_col = config['OAT_COL']
        self.supply_vfd_speed_col = config['SUPPLY_VFD_SPEED_COL']

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:
        if output_col is None:
            output_col = "fc2_flag"

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(25, 8))
        plt.title('Fault Conditions 2 Plot')

        ax1.plot(df.index, df[self.mat_col], color='r', label="Mix Temp")  # red
        ax1.plot(df.index, df[self.rat_col], color='b', label="Return Temp")  # blue
        ax1.plot(df.index, df[self.oat_col], color='g', label="Out Temp")  # green
        ax1.legend(loc='best')
        ax1.set_ylabel("°F")

        ax2.plot(df.index, df[output_col], label="Fault", color="k")
        ax2.set_xlabel('Date')
        ax2.set_ylabel('Fault Flags')
        ax2.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def create_report(self, path: str, df: pd.DataFrame, output_col: str = "fc2_flag", report_name: str = "report_fc2.docx") -> None:
        document = Document()
        document.add_heading("Fault Condition Two Report", 0)

        p = document.add_paragraph(
            """Fault condition two and three of ASHRAE Guideline 36 is related to flagging mixing air temperatures of the AHU that are out of acceptable ranges. Fault condition 2 flags mixing air temperatures that are too low and fault condition 3 flags mixing temperatures that are too high when in comparison to return and outside air data. The mixing air temperatures in theory should always be in between the return and outside air temperatures ranges. Fault condition two equation as defined by ASHRAE:"""
        )

        # Correcting the path to the image
        image_path = os.path.join(os.path.dirname(__file__), '..', 'images', 'fc2_definition.png')
        document.add_picture(image_path, width=Inches(6))

        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)
        document.add_picture(fan_plot_image, width=Inches(6))

        summary = self.summarize_fault_times(df, output_col)
        self.add_summary_to_document(document, summary)

        hours_motor_runtime = summary['hours_motor_runtime']
        total_hours = summary['total_hours']
        flag_true_mat = round(
            df[self.mat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_oat = round(
            df[self.oat_col].where(df[output_col] == 1).mean(), 2
        )
        flag_true_rat = round(
            df[self.rat_col].where(df[output_col] == 1).mean(), 2
        )

        if int(total_hours) == int(hours_motor_runtime):
            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"This fan system appears to run 24/7 consider implementing occupancy schedules to reduce building fuel use through HVAC"
            )

        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()
            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'When fault condition 2 is True the average mix air temp is {flag_true_mat}°F, outside air temp is {flag_true_oat}°F, and return air temp is {flag_true_rat}°F. This could possibly help with pin pointing AHU operating conditions for when this fault is True.'
            )

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")
            paragraph = document.add_paragraph()
            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset.')

        document.add_heading('Summary Statistics filtered for when the AHU is running', level=1)

        df_motor_on_filtered = df[df[self.supply_vfd_speed_col] > 0.1]

        document.add_heading('Mix Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.mat_col].describe()))

        document.add_heading('Return Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.rat_col].describe()))

        document.add_heading('Outside Temp', level=3)
        paragraph = document.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(str(df_motor_on_filtered[self.oat_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        percent_true = summary['percent_true']
        if percent_true > 5:
            paragraph.add_run(
                'The percent True of time in fault condition 2 is high indicating the AHU temperature temp sensors are out of calibration')
        else:
            paragraph.add_run(
                'The percent True of time is low indicating the AHU temperature sensors are within calibration')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        document.save(f"{path}/{report_name}")
