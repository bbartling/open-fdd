import os
from io import BytesIO
from docx import Document
from docx.shared import Inches
from air_handling_unit.reports.base_report import BaseReport
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import time


class FaultCodeOneReport(BaseReport):
    def __init__(self, config):
        super().__init__(config)
        self.vfd_speed_percent_err_thres = config['VFD_SPEED_PERCENT_ERR_THRES']
        self.duct_static_col = config['DUCT_STATIC_COL']
        self.supply_vfd_speed_col = config['SUPPLY_VFD_SPEED_COL']
        self.duct_static_setpoint_col = config['DUCT_STATIC_SETPOINT_COL']

    def create_fan_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:
        if output_col is None:
            output_col = "fc1_flag"

        fig, (ax1, ax2, ax3) = plt.subplots(3, 1, figsize=(25, 8))
        plt.title('Fault Conditions 1 Plot')

        ax1.plot(df.index, df[self.duct_static_col], label="STATIC")
        ax1.legend(loc='best')
        ax1.set_ylabel("Inch WC")

        ax2.plot(df.index, df[self.supply_vfd_speed_col],
                 color="g", label="FAN")
        ax2.legend(loc='best')
        ax2.set_ylabel('%')

        ax3.plot(df.index, df[output_col], label="Fault", color="k")
        ax3.set_xlabel('Date')
        ax3.set_ylabel('Fault Flags')
        ax3.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def create_report(self, path: str, df: pd.DataFrame, output_col: str = "fc1_flag", report_name: str = "report_fc1.docx") -> None:
        document = Document()
        document.add_heading("Fault Condition One Report", 0)

        p = document.add_paragraph(
            """Fault condition one is related to flagging poor performance of a AHU variable supply fan attempting to control to a duct pressure setpoint."""
        )

        # Correcting the path to the image
        image_path = os.path.join(os.path.dirname(__file__), '..', 'images', 'fc1_definition.png')
        document.add_picture(image_path, width=Inches(6))

        # Add bullet points under the math equation
        bullet_points = [
            "DSP: Duct Static Pressure",
            "DSPSP: Duct Static Pressure Setpoint",
            "VFDSPD: VFD Speed Reference in Percent",
            "eVFDSPD: VFD Speed Reference Error Threshold"
        ]

        for point in bullet_points:
            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(point)

        document.add_heading("Dataset Plot", level=2)

        fig = self.create_fan_plot(df, output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)
        document.add_picture(fan_plot_image, width=Inches(6))

        summary = self.summarize_fault_times(df, output_col)
        self.add_summary_to_document(document, summary)

        hours_motor_runtime = summary['hours_motor_runtime']
        total_hours = summary['total_hours']
        flag_true_duct_static = round(
            df[self.duct_static_col].where(df[output_col] == 1).mean(), 2
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Calculated motor runtime in hours based off of VFD signal > zero: {hours_motor_runtime}"
        )

        if int(total_hours) == int(hours_motor_runtime):
            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"This fan system appears to run 24/7 consider implementing occupancy schedules to reduce building fuel use through HVAC"
            )

        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f'Average duct system pressure for when in fault condition (fan VFD speed > 95%): {flag_true_duct_static}"WC'
            )

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")
            paragraph = document.add_paragraph()
            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset for the equation defined by ASHRAE.')

        document.add_heading(
            'Summary Statistics filtered for when the AHU is running', level=1)

        df_motor_on_filtered = df[df[self.supply_vfd_speed_col] > 0.1]

        document.add_heading("VFD Speed", level=3)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            str(df_motor_on_filtered[self.supply_vfd_speed_col].describe()))

        document.add_heading("Duct Pressure", level=3)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            str(df_motor_on_filtered[self.duct_static_col].describe()))

        document.add_heading("Duct Pressure Setpoint", level=3)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            str(df_motor_on_filtered[self.duct_static_setpoint_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        percent_true = summary['percent_true']
        if percent_true > 5.0:
            paragraph.add_run(
                "The percent True metric that represents the amount of time for when the fault flag is True is high indicating the fan is running at high speeds and appearing to not generate good duct static pressure"
            )
        else:
            paragraph.add_run(
                "The percent True metric that represents the amount of time for when the fault flag is True is low indicating the fan appears to generate good duct static pressure"
            )

        if df[self.duct_static_setpoint_col].std() < 0.1:
            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                "No duct static pressure setpoint reset detected consider implementing a reset strategy to save AHU fan energy"
            )
        else:
            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run("Duct pressure reset detected (Good)")

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        document.save(f"{path}/{report_name}")
