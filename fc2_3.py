import operator
import time
from datetime import datetime, timedelta

import pandas as pd
import numpy as np

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from docx import Document
from docx.shared import Inches
import argparse

# python 3.10 on Windows 10
# py .\fc2_3.py -i ./ahu_data/hvac_random_fake_data/fc2_3_fake_data1.csv -o 1_ahu_fc2_3_report

parser = argparse.ArgumentParser(add_help=False)
args = parser.add_argument_group('Options')

args.add_argument('-h', '--help', action='help',
                  help='Show this help message and exit.')
args.add_argument('-i', '--input', required=True, type=str,
                  help='CSV File Input')
args.add_argument('-o', '--output', required=True, type=str,
                  help='Word File Output Name')
'''
args.add_argument('--use-flask', default=False, action='store_true')
args.add_argument('--no-flask', dest='use-flask', action='store_false')
'''
args = parser.parse_args()

# required params taken from the screenshot above
OUTDOOR_DEGF_ERR_THRES = .05
MIX_DEGF_ERR_THRES = .05
RETURN_DEGF_ERR_THRES = .05


def fault_condition_two_(dataframe):
    return ((dataframe.mat + dataframe.mix_degf_err_thres) < np.minimum((dataframe.rat - dataframe.return_degf_err_thres),
                                                                        (dataframe.oat - dataframe.outdoor_degf_err_thres)))


def fault_condition_three_(dataframe):
    return ((dataframe.mat - dataframe.mix_degf_err_thres) > np.maximum((dataframe.rat + dataframe.return_degf_err_thres),
                                                                        (dataframe.oat + dataframe.outdoor_degf_err_thres)))


'''
mat = pd.read_csv(
    './ahu_data/MA-T.csv',
    index_col='Date',
    parse_dates=True).fillna(method='ffill').fillna(method='bfill').dropna()
#print(mat)
mat = mat.rolling('5T').mean()

rat = pd.read_csv(
    './ahu_data/RA-T.csv',
    index_col='Date',
    parse_dates=True).fillna(method='ffill').fillna(method='bfill').dropna()
#print(rat)
vfd_speed_avg = rat.rolling('5T').mean()

oat = pd.read_csv(
    './ahu_data/OA-T.csv',
    index_col='Date',
    parse_dates=True).fillna(method='ffill').fillna(method='bfill').dropna()
#print(oat)
vfd_speed_avg = rat.rolling('5T').mean()

mat_rat = mat.join(rat)
df = mat_rat.join(oat)
'''

df = pd.read_csv(args.input,
                 index_col='Date',
                 parse_dates=True).rolling('5T').mean()

# make an entire column out of these params in the Pandas Dataframe
df['outdoor_degf_err_thres'] = OUTDOOR_DEGF_ERR_THRES
df['mix_degf_err_thres'] = MIX_DEGF_ERR_THRES
df['return_degf_err_thres'] = RETURN_DEGF_ERR_THRES

start = df.head(1).index.date
print('Dataset start: ', start)

end = df.tail(1).index.date
print('Dataset end: ', end)
print('COLUMNS: ', print(df.columns))

df['fc2_flag'] = fault_condition_two_(df)
df['fc3_flag'] = fault_condition_three_(df)

df2 = df.copy().dropna()
df2['fc2_flag'] = df2['fc2_flag'].astype(int)
df2['fc3_flag'] = df2['fc3_flag'].astype(int)

# drop params column for better plot
df2 = df2.drop(['outdoor_degf_err_thres',
                'mix_degf_err_thres',
                'return_degf_err_thres'], axis=1)

print(df2)

fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(25, 8))
plt.title('Fault Conditions 2 and 3 Plots')

plot1a, = ax1.plot(df2.index, df2.mat, color='r')  # red
plot1b, = ax1.plot(df2.index, df2.rat, color='b')  # blue
plot1c, = ax1.plot(df2.index, df2.oat, color='g')  # green
ax1.set_ylabel('AHU Temp Sensors')

ax2.plot(df2.index, df2.fc2_flag, color='c')  # cyan
ax2.plot(df2.index, df2.fc3_flag, color='m')  # purple
ax2.set_xlabel('Date')
ax2.set_ylabel('Fault Flags')

red_patch = mpatches.Patch(color='red', label='MAT')
blue_patch = mpatches.Patch(color='blue', label='RAT')
green_patch = mpatches.Patch(color='green', label='OAT')
cyan_patch = mpatches.Patch(color='cyan', label='fc2_flag')
purple_patch = mpatches.Patch(color='purple', label='fc3_flag')
plt.legend(handles=[red_patch, blue_patch,
           green_patch, cyan_patch, purple_patch])
plt.tight_layout()
plt.savefig('./static/ahu_fc2_fans_plot.png')
# plt.show()

print("Starting ahu fc2 docx report")
document = Document()
document.add_heading('Fault Condition Two and Three Report', 0)

p = document.add_paragraph(
    'Fault condition two and three of ASHRAE Guideline 36 is related to flagging mixing air temperatures of the AHU that are out of acceptable ranges. Fault condition 2 flags mixing air temperatures that are too low and fault condition 3 flags mixing temperatures that are too high when in comparision to return and outside air data. The mixing air temperatures in theory should always be in between the return and outside air temperatures ranges. Fault condition two equation as defined by ASHRAE:')
document.add_picture('./images/fc2_definition.png', width=Inches(6))

p = document.add_paragraph(
    'Fault condition three equation as defined by ASHRAE:')
document.add_picture('./images/fc3_definition.png', width=Inches(6))
document.add_heading('Dataset Plot', level=2)

# ADD IN SUBPLOTS SECTION
document.add_picture('./static/ahu_fc2_fans_plot.png', width=Inches(6))
document.add_heading('Dataset Statistics', level=2)

# calculate dataset statistics
df2["timedelta_alldata"] = df2.index.to_series().diff()
seconds_alldata = df2.timedelta_alldata.sum().seconds
days_alldata = df2.timedelta_alldata.sum().days

hours_alldata = round(seconds_alldata/3600, 2)
minutes_alldata = round((seconds_alldata/60) % 60, 2)
total_hours_calc = days_alldata * 24.0 + hours_alldata

# fc2 stats for histogram plot
df2["timedelta_fddflag_fc2"] = df2.index.to_series(
).diff().where(df2["fc2_flag"] == 1)
seconds_fc2_mode = df2.timedelta_fddflag_fc2.sum().seconds
hours_fc2_mode = round(seconds_fc2_mode/3600, 2)

percent_true_fc2 = round(df2.fc2_flag.mean() * 100, 2)
percent_false_fc2 = round((100 - percent_true_fc2), 2)

df2['hour_of_the_day_fc2'] = df2.index.hour.where(df2["fc2_flag"] == 1)
flag_true_fc2 = round(
    df2.mat.where(df2["fc2_flag"] == 1).mean(), 2)

# fc3 stats for histogram plot
df2["timedelta_fddflag_fc3"] = df2.index.to_series(
).diff().where(df2["fc3_flag"] == 1)
seconds_fc3_mode = df2.timedelta_fddflag_fc3.sum().seconds
hours_fc3_mode = round(seconds_fc3_mode/3600, 2)

percent_true_fc3 = round(df2.fc3_flag.mean() * 100, 2)
percent_false_fc3 = round((100 - percent_true_fc3), 2)

df2['hour_of_the_day_fc3'] = df2.index.hour.where(df2["fc3_flag"] == 1)
flag_true_fc3 = round(
    df2.mat.where(df2["fc3_flag"] == 1).mean(), 2)

print('UNIQUE DF2 HOUR OF DAY: ', df2.hour_of_the_day_fc2.unique())
print('UNIQUE DF3 HOUR OF DAY: ', df2.hour_of_the_day_fc3.unique())

# make hist plots fc3
fig, axs = plt.subplots(1, 2, sharey=True, tight_layout=True, figsize=(25, 8))
axs[0].hist(df2.hour_of_the_day_fc2.dropna())
axs[1].hist(df2.hour_of_the_day_fc3.dropna())

fc = 2
for ax in axs:
    ax.set_xlabel('24 Hour Number in Day')
    ax.set_ylabel('Frequency')
    ax.set_title(f'Hour-Of-Day When Fault Flag {fc} is TRUE')
    fc += 1
fig.savefig('./static/ahu_fc23_histogram.png')

# add calcs to word doc
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time calculated in dataset: {df2.timedelta_alldata.sum()}')
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time in hours calculated in dataset: {total_hours_calc}')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time in hours for when fault flag 2 is True: {hours_fc2_mode}')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time in hours for when fault flag 3 is True: {hours_fc3_mode}')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Percent of time in the dataset when the fault flag 2 is True: {percent_true_fc2}%')
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Percent of time in the dataset when the fault flag 3 is True: {percent_true_fc3}%')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Percent of time in the dataset when fault flag 2 is False: {percent_false_fc2}%')
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Percent of time in the dataset when fault flag 3 is False: {percent_false_fc3}%')

paragraph = document.add_paragraph()
# ADD HIST Plots
document.add_heading('Time-of-day Histogram Plots', level=2)
document.add_picture('./static/ahu_fc23_histogram.png', width=Inches(6))

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Average mix air temp for when in fault condition 2 is True (mixing temp is LOW outside the ranges of return and outside temp): {flag_true_fc2} °F')
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Average mix air temp for when in fault condition 3 is True (mixing temp is HIGH and outside the ranges of return and outside temp): {flag_true_fc3} °F')
paragraph = document.add_paragraph()


# ADD in Summary Statistics
document.add_heading('Mix Temp Statistics', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(str(df2.mat.describe()))

# ADD in Summary Statistics
document.add_heading('Return Temp Statistics', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(str(df2.rat.describe()))

# ADD in Summary Statistics
document.add_heading('Outside Temp Statistics', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(str(df2.oat.describe()))


document.add_heading('Suggestions based on data analysis', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'

if percent_true_fc2 < 5 and percent_true_fc3 < 5:

    paragraph.add_run(
        'The percent True of time is very high indicating the AHU temperature sensors are out of calibration')

else:
    paragraph.add_run(
        'The percent True of time is low inidicating the AHU temperature sensors are within calibration')

print('df2.mat.std: ', df2.mat.std())
print('df2.mat.min: ', df2.mat.min())
print('df2.mat.max: ', df2.mat.max())

paragraph = document.add_paragraph()
run = paragraph.add_run(f'Report generated: {time.ctime()}')
run.style = 'Emphasis'

document.save(f'./final_report/{args.output}.docx')
print('All Done')

# df2.to_csv('test1.csv')
