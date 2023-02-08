import operator
import time
from datetime import datetime, timedelta

import pandas as pd
import numpy as np

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from docx import Document
from docx.shared import Inches

import argparse

# python 3.10 on Windows 10
# py .\fc4.py -i ./ahu_data/hvac_random_fake_data/fc4_fake_data1.csv -o fake1_ahu_fc4_report

parser = argparse.ArgumentParser(add_help=False)
args = parser.add_argument_group('Options')

args.add_argument('-h', '--help', action='help', help='Show this help message and exit.')
args.add_argument('-i', '--input', required=True, type=str,
                    help='CSV File Input')
args.add_argument('-o', '--output', required=True, type=str,
                    help='Word File Output Name')
'''
FUTURE 
 * incorporate an arg for SI units 
 * °C on temp sensors
 * piping pressure sensor PSI conversion
 * air flow CFM conversion
 * AHU duct static pressure "WC

args.add_argument('--use-SI-units', default=False, action='store_true')
args.add_argument('--no-SI-units', dest='use-SI-units', action='store_false')
'''
args = parser.parse_args()

# required params taken from the screenshot above
AHU_MIN_OA = 20
DELTA_OS_MAX = 7.


df = pd.read_csv(args.input,
    index_col='Date',
    parse_dates=True).rolling('5T').mean()

print(df)

# make an entire column out of these params in the Pandas Dataframe
df['delta_os_max'] = AHU_MIN_OA
df['ahu_min_oa'] = DELTA_OS_MAX

start = df.head(1).index.date
print('Dataset start: ', start)

end = df.tail(1).index.date
print('Dataset end: ', end)
print('COLUMNS: ', print(df.columns))


df['heating_mode'] = df['heating_sig'].gt(0.)
df['econ_mode'] = df['economizer_sig'].gt(df['ahu_min_oa']) & df['cooling_sig'].eq(0.)
df['econ_cooling_mode'] = df['economizer_sig'].gt(df['ahu_min_oa']) & df['cooling_sig'].gt(0.)
df['mech_cooling_mode'] = df['economizer_sig'].eq(df['ahu_min_oa']) & df['cooling_sig'].gt(0.)


df_to_plot = df[['heating_mode','econ_mode','econ_cooling_mode','mech_cooling_mode',
                 'heating_sig', 'economizer_sig', 'cooling_sig']]
df = df[['heating_mode','econ_mode','econ_cooling_mode','mech_cooling_mode']]

df = df.astype(int)
df_to_plot.heating_mode = df_to_plot.heating_mode.astype(int)
df_to_plot.econ_mode = df_to_plot.econ_mode.astype(int)
df_to_plot.econ_cooling_mode = df_to_plot.econ_cooling_mode.astype(int)
df_to_plot.mech_cooling_mode = df_to_plot.mech_cooling_mode.astype(int)

# calc changes per hour for modes
# https://stackoverflow.com/questions/69979832/pandas-consecutive-boolean-event-rollup-time-series
df = df.resample('H').apply(lambda x: (x.eq(1) & x.shift().ne(1)).sum())

df_debug =  df.copy()

df['fc4_flag'] = df[df.columns].gt(DELTA_OS_MAX).any(1)
df = df.astype(int)

print('df.fc4_flag.describe')
print(df.fc4_flag.describe())

flag4_max_val = df.fc4_flag.max()
print('df.fc4_flag.max = ',flag4_max_val)

fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(25, 8))
plt.title('Fault Conditions 4 Plots')

plot1a, = ax1.plot(df_to_plot.index, df_to_plot.heating_sig, color='r')  # red
plot1b, = ax1.plot(df_to_plot.index, df_to_plot.cooling_sig, color='b')  # blue
plot1c, = ax1.plot(df_to_plot.index, df_to_plot.economizer_sig, color='g')  # green
ax1.set_ylabel('AHU Output Signals In %')

ax2.plot(df_to_plot.index, df_to_plot.heating_mode, color='orange')  # orange
ax2.plot(df_to_plot.index, df_to_plot.econ_mode, color='olive')  # olive
ax2.plot(df_to_plot.index, df_to_plot.econ_cooling_mode, color='c')  # cyan
ax2.plot(df_to_plot.index, df_to_plot.mech_cooling_mode, color='m')  # black

ax2.set_xlabel('Date')
ax2.set_ylabel('Calculated AHU Operating States')

red_patch = mpatches.Patch(color='red', label='Heating Signal')
blue_patch = mpatches.Patch(color='blue', label='Cooling Signal')
green_patch = mpatches.Patch(color='green', label='Economizer Signal')
orange_patch = mpatches.Patch(color='orange', label='Heating Mode')
olive_patch = mpatches.Patch(color='olive', label='Econ Mode')
cyan_patch = mpatches.Patch(color='cyan', label='Econ + Mech Cooling Mode')
black_patch = mpatches.Patch(color='black', label='Mech Cooling Mode')
plt.legend(handles=[red_patch, blue_patch, green_patch,
           orange_patch, olive_patch, cyan_patch, black_patch])
plt.tight_layout()
plt.savefig('./static/ahu_fc4_oper_states.png')


print("Starting ahu fc4 docx report")
document = Document()
document.add_heading('Fault Condition Four Report', 0)

p = document.add_paragraph(
    'Fault condition four of ASHRAE Guideline 36 is related to flagging AHU control programming that is hunting between heating, economizing, economizing plus mechanical cooling, and mechanical cooling operating states. This fault diagnostic does NOT flag simultaneous heating and cooling, just excessive cycling between the states or operating modes the AHU maybe going in and out of. Fault condition four equation as defined by ASHRAE:')
document.add_picture('./images/fc4_definition.png', width=Inches(6))

document.add_heading('Calculated Operating States Plot', level=2)
document.add_picture('./static/ahu_fc4_oper_states.png', width=Inches(6))

document.add_heading('Dataset Statistics', level=2)

# calculate dataset statistics
delta_all_data = df.index.to_series().diff()
total_days_all_data = round(delta_all_data.sum() / pd.Timedelta(days=1),2)
print('DAYS ALL DATA: ',total_days_all_data)
total_hours_all_data = delta_all_data.sum() / pd.Timedelta(hours=1)
print('HOURS ALL DATA: ',total_hours_all_data)

# heating mode runtime stats
delta_heating = df.heating_mode.index.to_series().diff()
total_hours_heating = (delta_heating * df["heating_mode"]).sum() / pd.Timedelta(hours=1)
print('HOURS HEATING MODE: ',total_hours_heating)
percent_heating = round(df.heating_mode.mean() * 100, 2)
print('PERCENT TIME HEATING MODE: ',percent_heating,'%')

# econ mode runtime stats
delta_econ = df.econ_mode.index.to_series().diff()
total_hours_econ = (delta_econ * df["econ_mode"]).sum() / pd.Timedelta(hours=1)
print('HOURS ECON MODE: ',total_hours_econ)
percent_econ = round(df.econ_mode.mean() * 100, 2)
print('PERCENT TIME ECON MODE: ',percent_econ,'%')

# econ plus mech cooling mode runtime stats
delta_econ_clg = df.econ_cooling_mode.index.to_series().diff()
total_hours_econ_clg = (delta_econ_clg * df["econ_cooling_mode"]).sum() / pd.Timedelta(hours=1)
print('HOURS ECON AND MECH CLG MODE: ',total_hours_econ_clg)
percent_econ_clg = round(df.econ_cooling_mode.mean() * 100, 2)
print('PERCENT TIME ECON AND MECH CLG MODE: ',percent_econ_clg,'%')

# mech clg mode runtime stats
delta_clg = df.mech_cooling_mode.index.to_series().diff()
total_hours_clg = (delta_clg * df["mech_cooling_mode"]).sum() / pd.Timedelta(hours=1)
print('HOURS MECH CLG MODE: ',total_hours_clg)
percent_clg = round(df.mech_cooling_mode.mean() * 100, 2)
print('PERCENT TIME MECH CLG MODE: ',percent_clg,'%')

# add calcs to word doc
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time in days calculated in dataset: {total_days_all_data}')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time in hours calculated in dataset: {total_hours_all_data}')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time in hours while AHU is in a heating mode: {total_hours_heating}')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total percent time in while AHU is in a heating mode: {percent_heating}%')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time in hours while AHU is in a economizing mode: {total_hours_econ}')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total percent time in while AHU is in a economizing mode: {percent_econ}%')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time in hours while AHU is in a economizing plus mechanical cooling mode: {total_hours_econ_clg}')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total percent time in while AHU is in a economizing plus mechanical cooling mode: {percent_econ_clg}%')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total time in hours while AHU is in a mechanical cooling mode: {total_hours_clg}')

paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(
    f'Total percent time in while AHU is in a mechanical cooling mode: {percent_clg}%')

# skip calculating statistics if there is no fault 4's
# Pandas calcs alot of NaNs and errors out
if flag4_max_val != 0:
    
    # fc4 stats for histogram plot
    df["timedelta_fddflag_fc4"] = df.index.to_series().diff().where(df["fc4_flag"] == 1)
    percent_true_fc4 = round(df.fc4_flag.mean() * 100, 2)

    percent_false_fc4 = round((100 - percent_true_fc4), 2)
    df['hour_of_the_day_fc4'] = df.index.hour.where(df["fc4_flag"] == 1)

    # flag_true_fc4
    flag_true_heating_mode = round(
        df.heating_mode.where(df["fc4_flag"] == 1).mean(), 2)
    flag_true_econ_mode = round(
        df.econ_mode.where(df["fc4_flag"] == 1).mean(), 2)
    flag_true_econ_cooling_mode = round(
        df.econ_cooling_mode.where(df["fc4_flag"] == 1).mean(), 2)
    flag_true_mech_cooling_mode = round(
        df.mech_cooling_mode.where(df["fc4_flag"] == 1).mean(), 2)


    # make hist plots
    df['hour_of_the_day'] = df.index.hour.where(df["fc4_flag"] == 1)

    print('DF HOUR OF DAY')
    print(df.hour_of_the_day)
    print(df.hour_of_the_day.describe())

    # make hist plots fc3
    fig, ax = plt.subplots(tight_layout=True, figsize=(25,8))
    ax.hist(df.hour_of_the_day)
    ax.set_xlabel('24 Hour Number in Day')
    ax.set_ylabel('Frequency')
    ax.set_title(f'Hour-Of-Day When Fault Flag 4 is TRUE')
    fig.savefig('./static/ahu_fc4_histogram.png')

    paragraph = document.add_paragraph()
    paragraph.style = 'List Bullet'
    paragraph.add_run(
        f'Total time for when fault flag 4 is True: {df.timedelta_fddflag_fc4.sum()}')

    paragraph = document.add_paragraph()
    paragraph.style = 'List Bullet'
    paragraph.add_run(
        f'Percent of time in the dataset when the fault flag 4 is True: {percent_true_fc4}%')

    paragraph = document.add_paragraph()
    paragraph.style = 'List Bullet'
    paragraph.add_run(
        f'Percent of time in the dataset when fault flag 4 is False: {percent_false_fc4}%')

    paragraph = document.add_paragraph()
    # ADD HIST Plots
    document.add_heading('Time-of-day Histogram Plots', level=2)
    document.add_picture('./static/ahu_fc4_histogram.png', width=Inches(6))

    paragraph = document.add_paragraph()
    paragraph.style = 'List Bullet'
    paragraph.add_run(
        f'Control system in a heating mode while fault condition 4 is True: {flag_true_heating_mode} °F')
    paragraph = document.add_paragraph()

    paragraph = document.add_paragraph()
    paragraph.style = 'List Bullet'
    paragraph.add_run(
        f'Control system in a free economizer mode while fault condition 4 is True: {flag_true_econ_mode} °F')
    paragraph = document.add_paragraph()

    paragraph = document.add_paragraph()
    paragraph.style = 'List Bullet'
    paragraph.add_run(
        f'Control system in a economizer and mechanical cooling mode while fault condition 4 is True: {flag_true_econ_cooling_mode} °F')
    paragraph = document.add_paragraph()

    paragraph = document.add_paragraph()
    paragraph.style = 'List Bullet'
    paragraph.add_run(
        f'Control system in a mechanical cooling mode while fault condition 4 is True: {flag_true_mech_cooling_mode} °F')
    paragraph = document.add_paragraph()


else:
    paragraph = document.add_paragraph()
    paragraph.style = 'List Bullet'
    paragraph.add_run(
        f'No fault condition 4 calcualted in the dataset')
    paragraph = document.add_paragraph()



# ADD in Summary Statistics
document.add_heading('Heating Signal Statistics', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(str(df_to_plot.heating_sig.describe()))

# ADD in Summary Statistics
document.add_heading('Cooling Signal Statistics', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(str(df_to_plot.cooling_sig.describe()))

# ADD in Summary Statistics
document.add_heading('Economizer Free Cooling Statistics', level=2)
paragraph = document.add_paragraph()
paragraph.style = 'List Bullet'
paragraph.add_run(str(df_to_plot.economizer_sig.describe()))


paragraph = document.add_paragraph()
run = paragraph.add_run(f'Report generated: {time.ctime()}')
run.style = 'Emphasis'

document.save(f'./final_report/{args.output}.docx')
print('All Done WITH: ',args.output)
