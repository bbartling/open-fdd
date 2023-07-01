import os
import time
from io import BytesIO
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches
import openai


class FaultCodeThreeReport:
    """Class provides the definitions for Fault Code 3 Report."""

    def __init__(
        self,
        mix_degf_err_thres: float,
        return_degf_err_thres: float,
        outdoor_degf_err_thres: float,
        mat_col: str,
        rat_col: str,
        oat_col: str,
        fan_vfd_speed_col: str,
        api_key: str,
    ):
        self.mix_degf_err_thres = mix_degf_err_thres
        self.return_degf_err_thres = return_degf_err_thres
        self.outdoor_degf_err_thres = outdoor_degf_err_thres
        self.mat_col = mat_col
        self.rat_col = rat_col
        self.oat_col = oat_col
        self.fan_vfd_speed_col = fan_vfd_speed_col
        self.api_key = api_key
        openai.api_key = self.api_key
        self.max_tokens = 3000
        self.completion_model = "gpt-3.5-turbo"

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:
        if output_col is None:
            output_col = "fc3_flag"

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(25, 8))
        plt.title("Fault Conditions 3 Plot")

        (plot1a,) = ax1.plot(
            df.index, df[self.mat_col], color="r", label="Mix Temp"
        )  # red

        (plot1b,) = ax1.plot(
            df.index, df[self.rat_col], color="b", label="Return Temp"
        )  # blue

        (plot1c,) = ax1.plot(
            df.index, df[self.oat_col], color="g", label="Out Temp"
        )  # green

        ax1.legend(loc="best")
        ax1.set_ylabel("°F")

        ax2.plot(df.index, df[output_col], label="Fault", color="k")
        ax2.set_xlabel("Date")
        ax2.set_ylabel("Fault Flags")
        ax2.legend(loc="best")

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc3_flag"
        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)
        # print("DAYS ALL DATA: ", total_days)
        total_hours = delta.sum() / pd.Timedelta(hours=1)
        # print("TOTAL HOURS: ", total_hours)
        hours_fc3_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)
        # print("FALT FLAG TRUE TOTAL HOURS: ", hours_fc1_mode)
        percent_true = round(df[output_col].mean() * 100, 2)
        # print("PERCENT TIME WHEN FLAG IS TRUE: ", percent_true, "%")
        percent_false = round((100 - percent_true), 2)
        # print("PERCENT TIME WHEN FLAG 5 FALSE: ", percent_false, "%")

        flag_true_mat = round(df[self.mat_col].where(df[output_col] == 1).mean(), 2)
        flag_true_oat = round(df[self.oat_col].where(df[output_col] == 1).mean(), 2)
        flag_true_rat = round(df[self.rat_col].where(df[output_col] == 1).mean(), 2)

        motor_on = df[self.fan_vfd_speed_col].gt(0.01).astype(int)
        hours_motor_runtime = round((delta * motor_on).sum() / pd.Timedelta(hours=1), 2)

        # for summary stats on I/O data to make useful
        df_motor_on_filtered = df[df[self.fan_vfd_speed_col] > 0.1]

        return (
            total_days,
            total_hours,
            hours_fc3_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_oat,
            flag_true_rat,
            hours_motor_runtime,
            df_motor_on_filtered,
        )

    def create_hist_plot(
        self, df: pd.DataFrame, output_col: str = None, mat_col: str = None
    ) -> plt:
        if output_col is None:
            output_col = "fc3_flag"

        if mat_col is None:
            mat_col = "mat"

        # calculate dataset statistics
        df["hour_of_the_day_fc3"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc3
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc3.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 3 is TRUE")
        return fig

    def create_report(
        self, path: str, df: pd.DataFrame, output_col: str = None
    ) -> None:
        if output_col is None:
            output_col = "fc3_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Three Report", 0)

        p = document.add_paragraph(
            """Fault condition two and three of ASHRAE Guideline 36 is related to flagging mixing air temperatures of the AHU that are out of acceptable ranges. Fault condition 2 flags mixing air temperatures that are too low and fault condition 3 flags mixing temperatures that are too high when in comparision to return and outside air data. The mixing air temperatures in theory should always be in between the return and outside air temperatures ranges. Fault condition three equation as defined by ASHRAE:"""
        )

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc3_definition.png"),
            width=Inches(6),
        )
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc3_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_oat,
            flag_true_rat,
            hours_motor_runtime,
            df_motor_on_filtered,
        ) = self.summarize_fault_times(df, output_col=output_col)

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc3_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Calculated motor runtime in hours based off of VFD signal > zero: {hours_motor_runtime}"
        )

        if int(total_hours) == int(hours_motor_runtime):
            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"This fan system appears to run 24/7 consider implementing occupancy schedules to reduce building fuel use through HVAC"
            )

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:
            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()

            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"When fault condition 3 is True the average mix air temp is {flag_true_mat}°F, outside air temp is {flag_true_oat}°F, and return air temp is {flag_true_rat}°F. This could possibly help with pin pointing AHU operating conditions for when this fault is True."
            )

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"No faults were found in this given dataset for the equation defined by ASHRAE."
            )

        document.add_heading(
            "Summary Statistics filtered for when the AHU is running", level=1
        )

        # ADD in Summary Statistics
        document.add_heading("Mix Temp", level=3)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(str(df_motor_on_filtered[self.mat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading("Return Temp", level=3)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(str(df_motor_on_filtered[self.rat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading("Outside Temp", level=3)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(str(df_motor_on_filtered[self.oat_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        if percent_true > 5:
            paragraph.add_run(
                "The percent True of time in fault condition 3 is high indicating the AHU temp sensors are out of calibration"
            )

        else:
            paragraph.add_run(
                "The percent True of time is low inidicating the AHU temperature sensors are within calibration"
            )

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document
