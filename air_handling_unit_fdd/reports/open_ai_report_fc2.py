import os
import time
from io import BytesIO
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches
import openai
from docx.shared import Pt


class FaultCodeTwoReport:
    """Class provides the definitions for Fault Code 2 Report."""

    def __init__(self, dict_, api_key):
        attributes_dict = {
            'mix_degf_err_thres': float,
            'return_degf_err_thres': float,
            'outdoor_degf_err_thres': float,
            'mat_col': str,
            'rat_col': str,
            'oat_col': str,
            'supply_vfd_speed_col': str,
        }
        for attribute in attributes_dict:
            upper = attribute.upper()
            value = dict_[upper]
            self.__setattr__(upper, value)
        openai.api_key = api_key
        self.max_tokens = 3000
        self.completion_model = "gpt-3.5-turbo"

    def get_completion(self, messages):
        response = openai.ChatCompletion.create(
            model=self.completion_model, messages=messages, temperature=0
        )
        return response.choices[0].message["content"]

    def generate_insights(
        self,
        total_days,
        total_hours,
        hours_fc2_mode,
        percent_true,
        percent_false,
        flag_true_mat,
        flag_true_oat,
        flag_true_rat,
        hours_motor_runtime,
        mat_col_describe,
        rat_col_describe,
        oat_col_describe,
    ):
        insights_prompt = """
        The fault detection dataset tables an air handling unit (AHU) temperature sensors for the return, mixing, and outdoor air. 
        The fault for this dataset called fc2_flag was generated in the dataset when the fan is running mixing air temperature
        is too low indicating sensor error as the mixing temp should be between the outside air and return air sensor values. 
        Analyzing the data allows us to identify use patterns in how the air handling unit operates over time as well as any potential mechanical issues.

        total days of data: {total_days}
        total hours of data: {total_hours}
        percent of time when fault is true: {percent_true}
        percent of time when fault is false: {percent_false} 
        calculated fan motor run time: {hours_motor_runtime} 
        
        calculated time in hours for when the fault condition is True: {hours_fc2_mode}
        averaged mixing air temperature for when the fault is true: {flag_true_mat}
        averaged outside air temperature for when the fault is true: {flag_true_oat}
        averaged return air temperature for when the fault is true: {flag_true_rat}

        summary statistics of mixing air temperature filtered for when the fan system is running: {mat_col_describe}
        summary statistics of return air temperature filtered for when the fan system is running: {rat_col_describe}
        summary statistics of outside air temperature filtered for when the fan system is running: {oat_col_describe}
        
        Provide a paragraph of written insights into how a virtual AI-powered analyst that specializes in the HVAC systems operations might 
        describe the air handling unit operations in as well as trends provided in the faults calculated and summary statistics of the sensor data. Assume the reader
        does not know anything about HVAC and needs to be described how a variable fan systems operate at a beginner level of understanding.
        Describe that thermaldynamics dictates that the mixing air temperature should always be between the return and outside air as inside the air handling unit
        these two air streams are blended together. The mixing air temperature in theory could always be calculated given the volume of return air and outdoor quantateties
        at whatever temperatures are provided, in fact design engineers do this in design to prevent mixing air temperature that may be too low in cold weather operations
        and cause mechanical damage such as a frozen heating or cooling coil. Breifly describe in written words that based on the data the AHU appears to average a mixing temperature
        a return air temperature and given the outside air temperature conditions given in the dataset.
        If the faults are high, recommend this AHU may have mechanical issues of not being duct static pressure setpoint else if there is
        good data and faults are low, recommend this AHU appears to operate fine meeting duct static pressure requirements. If total_hours is approximately equal
        to the hours_motor_runtime, recommend a schedule placed on the AHU fan could save electrical energy consumption.
        """

        messages = [
            {"role": "system", "content": "You are an AI powered HVAC specialist."},
            {
                "role": "user",
                "content": insights_prompt.format(
                    total_days=total_days,
                    total_hours=total_hours,
                    hours_fc2_mode=hours_fc2_mode,
                    percent_true=percent_true,
                    percent_false=percent_false,
                    flag_true_mat=flag_true_mat,
                    flag_true_oat=flag_true_oat,
                    flag_true_rat=flag_true_rat,
                    hours_motor_runtime=hours_motor_runtime,
                    mat_col_describe=mat_col_describe,
                    rat_col_describe=rat_col_describe,
                    oat_col_describe=oat_col_describe,
                ),
            },
        ]

        response = self.get_completion(messages)
        start_index = response.find("AI:")
        insights = response[start_index + 4 :].strip()

        print(insights)
        return insights

    def create_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:
        if output_col is None:
            output_col = "fc2_flag"

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(25, 8))
        plt.title("Fault Conditions 2 Plot")

        (plot1a,) = ax1.plot(
            df.index, df[self.mat_col], color="r", label="Mix Temp"
        )  # red

        (plot1b,) = ax1.plot(
            df.index, df[self.rat_col], color="b", label="Return Temp"
        )  # blue

        (plot1c,) = ax1.plot(
            df.index, df[self.oat_col], color="g", label="Out Temp"
        )  # green

        ax1.legend(loc="best")
        ax1.set_ylabel("°F")

        ax2.plot(df.index, df[output_col], label="Fault", color="k")
        ax2.set_xlabel("Date")
        ax2.set_ylabel("Fault Flags")
        ax2.legend(loc="best")

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc2_flag"
        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)
        # print("DAYS ALL DATA: ", total_days)
        total_hours = delta.sum() / pd.Timedelta(hours=1)
        # print("TOTAL HOURS: ", total_hours)
        hours_fc2_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)
        # print("FALT FLAG TRUE TOTAL HOURS: ", hours_fc1_mode)
        percent_true = round(df[output_col].mean() * 100, 2)
        # print("PERCENT TIME WHEN FLAG IS TRUE: ", percent_true, "%")
        percent_false = round((100 - percent_true), 2)
        # print("PERCENT TIME WHEN FLAG 5 FALSE: ", percent_false, "%")

        flag_true_mat = round(df[self.mat_col].where(df[output_col] == 1).mean(), 2)
        flag_true_oat = round(df[self.oat_col].where(df[output_col] == 1).mean(), 2)
        flag_true_rat = round(df[self.rat_col].where(df[output_col] == 1).mean(), 2)
        motor_on = df[self.fan_vfd_speed_col].gt(0.01).astype(int)
        hours_motor_runtime = round((delta * motor_on).sum() / pd.Timedelta(hours=1), 2)

        # for summary stats on I/O data to make useful
        df_motor_on_filtered = df[df[self.fan_vfd_speed_col] > 0.1]
        return (
            total_days,
            total_hours,
            hours_fc2_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_oat,
            flag_true_rat,
            hours_motor_runtime,
            df_motor_on_filtered,
        )

    def create_hist_plot(
        self, df: pd.DataFrame, output_col: str = None, mat_col: str = None
    ) -> plt:
        if output_col is None:
            output_col = "fc2_flag"

        if mat_col is None:
            mat_col = "mat"

        # calculate dataset statistics
        df["hour_of_the_day_fc2"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc3
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc2.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 2 is TRUE")
        return fig

    def create_report(
        self, path: str, df: pd.DataFrame, output_col: str = None
    ) -> None:
        if output_col is None:
            output_col = "fc2_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition Two Report", 0)

        paragraph = document.add_paragraph()
        run = paragraph.add_run("Mix air temperature too low; should be between outside and return.")
        run.bold = True
        run.font.size = Pt(14)

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc2_definition.png"),
            width=Inches(6),
        )
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(f"MATavg: Mixing Air Temperature Rolling Average")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(f"eMAT: Mixing Air Temperature Error Threshold")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(f"RATavg: Return Air Temperature Rolling Average")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(f"eRAT: Return Air Temperature Error Threshold")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(f"OATavg: Outside Air Temperature Rolling Average")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(f"eOAT: Outside Air Temperature Error Threshold")

        document.add_heading("Dataset Plot", level=2)

        fig = self.create_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc2_mode,
            percent_true,
            percent_false,
            flag_true_mat,
            flag_true_oat,
            flag_true_rat,
            hours_motor_runtime,
            df_motor_on_filtered,
        ) = self.summarize_fault_times(df, output_col=output_col)

        mat_col_describe = str(df_motor_on_filtered[self.mat_col].describe())
        rat_col_describe = str(df_motor_on_filtered[self.rat_col].describe())
        oat_col_describe = str(df_motor_on_filtered[self.oat_col].describe())

        insights = self.generate_insights(
            total_days=total_days,
            total_hours=total_hours,
            hours_fc2_mode=hours_fc2_mode,
            percent_true=percent_true,
            percent_false=percent_false,
            flag_true_mat=flag_true_mat,
            flag_true_oat=flag_true_oat,
            flag_true_rat=flag_true_rat,
            hours_motor_runtime=hours_motor_runtime,
            mat_col_describe=mat_col_describe,
            rat_col_describe=rat_col_describe,
            oat_col_describe=oat_col_describe,
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc2_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Calculated motor runtime in hours based off of VFD signal > zero: {hours_motor_runtime}"
        )

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:
            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f"When fault condition 2 is True the average mix air temp is {flag_true_mat}°F, outside air temp is {flag_true_oat}°F, and return air temp is {flag_true_rat}°F. This could possibly help with pin pointing AHU operating conditions for when this fault is True."
            )

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = "List Bullet"
            paragraph.add_run(f"No faults were found in this given dataset.")

        document.add_heading(
            "Summary Statistics filtered for when the AHU is running", level=1
        )

        # ADD in Summary Statistics
        document.add_heading("Mix Temp", level=3)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(str(df_motor_on_filtered[self.mat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading("Return Temp", level=3)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(str(df_motor_on_filtered[self.rat_col].describe()))

        # ADD in Summary Statistics
        document.add_heading("Outside Temp", level=3)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(str(df_motor_on_filtered[self.oat_col].describe()))

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        try:
            paragraph.add_run(f"{insights}")
        except Exception as e:
            paragraph.add_run(f"Failed to connect to AI service for suggestions - {e}.")

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return document
