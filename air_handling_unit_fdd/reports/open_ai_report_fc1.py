import os
import time
from io import BytesIO
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches
import openai
from docx.shared import Pt


class FaultCodeOneReport:
    """Class provides the definitions for Fault Code 1 Report."""

    def __init__(self, dict_, api_key):
        """Passes dictionary into initialization of class instance, then uses the attributes called out below in
        attributes_dict to set only the attributes that match from dict_.

        :param dict_: dictionary of all possible class attributes (loaded from config file)
        """
        self.vfd_speed_percent_err_thres = float,
        self.vfd_speed_percent_max = float,
        self.duct_static_inches_err_thres = float,
        self.duct_static_col = str,
        self.supply_vfd_speed_col = str,
        self.duct_static_setpoint_col = str,

        for attribute in self.__dict__:
            upper = attribute.upper()
            value = dict_[upper]
            self.__setattr__(attribute, value)

        openai.api_key = api_key
        self.max_tokens = 3000
        self.completion_model = 'gpt-3.5-turbo'

    def get_completion(self, messages):
        response = openai.ChatCompletion.create(
            model=self.completion_model,
            messages=messages,
            temperature=0
        )
        return response.choices[0].message['content']

    def generate_insights(self,
                          total_days,
                          total_hours,
                          hours_fc1_mode,
                          percent_true,
                          percent_false,
                          flag_true_duct_static,
                          hours_motor_runtime,
                          fan_vfd_speed_describe,
                          duct_static_describe,
                          duct_static_setpoint_describe
                          ):
        insights_prompt = '''
        The fault detection dataset tables an air handling unit (AHU) supply fan and duct static pressure. The fault was generated in the dataset when the fan is
        running near 100 percent speed and the duct static pressure in the duct system is not meeting setpoint. 
        Analyzing the data allows us to identify use patterns in how the air handling unit operates over time as well as any potential mechanical issues.

        total days of data: {total_days}
        total hours of data: {total_hours}
        percent of time when fault is true: {percent_true}
        percent of time when fault is false: {percent_false}
        averaged duct static pressure when the fault is true: {flag_true_duct_static}
        calculated fan motor run time: {hours_motor_runtime}
        summary statistics of fan speed in percent data: {fan_vfd_speed_describe}
        summary statistics of duct static pressure in engineering units: {duct_static_describe}
        summary statistics of duct static pressure setpoint in engineering units: {duct_static_setpoint_describe}

        Provide a paragraph of written insights into how a virtual AI-powered analyst that specializes in the HVAC systems operations might 
        describe the air handling unit operations in as well as trends provided in the faults calculated and summary statistics of the sensor data. Assume the reader
        does not know anything about HVAC and needs to be described how a variable fan controls to a duct static pressure sensor at beginner level of understanding.
        If the faults are high, recommend this AHU may have mechanical issues of not being duct static pressure setpoint else if there is
        good data and faults are low, recommend this AHU appears to operate fine meeting duct static pressure requirements. If total_hours is approximately equal
        to the hours_motor_runtime, recommend a schedule placed on the AHU fan could save electrical energy consumption. If the duct_static_setpoint_col
        summary statistics reveal the duct static setpoint not changing, recommend a duct static pressure reset to save electrical energy
        from the fan motor consumption.
        '''

        messages = [
            {"role": "system", "content": "You are an AI powered HVAC specialist."},
            {"role": "user", "content": insights_prompt.format(
                total_days=total_days,
                total_hours=total_hours,
                hours_fc1_mode=hours_fc1_mode,
                percent_true=percent_true,
                percent_false=percent_false,
                flag_true_duct_static=flag_true_duct_static,
                hours_motor_runtime=hours_motor_runtime,
                fan_vfd_speed_describe=fan_vfd_speed_describe,
                duct_static_describe=duct_static_describe,
                duct_static_setpoint_describe=duct_static_setpoint_describe
            )}
        ]

        response = self.get_completion(messages)
        start_index = response.find('AI:')
        insights = response[start_index + 4:].strip()

        print(insights)
        return insights

    def create_fan_plot(self, df: pd.DataFrame, output_col: str = None) -> plt:
        if output_col is None:
            output_col = "fc1_flag"

        fig, (ax1, ax2, ax3) = plt.subplots(3, 1, figsize=(25, 8))
        plt.title('Fault Conditions 1 Plot')

        ax1.plot(df.index, df[self.duct_static_col], label="STATIC")
        ax1.legend(loc='best')
        ax1.set_ylabel("Inch WC")

        ax2.plot(df.index, df[self.supply_vfd_speed_col],
                 color="g", label="FAN")
        ax2.legend(loc='best')
        ax2.set_ylabel('%')

        ax3.plot(df.index, df[output_col], label="Fault", color="k")
        ax3.set_xlabel('Date')
        ax3.set_ylabel('Fault Flags')
        ax3.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    def summarize_fault_times(self, df: pd.DataFrame, output_col: str = None) -> str:
        if output_col is None:
            output_col = "fc1_flag"
        delta = df.index.to_series().diff()
        total_days = round(delta.sum() / pd.Timedelta(days=1), 2)

        total_hours = delta.sum() / pd.Timedelta(hours=1)

        hours_fc1_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)

        percent_true = round(df[output_col].mean() * 100, 2)

        percent_false = round((100 - percent_true), 2)

        flag_true_duct_static = round(
            df[self.duct_static_col].where(df[output_col] == 1).mean(), 2
        )

        motor_on = df[self.supply_vfd_speed_col].gt(.01).astype(int)
        hours_motor_runtime = round(
            (delta * motor_on).sum() / pd.Timedelta(hours=1), 2)

        # for summary stats on I/O data to make useful
        df_motor_on_filtered = df[df[self.supply_vfd_speed_col] > 0.1]

        return (
            total_days,
            total_hours,
            hours_fc1_mode,
            percent_true,
            percent_false,
            flag_true_duct_static,
            hours_motor_runtime,
            df_motor_on_filtered
        )

    def create_hist_plot(
        self, df: pd.DataFrame, output_col: str = None, duct_static_col: str = None
    ) -> plt:
        if output_col is None:
            output_col = "fc1_flag"
        if duct_static_col is None:
            duct_static_col = "duct_static"
        # calculate dataset statistics
        df["hour_of_the_day_fc1"] = df.index.hour.where(df[output_col] == 1)

        # make hist plots fc3
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(df.hour_of_the_day_fc1.dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag 1 is TRUE")
        return fig

    def create_report(
        self,
        path: str,
        df: pd.DataFrame,
        output_col: str = None
    ) -> None:

        if output_col is None:
            output_col = "fc1_flag"

        print(f"Starting {path} docx report!")
        document = Document()
        document.add_heading("Fault Condition One Report", 0)
        
        paragraph = document.add_paragraph()
        run = paragraph.add_run("Duct static pressure too low with fan at full speed.")
        run.bold = True
        run.font.size = Pt(14)

        document.add_picture(
            os.path.join(os.path.curdir, "images", "fc1_definition.png"),
            width=Inches(6),
        )
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"DSP: Duct Static Pressure")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"DPSP: Duct Static Pressure Setpoint")
        
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"VFDSPD: VFD Speed Reference in Percent")
        
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"eVFDSPD: VFD Speed Reference Error Threshold")
        
        document.add_heading("Dataset Plot", level=2)

        fig = self.create_fan_plot(df, output_col=output_col)
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)
        # ADD IN SUBPLOTS SECTION
        document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )
        document.add_heading("Dataset Statistics", level=2)

        (
            total_days,
            total_hours,
            hours_fc1_mode,
            percent_true,
            percent_false,
            flag_true_duct_static,
            hours_motor_runtime,
            df_motor_on_filtered
        ) = self.summarize_fault_times(df, output_col=output_col)

        fan_vfd_speed_describe = str(df_motor_on_filtered[self.supply_vfd_speed_col].describe())
        duct_static_describe = str(df_motor_on_filtered[self.duct_static_col].describe())
        duct_static_setpoint_describe = str(df_motor_on_filtered[self.duct_static_setpoint_col].describe())

        insights = self.generate_insights(
            total_days=total_days,
            total_hours=total_hours,
            hours_fc1_mode=hours_fc1_mode,
            percent_true=percent_true,
            percent_false=percent_false,
            flag_true_duct_static=flag_true_duct_static,
            hours_motor_runtime=hours_motor_runtime,
            fan_vfd_speed_describe=fan_vfd_speed_describe,
            duct_static_describe=duct_static_describe,
            duct_static_setpoint_describe=duct_static_setpoint_describe
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in days calculated in dataset: {total_days}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours calculated in dataset: {total_hours}")

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Total time in hours for when fault flag is True: {hours_fc1_mode}"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is True: {percent_true}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Percent of time in the dataset when the fault flag is False: {percent_false}%"
        )

        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            f"Calculated motor runtime in hours based off of VFD signal > zero: {hours_motor_runtime}"
        )

        paragraph = document.add_paragraph()

        # if there is no faults skip the histogram plot
        fc_max_faults_found = df[output_col].max()
        if fc_max_faults_found != 0:

            # ADD HIST Plots
            document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot(df, output_col=output_col)
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(
                f'Average duct system pressure for when in fault condition (fan VFD speed > 95%): {flag_true_duct_static}"WC'
            )

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")

            paragraph.style = 'List Bullet'
            paragraph.add_run(
                f'No faults were found in this given dataset.')

        document.add_heading(
            'Summary Statistics filtered for when the AHU is running', level=1)

        # ADD in Summary Statistics of fan operation
        document.add_heading("VFD Speed", level=3)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            fan_vfd_speed_describe
        )

        # ADD in Summary Statistics of duct pressure
        document.add_heading("Duct Pressure", level=3)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            duct_static_describe
        )

        # ADD in Summary Statistics of duct pressure
        document.add_heading("Duct Pressure Setpoint", level=3)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"
        paragraph.add_run(
            duct_static_setpoint_describe
        )

        document.add_heading("Suggestions based on data analysis", level=2)
        paragraph = document.add_paragraph()
        paragraph.style = "List Bullet"

        try:
            paragraph.add_run(
                f"{insights}"
            )
        except Exception as e:
            paragraph.add_run(
                f"Failed to connect to AI service for suggestions - {e}."
            )

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        print("REPORT SUCCESS FC1")
        return document